"""app.py 9:02 PM, 9/7/25
"""

import random
import time
from flask import Flask, render_template, url_for, request, session, redirect, abort, jsonify, flash
from database import mongo
from werkzeug.utils import secure_filename
import os, re, logging, sys
import spacy, fitz, io
from bson.objectid import ObjectId
from google.oauth2 import id_token
from google_auth_oauthlib.flow import Flow
from pip._vendor import cachecontrol
import google.auth.transport.requests
import pathlib
import requests
from werkzeug.security import generate_password_hash, check_password_hash
from flask_pymongo import PyMongo
from datetime import datetime, timedelta
import shutil
# from resume_utils import GeneralResumeProcessor
import warnings
from flask import make_response  # Add this if not already there
import csv
from io import StringIO

import threading
# Add this import at the top
import re
from collections import defaultdict

# Add this import at the top of app.py
import secrets
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
import smtplib
from PIL import Image
import uuid
from flask import Flask, send_file, request
from transformers import pipeline
# from transformers import AutoTokenizer, AutoModelForTokenClassification, pipeline
# import torch
from dotenv import load_dotenv
load_dotenv()
import numpy as np
import docx
from spacy.tokens import Doc
from resume_data_cleaner import clean_resume_data
import logging
from typing import Dict, List, Any, Optional
from sentence_transformer_ranker import (
    get_semantic_ranker, 
    enhance_application_with_semantic_score,
    get_enhanced_candidate_ranking,
    
)

from resume_utils import GeneralResumeProcessor
from resume_processor import (
    IntelligentResumeProcessor, 
    ResumeTextExtractor,
    process_resume_with_nlp
)


from notifications import (
    create_notification,
    create_application_notification,
    create_interview_notification,
    create_status_update_notification,
    get_user_notifications,
    get_unread_count,
    mark_notification_read,
    mark_all_notifications_read
)



from flask import request, jsonify, send_file, render_template
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
from docx.oxml import parse_xml




warnings.filterwarnings("ignore", message=".*was trained with spaCy.*", category=UserWarning)

# Also suppress the dotenv warning
warnings.filterwarnings("ignore", message=".*could not parse statement.*", category=UserWarning)



# Force UTF-8 encoding for console output on Windows
if sys.platform.startswith('win'):
    # Set environment variable for Python to use UTF-8
    os.environ['PYTHONIOENCODING'] = 'utf-8'
    
    # Configure logging with UTF-8 encoding
    logging.basicConfig(
        level=logging.INFO,
        format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
        handlers=[
            logging.StreamHandler(sys.stdout)
        ]
    )
    
    # Ensure stdout uses UTF-8
    if hasattr(sys.stdout, 'reconfigure'):
        sys.stdout.reconfigure(encoding='utf-8')
    
else:
    # Standard logging for non-Windows systems
    logging.basicConfig(
        level=logging.INFO,
        format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
    )





# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('app.log'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)





def allowedExtension(filename):
    """Validate resume file extension"""
    if not filename or '.' not in filename:
        return False
    return filename.rsplit('.', 1)[1].lower() in ['docx', 'pdf']

def allowedExtensionPdf(filename):
    """Validate PDF file extension"""
    if not filename or '.' not in filename:
        return False
    return filename.rsplit('.', 1)[1].lower() in ['pdf']

def validate_email(email):
    """Validate email format"""
    pattern = r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$'
    return re.match(pattern, email) is not None

def validate_password(password):
    """Validate password strength"""
    if len(password) < 6:
        return False, "Password must be at least 6 characters long"
    if not re.search(r'[A-Za-z]', password):
        return False, "Password must contain at least one letter"
    if not re.search(r'\d', password):
        return False, "Password must contain at least one number"
    return True, "Valid password"

def cleanup_user_files(user_id, keep_file=None):
    """Clean up old resume files for a user"""
    try:
        user_resumes = resumeFetchedData.find({"UserId": ObjectId(user_id)})
        for resume in user_resumes:
            if resume.get('ResumeTitle') and resume['ResumeTitle'] != keep_file:
                old_path = os.path.join(app.config['UPLOAD_FOLDER'], resume['ResumeTitle'])
                if os.path.exists(old_path):
                    os.remove(old_path)
                    logger.info(f"Cleaned up old resume file: {old_path}")
    except Exception as e:
        logger.error(f"Error cleaning up files for user {user_id}: {e}")

app = Flask(__name__)





# Environment variable configuration
app.secret_key = os.environ.get("SECRET_KEY", "Resume_screening_default_key_change_in_production")

# OAuth Configuration
os.environ["OAUTHLIB_INSECURE_TRANSPORT"] = "1"
GOOGLE_CLIENT_ID = os.environ.get("GOOGLE_CLIENT_ID", "114980943385-4rsv0t3ck6878gf4j6mtd244i5o8ti5l.apps.googleusercontent.com")
client_secrets_file = os.path.join(pathlib.Path(__file__).parent, "client_secret.json")

try:
    flow = Flow.from_client_secrets_file(
        client_secrets_file=client_secrets_file,
        scopes=["https://www.googleapis.com/auth/userinfo.profile", 
                "https://www.googleapis.com/auth/userinfo.email", "openid"],
        redirect_uri=os.environ.get("REDIRECT_URI", "http://127.0.0.1:5000/callback")
    )
    logger.info("OAuth flow configured successfully")
except Exception as e:
    logger.error(f"Failed to configure OAuth flow: {e}")
    flow = None

# File upload configuration
UPLOAD_FOLDER = os.environ.get('UPLOAD_FOLDER', 'static/uploaded_resumes')
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size
PROFILE_UPLOAD_FOLDER = os.environ.get('PROFILE_UPLOAD_FOLDER', 'static/profile_pictures')
DOCUMENT_UPLOAD_FOLDER = os.environ.get('DOCUMENT_UPLOAD_FOLDER', 'static/user_documents')
ALLOWED_PROFILE_EXTENSIONS = {'png', 'jpg', 'jpeg', 'gif'}
ALLOWED_DOCUMENT_EXTENSIONS = {'pdf', 'doc', 'docx', 'jpg', 'jpeg', 'png'}

# Ensure upload directory exists
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(PROFILE_UPLOAD_FOLDER, exist_ok=True)
os.makedirs(DOCUMENT_UPLOAD_FOLDER, exist_ok=True)

MONGO_URI = os.environ.get(
    'MONGO_URI',
    'mongodb+srv://admin:admin@cluster0.6qbxdwl.mongodb.net/HireArchy'
)
app.config['MONGO_URI'] = MONGO_URI

# MongoDB configuration
try:
    mongo.init_app(app)

    # Define collections first
    resumeFetchedData = mongo.db.resumeFetchedData
    Applied_EMP = mongo.db.Applied_EMP
    IRS_USERS = mongo.db.IRS_USERS
    JOBS = mongo.db.JOBS


    logger.info("MongoDB initialized successfully")
except Exception as e:
    logger.error(f"Failed to initialize MongoDB: {e}")
    raise


try:
    USER_DOCUMENTS = mongo.db.USER_DOCUMENTS
    logger.info("USER_DOCUMENTS collection initialized")
except Exception as e:
    logger.error(f"Failed to initialize USER_DOCUMENTS collection: {e}")



try:
    NOTIFICATIONS = mongo.db.NOTIFICATIONS
    logger.info("NOTIFICATIONS collection initialized")
except Exception as e:
    logger.error(f"Failed to initialize NOTIFICATIONS collection: {e}")



# # Load SpaCy model
# try:
#     logger.info("Loading Resume Parser model...")
#     nlp = spacy.load("./output/model-best")
#     logger.info("Resume Parser model loaded successfully")
# except Exception as e:
#     logger.error(f"Failed to load SpaCy model: {e}")
#     nlp = None

if not Doc.has_extension("trf_data"):
    Doc.set_extension("trf_data", default=None)

# Load BOTH models
try:
    logger.info("Loading NEW Resume Parser model...")
    nlp_new = spacy.load("./output/model-best") 
    logger.info("New model loaded")
except Exception as e:
    logger.error(f"Failed to load new model: {e}")
    nlp_new = None

try:
    logger.info("Loading OLD Resume Parser model...")
    nlp_old = spacy.load("assets/ResumeModel/output/model-best")
    logger.info("Old model loaded")
except Exception as e:
    logger.error(f"Failed to load old model: {e}")
    nlp_old = None




# try:
#     logger.info("Initializing BERT Resume Extractor...")
#     bert_extractor = create_bert_extractor()
#     if bert_extractor.is_available():
#         logger.info("✅ BERT resume extraction available")
#     else:
#         logger.warning("⚠️ BERT extraction not available - install transformers")
#         bert_extractor = None
# except Exception as e:
#     logger.error(f"Failed to initialize BERT extractor: {e}")
#     bert_extractor = None




def migrate_existing_users():
    """Add new fields to existing user profiles"""
    try:
        # Add ProfilePicture field to users who don't have it
        result = IRS_USERS.update_many(
            {"ProfilePicture": {"$exists": False}},
            {"$set": {"ProfilePicture": None}}
        )
        logger.info(f"Added ProfilePicture field to {result.modified_count} users")
        
        # Add ProfessionalSummary field
        result = IRS_USERS.update_many(
            {"ProfessionalSummary": {"$exists": False}},
            {"$set": {"ProfessionalSummary": None}}
        )
        logger.info(f"Added ProfessionalSummary field to {result.modified_count} users")
        
    except Exception as e:
        logger.error(f"Error in user migration: {e}")

        
def create_indexes():
    """Create database indexes for better performance"""
    try:
        # User documents indexes
        USER_DOCUMENTS.create_index([("UserId", 1)])
        USER_DOCUMENTS.create_index([("Category", 1)])
        USER_DOCUMENTS.create_index([("UploadedAt", -1)])
        USER_DOCUMENTS.create_index([("IsActive", 1)])
        
        # User profile indexes
        IRS_USERS.create_index([("Email", 1)], unique=True)
        IRS_USERS.create_index([("Role", 1)])
        
        logger.info("Database indexes created successfully")
    except Exception as e:
        logger.error(f"Error creating indexes: {e}")

# Call this function after your MongoDB initialization
with app.app_context():
    create_indexes()
    
with app.app_context():
    migrate_existing_users()

with app.app_context():
    # Initialize legitimate AI tracking fields only
    Applied_EMP.update_many(
        {},
        {
            "$set": {
                "semantic_similarity": None,
               
                "enhanced_version": None
            }
        }
    )
    logger.info("Ensured Applied_EMP has legitimate AI tracking fields")





def allowed_profile_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_PROFILE_EXTENSIONS

def allowed_document_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_DOCUMENT_EXTENSIONS


def get_file_icon_class(file_extension):
    """Get CSS class for file type icon"""
    extension = file_extension.lower()
    
    if extension == 'pdf':
        return 'doc-pdf'
    elif extension in ['jpg', 'jpeg', 'png', 'gif']:
        return 'doc-jpg'
    elif extension in ['doc', 'docx']:
        return 'doc-doc'
    else:
        return 'doc-default'

# Add this to your template filters (add to app.py)
@app.template_filter('file_icon')
def file_icon_filter(file_extension):
    return get_file_icon_class(file_extension)





# Register blueprints
from Job_post import APPLY_JOB, job_post
from Job_post import APPLY_JOB_ENHANCED, job_post
app.register_blueprint(job_post, url_prefix="/HR1")


# Updated model loading with fallback chain





# Helper functions for dashboard functionality
def get_time_ago(timestamp):
    """Calculate human-readable time difference"""
    try:
        if not timestamp:
            return "Recently"
        
        now = datetime.now()
        diff = now - timestamp
        
        if diff.days > 30:
            return f"{diff.days // 30} month{'s' if diff.days // 30 > 1 else ''} ago"
        elif diff.days > 0:
            return f"{diff.days} day{'s' if diff.days > 1 else ''} ago"
        elif diff.seconds > 3600:
            hours = diff.seconds // 3600
            return f"{hours} hour{'s' if hours > 1 else ''} ago"
        elif diff.seconds > 60:
            minutes = diff.seconds // 60
            return f"{minutes} minute{'s' if minutes > 1 else ''} ago"
        else:
            return "Just now"
    except Exception:
        return "Recently"


def get_activity_icon(activity_type):
    """Get icon for activity type"""
    icons = {
        'new_application': '📄',
        'interview_scheduled': '📅',
        'candidate_hired': '✅',
        'job_posted': '📝',
        'status_update': '🔄'
    }
    return icons.get(activity_type, '📌')


def get_activity_color(activity_type):
    """Get color for activity type"""
    colors = {
        'new_application': '#3b82f6',
        'interview_scheduled': '#f59e0b',
        'candidate_hired': '#10b981',
        'job_posted': '#8b5cf6',
        'status_update': '#6b7280'
    }
    return colors.get(activity_type, '#6b7280')


def get_activity_title(activity_type):
    """Get title for activity type"""
    titles = {
        'new_application': 'New Application',
        'interview_scheduled': 'Interview Scheduled',
        'candidate_hired': 'Candidate Hired',
        'job_posted': 'Job Posted',
        'status_update': 'Status Updated'
    }
    return titles.get(activity_type, 'Activity')


def get_performance_assessment(percentage):
    """Assess performance percentage"""
    if percentage >= 20:
        return "excellent performance"
    elif percentage >= 15:
        return "above average"
    elif percentage >= 10:
        return "good performance"
    elif percentage >= 5:
        return "average performance"
    else:
        return "room for improvement"
    

@app.route('/profile')
def profile_page():
    """Enhanced profile page with documents and resume management"""
    try:
        if 'user_id' not in session or session.get('role') != 'applicant':
            flash("Access denied. Applicant login required.", "error")
            return redirect(url_for('loginpage'))
        
        user_id = session['user_id']
        
        # Get user profile information
        try:
            user_profile = IRS_USERS.find_one({"_id": ObjectId(user_id)})
            if not user_profile:
                flash("User profile not found", "error")
                return redirect(url_for('loginpage'))
        except Exception as e:
            logger.error(f"Error fetching user profile for {user_id}: {e}")
            user_profile = {}
        
        # Get user documents
        try:
            user_documents = list(mongo.db.USER_DOCUMENTS.find({"UserId": ObjectId(user_id), "Category": {"$ne": "resume"}  }).sort("UploadedAt", -1))
        except Exception as e:
            logger.error(f"Error fetching user documents for {user_id}: {e}")
            user_documents = []
        
        # Get user resumes
        try:
            resumes = list(resumeFetchedData.find({"UserId": ObjectId(user_id)}).sort("UploadedAt", -1))
        except Exception as e:
            logger.error(f"Error fetching resumes for {user_id}: {e}")
            resumes = []
        
        # Calculate profile completion
        try:
            completion_factors = {
                'basic_info': bool(user_profile.get('Name') and user_profile.get('Email')),
                'profile_picture': bool(user_profile.get('ProfilePicture')),
                'contact_info': bool(user_profile.get('PhoneNumber')),
                'location': bool(user_profile.get('Location')),
                'professional_info': bool(user_profile.get('CurrentTitle') and user_profile.get('ExperienceLevel')),
                'summary': bool(user_profile.get('ProfessionalSummary')),
                'linkedin': bool(user_profile.get('LinkedInProfile')),
                'resume_uploaded': bool(resumes),
            }
            
            completed = sum(completion_factors.values())
            profile_completion = round((completed / len(completion_factors)) * 100)
        except Exception as e:
            logger.error(f"Error calculating profile completion: {e}")
            profile_completion = 0
        
        logger.info(f"Profile page loaded for user {user_id}: {len(user_documents)} documents, {len(resumes)} resumes")
        
        return render_template("enhanced_profile.html",
                             user_profile=user_profile,
                             user_documents=user_documents,
                             resumes=resumes,
                             profile_completion=profile_completion)
                             
    except Exception as e:
        logger.error(f"Error in profile_page: {e}")
        return render_template("enhanced_profile.html",
                             user_profile={},
                             user_documents=[],
                             resumes=[],
                             profile_completion=0,
                             errorMsg="An error occurred loading your profile")

@app.route('/update_profile', methods=['POST'])
def update_profile():
    """Update user profile information"""
    try:
        if 'user_id' not in session or session.get('role') != 'applicant':
            return jsonify({"success": False, "message": "Access denied"}), 403
        
        user_id = session['user_id']
        
        # Get form data
        update_data = {
            'Name': request.form.get('full_name', '').strip(),
            'PhoneNumber': request.form.get('phone_number', '').strip(),
            'Location': request.form.get('location', '').strip(),
            'CurrentTitle': request.form.get('current_title', '').strip(),
            'ExperienceLevel': request.form.get('experience_level', '').strip(),
            'IndustryPreference': request.form.get('industry_preference', '').strip(),
            'LinkedInProfile': request.form.get('linkedin_profile', '').strip(),
            'ProfessionalSummary': request.form.get('professional_summary', '').strip(),
            'ProfileUpdatedAt': datetime.now()
        }
        
        # Remove empty values
        update_data = {k: v for k, v in update_data.items() if v}
        
        # Update user profile
        result = IRS_USERS.update_one(
            {"_id": ObjectId(user_id)},
            {"$set": update_data}
        )
        
        if result.modified_count > 0:
            logger.info(f"Profile updated for user {user_id}")
            flash("Profile updated successfully!", "success")
        else:
            flash("No changes made to profile", "info")
            
        return redirect(url_for('profile_page'))
        
    except Exception as e:
        logger.error(f"Error updating profile: {e}")
        flash("Error updating profile", "error")
        return redirect(url_for('profile_page'))

@app.route('/upload_profile_picture', methods=['POST'])
def upload_profile_picture():
    """Upload and update profile picture"""
    try:
        if 'user_id' not in session or session.get('role') != 'applicant':
            return jsonify({"success": False, "message": "Access denied"}), 403
        
        if 'profile_picture' not in request.files:
            return jsonify({"success": False, "message": "No file selected"}), 400
        
        file = request.files['profile_picture']
        
        if file.filename == '':
            return jsonify({"success": False, "message": "No file selected"}), 400
        
        if not allowed_profile_file(file.filename):
            return jsonify({"success": False, "message": "Invalid file type. Only PNG, JPG, JPEG, GIF allowed."}), 400
        
        user_id = session['user_id']
        
        # Generate unique filename
        file_extension = file.filename.rsplit('.', 1)[1].lower()
        filename = f"profile_{user_id}_{uuid.uuid4().hex}.{file_extension}"
        filepath = os.path.join(PROFILE_UPLOAD_FOLDER, filename)
        
        try:
            # Save and resize image
            file.save(filepath)
            
            # Resize image to standard size (400x400)
            with Image.open(filepath) as img:
                # Convert to RGB if necessary
                if img.mode in ('RGBA', 'P'):
                    img = img.convert('RGB')
                
                # Resize maintaining aspect ratio
                img.thumbnail((400, 400), Image.Resampling.LANCZOS)
                
                # Create square image with white background
                if img.size[0] != img.size[1]:
                    size = max(img.size)
                    background = Image.new('RGB', (size, size), (255, 255, 255))
                    background.paste(img, ((size - img.size[0]) // 2, (size - img.size[1]) // 2))
                    img = background
                
                img.save(filepath, 'JPEG', quality=90)
            
            # Remove old profile picture if exists
            old_user = IRS_USERS.find_one({"_id": ObjectId(user_id)}, {"ProfilePicture": 1})
            if old_user and old_user.get('ProfilePicture'):
                old_path = os.path.join('static', old_user['ProfilePicture'].lstrip('/'))
                if os.path.exists(old_path):
                    os.remove(old_path)
            
            # Update user profile
            profile_picture_url = f"/static/profile_pictures/{filename}"
            result = IRS_USERS.update_one(
                {"_id": ObjectId(user_id)},
                {"$set": {"ProfilePicture": profile_picture_url, "ProfileUpdatedAt": datetime.now()}}
            )
            
            if result.modified_count > 0:
                logger.info(f"Profile picture updated for user {user_id}")
                return jsonify({
                    "success": True,
                    "message": "Profile picture updated successfully",
                    "image_url": profile_picture_url
                })
            else:
                return jsonify({"success": False, "message": "Failed to update profile picture"}), 500
                
        except Exception as e:
            logger.error(f"Error processing profile picture: {e}")
            # Clean up file if there was an error
            if os.path.exists(filepath):
                os.remove(filepath)
            return jsonify({"success": False, "message": "Error processing image"}), 500
        
    except Exception as e:
        logger.error(f"Error in upload_profile_picture: {e}")
        return jsonify({"success": False, "message": "Server error"}), 500

@app.route('/upload_document', methods=['POST'])
def upload_document():
    """Upload user document (certification, achievement, etc.)"""
    try:
        if 'user_id' not in session or session.get('role') != 'applicant':
            return jsonify({"success": False, "message": "Access denied"}), 403
        
        if 'document_file' not in request.files:
            return jsonify({"success": False, "message": "No file selected"}), 400
        
        file = request.files['document_file']
        document_name = request.form.get('document_name', '').strip()
        category = request.form.get('category', '').strip()
        description = request.form.get('description', '').strip()
        
        if file.filename == '':
            return jsonify({"success": False, "message": "No file selected"}), 400
        
        if not document_name:
            return jsonify({"success": False, "message": "Document name is required"}), 400
        
        if not category:
            return jsonify({"success": False, "message": "Category is required"}), 400
        
        if not allowed_document_file(file.filename):
            return jsonify({"success": False, "message": "Invalid file type"}), 400
        
        user_id = session['user_id']
        
        # Generate unique filename
        file_extension = file.filename.rsplit('.', 1)[1].lower()
        filename = f"doc_{user_id}_{uuid.uuid4().hex}.{file_extension}"
        filepath = os.path.join(DOCUMENT_UPLOAD_FOLDER, filename)
        
        try:
            # Save file
            file.save(filepath)
            
            # Get file size
            file_size = os.path.getsize(filepath)
            
            # Create document record
            document_data = {
                "UserId": ObjectId(user_id),
                "DocumentName": document_name,
                "FileName": filename,
                "OriginalFileName": file.filename,
                "Category": category,
                "Description": description,
                "FileExtension": file_extension,
                "FileSize": file_size,
                "FilePath": f"/static/user_documents/{filename}",
                "UploadedAt": datetime.now(),
                "IsActive": True
            }
            
            # Insert document record
            result = mongo.db.USER_DOCUMENTS.insert_one(document_data)
            
            if result.inserted_id:
                logger.info(f"Document uploaded for user {user_id}: {document_name} ({category})")
                return jsonify({
                    "success": True,
                    "message": "Document uploaded successfully",
                    "document_id": str(result.inserted_id)
                })
            else:
                # Clean up file if database insert failed
                if os.path.exists(filepath):
                    os.remove(filepath)
                return jsonify({"success": False, "message": "Failed to save document record"}), 500
                
        except Exception as e:
            logger.error(f"Error saving document: {e}")
            # Clean up file if there was an error
            if os.path.exists(filepath):
                os.remove(filepath)
            return jsonify({"success": False, "message": "Error saving document"}), 500
        
    except Exception as e:
        logger.error(f"Error in upload_document: {e}")
        return jsonify({"success": False, "message": "Server error"}), 500

@app.route('/view_document/<document_id>')
def view_document(document_id):
    """View uploaded document"""
    try:
        if 'user_id' not in session or session.get('role') != 'applicant':
            return "Access denied", 403
        
        user_id = session['user_id']
        
        # Find document
        document = mongo.db.USER_DOCUMENTS.find_one({
            "_id": ObjectId(document_id),
            "UserId": ObjectId(user_id)
        })
        
        if not document:
            return "Document not found", 404
        
        # Serve file
        filepath = os.path.join(DOCUMENT_UPLOAD_FOLDER, document['FileName'])
        
        if not os.path.exists(filepath):
            return "File not found", 404
        
        return send_file(filepath, as_attachment=False)
        
    except Exception as e:
        logger.error(f"Error viewing document {document_id}: {e}")
        return "Error viewing document", 500

@app.route('/delete_document/<document_id>', methods=['DELETE'])
def delete_document(document_id):
    """Delete user document"""
    try:
        if 'user_id' not in session or session.get('role') != 'applicant':
            return jsonify({"success": False, "message": "Access denied"}), 403
        
        user_id = session['user_id']
        
        # Find document
        document = mongo.db.USER_DOCUMENTS.find_one({
            "_id": ObjectId(document_id),
            "UserId": ObjectId(user_id)
        })
        
        if not document:
            return jsonify({"success": False, "message": "Document not found"}), 404
        
        # Delete file
        filepath = os.path.join(DOCUMENT_UPLOAD_FOLDER, document['FileName'])
        if os.path.exists(filepath):
            os.remove(filepath)
        
        # Delete database record
        result = mongo.db.USER_DOCUMENTS.delete_one({"_id": ObjectId(document_id)})
        
        if result.deleted_count > 0:
            logger.info(f"Document deleted: {document_id} for user {user_id}")
            return jsonify({"success": True, "message": "Document deleted successfully"})
        else:
            return jsonify({"success": False, "message": "Failed to delete document"}), 500
            
    except Exception as e:
        logger.error(f"Error deleting document {document_id}: {e}")
        return jsonify({"success": False, "message": "Server error"}), 500

@app.route('/get_user_documents')
def get_user_documents():
    """Get user's documents for application selection"""
    try:
        if 'user_id' not in session or session.get('role') != 'applicant':
            return jsonify({"error": "Access denied"}), 403
        
        user_id = session['user_id']
        
        documents = list(mongo.db.USER_DOCUMENTS.find(
            {"UserId": ObjectId(user_id), "IsActive": True},
            {"DocumentName": 1, "Category": 1, "FileName": 1, "FileExtension": 1}
        ))
        
        result = []
        for doc in documents:
            result.append({
                "_id": str(doc['_id']),
                "DocumentName": doc['DocumentName'],
                "Category": doc['Category'],
                "FileExtension": doc.get('FileExtension', ''),
                "FileName": doc['FileName']
            })
        
        return jsonify(result)
        
    except Exception as e:
        logger.error(f"Error fetching user documents: {e}")
        return jsonify({"error": "Server error"}), 500






@app.route('/home')
def home():
    """Render the public home page with features"""
    try:
        return render_template('home.html')
    except Exception as e:
        logger.error(f"Error rendering home page: {e}")
        return "Error loading page", 500

# How It Works page route
@app.route('/how-it-works')
def how_it_works():
    """Render the how it works page"""
    try:
        return render_template('how_it_works.html')
    except Exception as e:
        logger.error(f"Error rendering how it works page: {e}")
        return "Error loading page", 500

# Contact Us page route (GET request - display form)
@app.route('/contact', methods=['GET'])
def contact_get():
    """Render the contact us page"""
    try:
        return render_template('contact_us.html')
    except Exception as e:
        logger.error(f"Error rendering contact page: {e}")
        return "Error loading page", 500

@app.route('/')
def index():
    try:
        return render_template("index.html")
    except Exception as e:
        logger.error(f"Error rendering index page: {e}")
        return "Error loading page", 500

@app.route('/emp')
def emp():
    try:
        if 'user_id' not in session or 'user_name' not in session:
            flash("Please log in first", "error")
            return render_template("index.html", errMsg="Login First")
        
        logger.info(f"User {session['user_id']} accessing employee dashboard")
        
        # Fetch user's resumes with error handling
        try:
            resumes = list(resumeFetchedData.find({"UserId": ObjectId(session['user_id'])}))
            logger.info(f"Found {len(resumes)} resumes for user {session['user_id']}")
        except Exception as e:
            logger.error(f"Error fetching resumes for user {session['user_id']}: {e}")
            resumes = []
            flash("Error loading resumes", "error")
        
        return render_template("applicant_dashboard_enhanced.html", resumes=resumes)
        
    except Exception as e:
        logger.error(f"Error in emp route: {e}")
        return render_template("index.html", errMsg="An error occurred")

@app.route('/login')
def login():
    try:
        if not flow:
            flash("OAuth not configured properly", "error")
            return redirect(url_for('loginpage'))
            
        authorization_url, state = flow.authorization_url()
        session["state"] = state
        logger.info("OAuth login initiated")
        return redirect(authorization_url)
        
    except Exception as e:
        logger.error(f"Error initiating OAuth login: {e}")
        flash("Login service unavailable", "error")
        return redirect(url_for('loginpage'))

@app.route('/loginpage', methods=['GET', 'POST'])
def loginpage():
    """Login page with email verification handling"""
    try:
        if request.method == 'POST':
            email = request.form.get('email', '').strip()
            password = request.form.get('password', '')

            if not email or not password:
                flash('Email and password are required', 'error')
                return render_template('login.html', email=email)
            
            if not validate_email(email):
                flash('Invalid email format', 'error')
                return render_template('login.html', email=email)

            try:
                user = IRS_USERS.find_one({'Email': email})
                
                if user and check_password_hash(user.get('Password', ''), password):
                    # ✅ Check email verification BEFORE setting session
                    if not user.get('EmailVerified', False):
                        flash('Please verify your email before logging in. Check your inbox for the verification link.', 'warning')
                        logger.warning(f"Login attempt with unverified email: {email}")
                        # ✅ Pass email to template to show resend form
                        return render_template('login.html', email=email)
                    
                    # ✅ Email verified - proceed with login
                    session['user_id'] = str(user['_id'])
                    session['user_name'] = user.get('Name', 'User')
                    session['role'] = user.get('Role', 'applicant')
                    session['email'] = email
                    
                    logger.info(f"Successful login: {email} (Role: {session['role']})")
                    
                    # Redirect based on role
                    if user.get('Role') == 'applicant':
                        flash(f'Welcome back, {session["user_name"]}!', 'success')
                        return redirect(url_for('Applicant_Dashboard_Enhanced'))
                    elif user.get('Role') == 'employer':
                        flash(f'Welcome back, {session["user_name"]}!', 'success')
                        return redirect(url_for('HR_Dashboard_Enhanced'))
                    else:
                        flash('Invalid user role', 'error')
                        return render_template('login.html')
                else:
                    logger.warning(f"Failed login attempt for {email}")
                    flash('Invalid email or password', 'error')
                    return render_template('login.html', email='')
                    
            except Exception as e:
                logger.error(f"Database error during login for {email}: {e}")
                import traceback
                logger.error(traceback.format_exc())
                flash('Login service unavailable. Please try again later.', 'error')
                return render_template('login.html', email=email)

        # GET request - check if email parameter exists (from verification redirect)
        email = request.args.get('email', '')
        return render_template('login.html', email=email)
        
    except Exception as e:
        logger.error(f"Error in loginpage route: {e}")
        import traceback
        logger.error(traceback.format_exc())
        return render_template('login.html', errorMsg="An error occurred")


def get_match_score(application):
    """
    Get the best available match score for an application.
    Priority: combined_score > calculated hybrid > traditional
    
    Args:
        application: MongoDB application document
        
    Returns:
        float: Match score (0-100)
    """
    # 1. Try stored combined score (fastest)
    combined = application.get('combined_score')
    if combined is not None:
        try:
            return float(combined)
        except (ValueError, TypeError):
            pass
    
    # 2. Get traditional score
    traditional = application.get('Matching_percentage', 0)
    if isinstance(traditional, dict):
        traditional = traditional.get('overall_score', 0)
    try:
        traditional = float(traditional)
    except (ValueError, TypeError):
        traditional = 0
    
    # 3. Try calculating hybrid with semantic
    semantic = application.get('semantic_similarity')
    if semantic is not None:
        try:
            semantic = float(semantic)
            if semantic > 0:
                from sentence_transformer_ranker import get_semantic_ranker
                ranker = get_semantic_ranker()
                if ranker.is_available():
                    return ranker.calculate_combined_score(traditional, semantic * 100)
        except Exception as e:
            logger.warning(f"Semantic calculation failed: {e}")
    
    # 4. Fallback to traditional
    return traditional

@app.route("/callback")
def callback():
    try:
        if not flow:
            logger.error("OAuth flow not configured")
            return redirect(url_for('index'))
            
        if not session.get("state") or session["state"] != request.args.get("state"):
            logger.warning("OAuth state mismatch")
            abort(500)

        flow.fetch_token(authorization_response=request.url)

        credentials = flow.credentials
        request_session = requests.session()
        cached_session = cachecontrol.CacheControl(request_session)
        token_request = google.auth.transport.requests.Request(session=cached_session)

        id_info = id_token.verify_oauth2_token(
            id_token=credentials._id_token,
            request=token_request,
            audience=GOOGLE_CLIENT_ID
        )
        
        email = id_info.get("email")
        name = id_info.get("name")
        google_id = id_info.get("sub")
        
        if not email:
            logger.error("No email received from Google OAuth")
            flash("Authentication failed", "error")
            return redirect(url_for('index'))

        try:
            result = IRS_USERS.find_one({"Email": email}, {"_id": 1, "Role": 1})
            
            if not result:
                # Create new user
                user_data = {
                    "Name": name,
                    "Email": email,
                    "Google_id": google_id,
                    "Role": "applicant",  # Default role
                    "CreatedAt": datetime.now()
                }
                new_user_id = IRS_USERS.insert_one(user_data).inserted_id
                session['user_id'] = str(new_user_id)
                session['role'] = "applicant"
                logger.info(f"Created new user via OAuth: {email}")
            else:
                session['user_id'] = str(result['_id'])
                session['role'] = result.get('Role', 'applicant')
                logger.info(f"Existing user logged in via OAuth: {email}")
            
            session['user_name'] = name
            # flash("Logged in successfully!", "success")
            return redirect(url_for("Applicant_Dashboard"))
            
        except Exception as e:
            logger.error(f"Database error during OAuth callback for {email}: {e}")
            flash("Authentication failed", "error")
            return redirect(url_for('index'))
            
    except Exception as e:
        logger.error(f"Error in OAuth callback: {e}")
        flash("Authentication failed", "error")
        return redirect(url_for('index'))

@app.route("/logout")
def logout():
    try:
        user_id = session.get('user_id')
        logger.info(f"User {user_id} logging out")
        session.clear()
        flash('You have been logged out successfully', 'success')
        return redirect(url_for("index"))
    except Exception as e:
        logger.error(f"Error during logout: {e}")
        return redirect(url_for("index"))
    

# def generate_ai_insight_badge(app_data):
#     """Convert technical AI scores into user-friendly insights with match score validation"""
    
#     # DEBUG: Log the actual data
#     logger.info(f"DEBUG BADGE - App data: {app_data}")
    
#     # Get scores using the helper function
#     match_score = get_match_score(app_data)  # ✅ This gets combined score
#     semantic_score = app_data.get('semantic_similarity', 0) or 0
#     enhanced_version = app_data.get('enhanced_version')
    
#     # Convert None values to safe defaults
#     if semantic_score is None:
#         semantic_score = 0
#     if match_score is None:
#         match_score = 0
    
#     logger.info(f"DEBUG BADGE - semantic: {semantic_score}, enhanced: {enhanced_version}, match: {match_score}")
    
#     insights = []
    
#     # FIXED: Semantic matching insights with match score validation
#     if semantic_score >= 0.8 and match_score >= 70:  # Both must be good
#         insights.append({
#             'type': 'skill_intelligence',
#             'message': 'Found strong skill relationships',
#             'icon': 'brain',
#             'color': 'success'
#         })
#     elif semantic_score >= 0.6 and match_score >= 60:  # Both moderate
#         insights.append({
#             'type': 'skill_intelligence', 
#             'message': 'Related skills detected',
#             'icon': 'link',
#             'color': 'info'
#         })
#     elif semantic_score >= 0.4 and match_score >= 40:  # Both low but some connection
#         insights.append({
#             'type': 'skill_intelligence',
#             'message': 'Some skill connections found',
#             'icon': 'search',
#             'color': 'warning'
#         })
    
#     # Enhanced processing insight
#     if enhanced_version:
#         insights.append({
#             'type': 'ai_processing',
#             'message': 'Advanced matching applied',
#             'icon': 'cpu',
#             'color': 'primary'
#         })
    
#     # FIXED: Overall quality insight with proper thresholds
#     if match_score >= 85 and semantic_score >= 0.7:  # Very high bar for "excellent"
#         insights.append({
#             'type': 'overall_quality',
#             'message': 'Excellent profile match',
#             'icon': 'star',
#             'color': 'success'
#         })
#     elif match_score >= 75 and semantic_score >= 0.6:  # High bar for "strong"
#         insights.append({
#             'type': 'overall_quality',
#             'message': 'Strong candidate potential',
#             'icon': 'thumbs-up',
#             'color': 'info'
#         })
    
#     return insights


# def generate_application_insights(app_data):
#     """Generate user-friendly insights for application cards"""
    
#     # Get scores using the helper function
#     match_score = get_match_score(app_data)  # ✅ This gets combined score
#     semantic_score = app_data.get('semantic_similarity', 0) or 0
#     enhanced_version = app_data.get('enhanced_version')
    
#     # Convert None values to safe defaults
#     if semantic_score is None:
#         semantic_score = 0
#     if match_score is None:
#         match_score = 0
    
#     # FIXED: Primary insight message with match score validation
#     if enhanced_version and semantic_score >= 0.7 and match_score >= 70:
#         primary_insight = "Detected strong skill alignment beyond keywords"
#     elif enhanced_version and semantic_score >= 0.5 and match_score >= 60:
#         primary_insight = "Found related skills and experience patterns"
#     elif enhanced_version and match_score < 60:
#         primary_insight = "Analysis suggests developing skills mentioned in job requirements"
#     elif enhanced_version:
#         primary_insight = "Matching applied"
#     else:
#         primary_insight = None
    
#     # FIXED: Secondary insights with proper thresholds
#     secondary_insights = []
    
#     if semantic_score >= 0.8 and match_score >= 75:  # Both must be high
#         secondary_insights.append("Exceptional skill relationship match")
#     elif semantic_score >= 0.6 and match_score >= 65:  # Both must be decent
#         secondary_insights.append("Good conceptual skill alignment")
    
#     if match_score >= 85 and semantic_score >= 0.7:
#         secondary_insights.append("Top-tier candidate profile")
    
#     return {
#         'primary': primary_insight,
#         'secondary': secondary_insights,
#         'has_ai_enhancement': bool(enhanced_version)
#     }

def calculate_enhanced_profile_completion(user_id):
    """Enhanced profile completion calculation including documents"""
    try:
        user_profile = IRS_USERS.find_one({"_id": ObjectId(user_id)})
        user_resume = resumeFetchedData.find_one({"UserId": ObjectId(user_id)})
        user_documents = USER_DOCUMENTS.count_documents({"UserId": ObjectId(user_id), "IsActive": True})
        
        completion_factors = {
            'basic_info': bool(user_profile and user_profile.get('Name') and user_profile.get('Email')),
            'profile_picture': bool(user_profile and user_profile.get('ProfilePicture')),
            'contact_info': bool(user_profile and user_profile.get('PhoneNumber')),
            'location': bool(user_profile and user_profile.get('Location')),
            'professional_info': bool(user_profile and user_profile.get('CurrentTitle') and user_profile.get('ExperienceLevel')),
            'professional_summary': bool(user_profile and user_profile.get('ProfessionalSummary')),
            'linkedin_profile': bool(user_profile and user_profile.get('LinkedInProfile')),
            'resume_uploaded': bool(user_resume),
            'skills_extracted': bool(user_resume and user_resume.get('SKILLS')),
            'documents_uploaded': bool(user_documents > 0)
        }
        
        completed = sum(completion_factors.values())
        return round((completed / len(completion_factors)) * 100)
    except Exception as e:
        logger.error(f"Error calculating enhanced profile completion: {e}")
        return 0
    
def get_missing_fields(user_id):
    """Get list of missing profile fields"""
    try:
        user_profile = IRS_USERS.find_one({"_id": ObjectId(user_id)})
        user_resume = resumeFetchedData.find_one({"UserId": ObjectId(user_id)})
        
        missing = []
        
        if not (user_profile and user_profile.get('Name')):
            missing.append('Name')
        if not (user_profile and user_profile.get('Email')):
            missing.append('Email')
        if not user_resume:
            missing.append('Resume')
        elif user_resume:
            if not user_resume.get('SKILLS'):
                missing.append('Skills')
            if not user_resume.get('CALCULATED_EXPERIENCE_YEARS', 0):
                missing.append('Experience')
            if not user_resume.get('EDUCATION'):
                missing.append('Education')
        
        return missing
    except Exception as e:
        logger.error(f"Error getting missing fields: {e}")
        return []


@app.route('/Applicant_Dashboard_Enhanced')
def Applicant_Dashboard_Enhanced():
    """CLEAN: Enhanced Applicant Dashboard with legitimate AI insights only"""
    try:
        if 'user_id' not in session or session.get('role') != 'applicant':
            flash("Access denied. Applicant login required.", "error")
            return redirect(url_for('loginpage'))
        
        user_id = session['user_id']
        current_date = datetime.now()
        current_time = current_date
        
        # Calculate profile completion and missing fields ONCE
        profile_completion_percentage = calculate_enhanced_profile_completion(user_id)
        missing_fields_list = get_missing_fields(user_id)
        
       # User Statistics - UPDATED to show AI insights instead of technical scores
        try:
            user_stats = {
                'total_applications': Applied_EMP.count_documents({"user_id": ObjectId(user_id)}),
                'interviews_scheduled': Applied_EMP.count_documents({
                    "user_id": ObjectId(user_id),
                    "status": {"$in": ["interviewed", "interview_scheduled"]}
                }),
                'applications_under_review': Applied_EMP.count_documents({
                    "user_id": ObjectId(user_id),
                    "status": "under_review"
                }),
                'applications_shortlisted': Applied_EMP.count_documents({
                    "user_id": ObjectId(user_id),
                    "status": "shortlisted"
                }),
                # ADD MISSING STATUS CALCULATIONS:
                'applications_pending': Applied_EMP.count_documents({
                    "user_id": ObjectId(user_id),
                    "status": {"$in": ["pending", "applied"]}
                }),
                'applications_interview_pending': Applied_EMP.count_documents({
                    "user_id": ObjectId(user_id),
                    "status": "interview_pending"
                }),
                'applications_rejected': Applied_EMP.count_documents({
                    "user_id": ObjectId(user_id),
                    "status": {"$in": ["rejected", "not_selected"]}
                }),
                'applications_hired': Applied_EMP.count_documents({
                    "user_id": ObjectId(user_id),
                    "status": "hired"
                }),
                'ai_enhanced_applications': Applied_EMP.count_documents({
                    "user_id": ObjectId(user_id),
                    "enhanced_version": {"$exists": True}
                })
            }
                        
            # Calculate averages for internal use (not displayed to user)
            avg_result = list(Applied_EMP.aggregate([
                {"$match": {"user_id": ObjectId(user_id)}},
                {"$group": {
                    "_id": None, 
                    "avg_combined": {"$avg": "$combined_score"},  # ✅ Use combined_score
                    "avg_traditional": {"$avg": "$Matching_percentage"},
                    "avg_semantic": {"$avg": "$semantic_similarity"}
                }}
            ]))

            if avg_result:
                # Use combined_score average (the actual displayed score)
                user_stats['avg_match_score'] = round(avg_result[0].get('avg_combined', 0) or 0, 1)
                
                # Store individual averages for internal use if needed
                avg_traditional = avg_result[0].get('avg_traditional', 0) or 0
                avg_semantic = avg_result[0].get('avg_semantic', 0) or 0
            else:
                user_stats['avg_match_score'] = 0
                avg_semantic = 0
                avg_traditional = 0
            
            # Generate user-friendly AI performance summary
            if user_stats['ai_enhanced_applications'] > 0:
                ai_percentage = (user_stats['ai_enhanced_applications'] / user_stats['total_applications']) * 100
                if ai_percentage >= 80:
                    user_stats['ai_summary'] = "Most applications use advanced AI matching"
                elif ai_percentage >= 50:
                    user_stats['ai_summary'] = "Half your applications benefit from AI enhancement"
                else:
                    user_stats['ai_summary'] = f"{user_stats['ai_enhanced_applications']} applications use AI matching"
            else:
                user_stats['ai_summary'] = "Upload resume to enable AI matching"
            
        except Exception as e:
            logger.error(f"Error calculating user stats: {e}")
            user_stats = {
                'total_applications': 0, 'interviews_scheduled': 0,
                'applications_under_review': 0, 'applications_shortlisted': 0,
                'avg_match_score': 0, 'ai_enhanced_applications': 0,
                'ai_summary': 'Data unavailable'
            }
        
        # Recent Applications with User-Friendly AI Insights
        recent_applications = []
        try:
            recent_apps = list(Applied_EMP.find(
                {"user_id": ObjectId(user_id)},
                {
                    "job_id": 1, "Matching_percentage": 1, "applied_at": 1, "status": 1, 
                    "enhanced_version": 1, "semantic_similarity": 1, "combined_score": 1  # ✅ Add this
                }
            ).sort([("applied_at", -1)]).limit(8))
            
            for app in recent_apps:
                try:
                    job = JOBS.find_one({"_id": app.get('job_id')}, {"Job_Profile": 1, "CompanyName": 1})
                    
                    if job:
                        recent_applications.append({
                            'job_title': job.get('Job_Profile', 'Unknown Job'),
                            'company_name': job.get('CompanyName', 'Company'),
                            'match_score': round(get_match_score(app), 1),  # ✅ Round for consistency
                            'applied_date': app.get('applied_at'),
                            'status': app.get('status', 'pending'),
                            'ai_enhanced': bool(app.get('enhanced_version')),
                        })
                except Exception as e:
                    logger.error(f"Error processing recent application: {e}")
                    continue
                    
        except Exception as e:  # ✅ Add outer except
            logger.error(f"Error fetching recent applications: {e}")
                    
        except Exception as e:
            logger.error(f"Error fetching recent applications: {e}")
        
        # Upcoming Interviews
        upcoming_interviews = []
        try:
            interview_apps = list(Applied_EMP.find({
                "user_id": ObjectId(user_id),
                "status": {"$in": ["interviewed", "interview_scheduled"]}
            }).limit(5))
            
            for app in interview_apps:
                try:
                    job = JOBS.find_one({"_id": app.get('job_id')}, {"Job_Profile": 1, "CompanyName": 1})
                    
                    if job:
                        interview_date = datetime.now() + timedelta(days=random.randint(1, 14))
                        
                        upcoming_interviews.append({
                            'job_title': job.get('Job_Profile', 'Unknown Job'),
                            'company_name': job.get('CompanyName', 'Company'),
                            'interview_date': interview_date,
                            'interview_type': 'Technical Interview',
                            'meeting_link': None
                        })
                except Exception as e:
                    logger.error(f"Error processing interview: {e}")
                    continue
                    
        except Exception as e:
            logger.error(f"Error fetching upcoming interviews: {e}")
        
        # Job Recommendations
        recommended_jobs = []
        try:
            user_resume = resumeFetchedData.find_one({"UserId": ObjectId(user_id)})
            if user_resume:
                user_skills = user_resume.get('SKILLS', [])
                user_experience = user_resume.get('CALCULATED_EXPERIENCE_YEARS', 0)
                
                if user_skills:
                    skill_pattern = '|'.join(user_skills[:15])
                    
                    matching_jobs = list(JOBS.find({
                        "Status": "Open",
                        "$or": [
                            {"Job_Description": {"$regex": skill_pattern, "$options": "i"}},
                            {"Job_Profile": {"$regex": skill_pattern, "$options": "i"}}
                        ]
                    }).limit(6))
                    
                    for job in matching_jobs:
                        try:
                            job_text = f"{job.get('Job_Description', '')} {job.get('Job_Profile', '')}".lower()
                            skill_matches = sum(1 for skill in user_skills if skill.lower() in job_text)
                            match_score = min(95, round((skill_matches / len(user_skills)) * 100)) if user_skills else 0
                            
                            if user_experience > 0:
                                match_score = min(100, match_score + random.randint(5, 15))
                            
                            recommended_jobs.append({
                                'id': str(job['_id']),
                                'title': job.get('Job_Profile', 'Unknown Job'),
                                'company': job.get('CompanyName', 'Company'),
                                'description': job.get('Job_Description', '')[:150] + "...",
                                'match_score': max(60, match_score),
                                'salary_range': job.get('Salary'),
                                'key_skills': [skill for skill in user_skills[:6] if skill.lower() in job_text.lower()][:4],
                                'posted_date': job.get('CreatedAt'),
                                'deadline': job.get('LastDate')
                            })
                        except Exception as e:
                            logger.error(f"Error processing job recommendation: {e}")
                            continue
                    
                    recommended_jobs.sort(key=lambda x: x['match_score'], reverse=True)
                        
        except Exception as e:
            logger.error(f"Error fetching job recommendations: {e}")
        
        # User Resumes with Analytics
        try:
            user_resumes = list(resumeFetchedData.find({"UserId": ObjectId(user_id)}))
            
            for resume in user_resumes:
                resume_applications = Applied_EMP.count_documents({
                    "user_id": ObjectId(user_id),
                    "resume_id": resume.get('_id')
                })
                resume['application_count'] = resume_applications
                
        except Exception as e:
            logger.error(f"Error fetching user resumes: {e}")
            user_resumes = []
        
        # LEGITIMATE AI Insights only
        ai_insights = []
        
        try:
            total_apps = user_stats['total_applications']
            avg_score = user_stats['avg_match_score']
            interviews = user_stats['interviews_scheduled']
            
            if total_apps == 0:
                ai_insights.append({
                    'title': 'Ready to Start',
                    'description': 'Upload your resume and begin applying to jobs that match your skills and experience level.',
                    'action': 'Upload Resume',
                    'action_url': '/uploadResume'
                })
            elif total_apps < 5:
                ai_insights.append({
                    'title': 'Building Your Track Record',
                    'description': f'You\'ve applied to {total_apps} jobs. Apply to 5-8 more similar positions to understand your competitive range.',
                    'action': 'Find Similar Jobs',
                    'action_url': '/HR1/show_job'
                })
            else:
                # Realistic performance assessment based on actual data
                interview_rate = (interviews / total_apps) * 100 if total_apps > 0 else 0
                
                if avg_score >= 75 and interview_rate >= 15:
                    ai_insights.append({
                        'title': 'Strong Performance',
                        'description': f'Your {avg_score:.0f}% average match and {interview_rate:.0f}% interview rate show you\'re targeting appropriate roles.',
                        'action': 'Continue Strategy',
                        'action_url': '/HR1/show_job'
                    })
                elif avg_score >= 60 and interview_rate < 10:
                    ai_insights.append({
                        'title': 'Good Matches, Few Interviews',
                        'description': f'Your {avg_score:.0f}% scores are solid, but {interview_rate:.0f}% interview rate suggests improving application timing and customization.',
                        'action': 'Improve Applications',
                        'action_url': '/uploadResume'
                    })
                elif avg_score < 60:
                    ai_insights.append({
                        'title': 'Adjust Your Targeting',
                        'description': f'Your {avg_score:.0f}% average suggests applying to roles requiring 1-2 years less experience than you\'re currently targeting.',
                        'action': 'Find Entry-Level Roles',
                        'action_url': '/HR1/show_job'
                    })
                else:
                    ai_insights.append({
                        'title': 'Steady Progress',
                        'description': f'Your {avg_score:.0f}% average shows you\'re on the right track. Keep applying consistently.',
                        'action': 'Browse More Jobs',
                        'action_url': '/HR1/show_job'
                    })
            
            # Profile completion insight (realistic impact)
            if profile_completion_percentage < 80:
                ai_insights.append({
                    'title': 'Incomplete Profile',
                    'description': f'Your {profile_completion_percentage}% complete profile may be limiting your visibility to employers.',
                    'action': 'Complete Profile',
                    'action_url': '/uploadResume'
                })
            
            # Market activity insight (based on available jobs)
            if len(recommended_jobs) >= 4:
                ai_insights.append({
                    'title': 'Active Job Market',
                    'description': f'Found {len(recommended_jobs)} relevant openings. Focus on applying to 2-3 per week with customized applications.',
                    'action': 'View Recommendations',
                    'action_url': '/HR1/show_job'
                })
            elif len(recommended_jobs) <= 1:
                ai_insights.append({
                    'title': 'Limited Matches',
                    'description': 'Few jobs match your current profile. Consider expanding your skill set or broadening your search criteria.',
                    'action': 'Explore Skills',
                    'action_url': '/uploadResume'
                })
            
            # Interview preparation insight
            if interviews > 0 and total_apps >= 5:
                ai_insights.append({
                    'title': 'Interview Stage',
                    'description': f'You\'ve reached {interviews} interview(s). Focus on interview preparation and follow-up communication.',
                    'action': 'Interview Tips',
                    'action_url': '/interviews'
                })
            
            # Fallback if no specific insights generated
            if not ai_insights:
                ai_insights.append({
                    'title': 'Keep Applying',
                    'description': 'Continue submitting quality applications to build your job search momentum.',
                    'action': 'Browse Jobs',
                    'action_url': '/HR1/show_job'
                })
                
        except Exception as e:
            logger.error(f"Error generating AI insights: {e}")
            ai_insights = [{
                'title': 'Analysis Unavailable',
                'description': 'Unable to generate insights at this time. Continue applying to relevant positions.',
                'action': 'Browse Jobs',
                'action_url': '/HR1/show_job'
            }]

            # ADD DEBUGGING RIGHT BEFORE RETURN STATEMENT
        
            # Debug logging to see what's actually being passed
            logger.info(f"DEBUG - ai_insights count: {len(ai_insights)}")
            logger.info(f"DEBUG - ai_insights content: {ai_insights}")
            
            # Debug recent applications
            if not ai_insights:
                ai_insights = [{
                    'title': 'Test Insight',
                    'description': 'This is a test insight to verify the template is working',
                    'action': 'Test Action',
                    'action_url': '/test'
                }]
                logger.info("DEBUG - Added test insight")
        
        return render_template("applicant_dashboard_enhanced.html",
                             current_date=current_date,
                             current_time=current_time,
                             profile_completion=profile_completion_percentage,
                             missing_fields=missing_fields_list,
                             user_stats=user_stats,
                             recent_applications=recent_applications,
                             upcoming_interviews=upcoming_interviews,
                             recommended_jobs=recommended_jobs,
                             user_resumes=user_resumes,
                             ai_insights=ai_insights)
                             
    except Exception as e:
        logger.error(f"Error in Applicant_Dashboard_Enhanced: {e}")
        return render_template("applicant_dashboard_enhanced.html",
                             current_date=datetime.now(),
                             current_time=datetime.now(),
                             profile_completion=0,
                             missing_fields=[],
                             user_stats={},
                             recent_applications=[],
                             upcoming_interviews=[],
                             recommended_jobs=[],
                             user_resumes=[],
                             ai_insights=[],
                             errorMsg="An error occurred loading the dashboard")






    
@app.route('/Applicant_Dashboard', methods=['GET', 'POST'])
def Applicant_Dashboard():
    """Redirect to enhanced applicant dashboard"""
    try:
        if 'user_id' not in session or session.get('role') != 'applicant':
            flash("Access denied. Applicant login required.", "error")
            return redirect(url_for('loginpage'))
        
        return redirect(url_for('Applicant_Dashboard_Enhanced'))
    except Exception as e:
        logger.error(f"Error in Applicant_Dashboard: {e}")
        return redirect(url_for('Applicant_Dashboard_Enhanced'))


def generate_hr_candidate_insights(app_data):
    """Generate HR-friendly insights for candidates"""
    semantic_score = app_data.get('semantic_similarity', 0)
    match_score = app_data.get('Matching_percentage', 0)
    enhanced_version = app_data.get('enhanced_version')
    
    insights = []
    
    # Skill assessment insight
    if semantic_score >= 0.8:
        insights.append("Strong skill relationships identified")
    elif semantic_score >= 0.6:
        insights.append("Good skill transferability detected")
    elif semantic_score >= 0.4:
        insights.append("Some relevant skills found")
    
    # Overall recommendation
    if match_score >= 80 and semantic_score >= 0.7:
        recommendation = "Highly recommended - excellent overall fit"
    elif match_score >= 70 and semantic_score >= 0.6:
        recommendation = "Recommended - strong potential"
    elif match_score >= 60:
        recommendation = "Consider - decent baseline match"
    else:
        recommendation = "Review carefully - limited alignment"
    
    return {
        'insights': insights,
        'recommendation': recommendation,
        'ai_processed': bool(enhanced_version)
    }

@app.route('/HR_Dashboard_Enhanced')
def HR_Dashboard_Enhanced():
    """Enhanced HR Dashboard with USER-SPECIFIC data only"""
    try:
        if 'user_id' not in session or session.get('role') != 'employer':
            flash("Access denied. Employer login required.", "error")
            return redirect(url_for('loginpage'))
        
        current_hr_user_id = session['user_id']
        current_date = datetime.now()
        current_time = current_date
        
        # Calculate date ranges
        week_ago = current_date - timedelta(days=7)
        month_ago = current_date - timedelta(days=30)
        today_start = current_date.replace(hour=0, minute=0, second=0, microsecond=0)
        
        # Get HR jobs (same as HR Profile)
        hr_jobs_query = {"created_by": ObjectId(current_hr_user_id)}
        hr_jobs = list(JOBS.find(hr_jobs_query, {"_id": 1, "Job_Profile": 1, "Status": 1, "CreatedAt": 1}))
        hr_job_ids = [job['_id'] for job in hr_jobs]
        
        # Initialize stats
        dashboard_stats = {
            'total_jobs': len(hr_jobs),
            'new_jobs_this_week': 0,
            'active_jobs': 0,
            'total_applications': 0,
            'new_applications_today': 0,
            'new_applications_week': 0,
            'interviews_scheduled': 0,
            'interviews_today': 0,
            'avg_match_score': 0,
            'ai_enhanced_applications': 0
        }
        
        pipeline_stats = {
            'applied': 0,
            'screening': 0,
            'interview': 0,
            'offer': 0,
            'hired': 0
        }
        
        # Calculate stats if HR has jobs
        if hr_job_ids:
            # Job stats
            dashboard_stats['new_jobs_this_week'] = sum(1 for job in hr_jobs if job.get('CreatedAt', datetime.min) >= week_ago)
            dashboard_stats['active_jobs'] = sum(1 for job in hr_jobs if job.get('Status') == 'Open')
            
            # Application query for this HR's jobs
            hr_app_query = {"job_id": {"$in": hr_job_ids}}
            
            # Get all applications for stats
            all_applications = list(Applied_EMP.find(hr_app_query))
            
            dashboard_stats['total_applications'] = len(all_applications)
            
            # Process each application for detailed stats
            match_scores = []
            for app in all_applications:
                applied_at = app.get('applied_at', datetime.min)
                
                # Today's applications
                if applied_at >= today_start:
                    dashboard_stats['new_applications_today'] += 1
                
                # Week's applications
                if applied_at >= week_ago:
                    dashboard_stats['new_applications_week'] += 1
                
                # Interview stats
                status = app.get('status', '')
                if status in ['interview_scheduled', 'interviewed']:
                    dashboard_stats['interviews_scheduled'] += 1
                    
                    interview_date = app.get('interview_date')
                    if interview_date and today_start <= interview_date < (today_start + timedelta(days=1)):
                        dashboard_stats['interviews_today'] += 1
                
                # Pipeline stats
                if status in ['pending', 'applied', 'under_review']:
                    pipeline_stats['applied'] += 1
                elif status in ['shortlisted', 'interview_pending']:
                    pipeline_stats['screening'] += 1
                elif status in ['interview_scheduled', 'interviewed']:
                    pipeline_stats['interview'] += 1
                elif status == 'offer_made':
                    pipeline_stats['offer'] += 1
                elif status == 'hired':
                    pipeline_stats['hired'] += 1
                
                # AI enhancement check
                if app.get('semantic_similarity') is not None:
                    dashboard_stats['ai_enhanced_applications'] += 1
                
                # Calculate match score
                match_score = get_match_score(app)
                if match_score > 0:
                    match_scores.append(match_score)
            
            # Calculate average match score
            if match_scores:
                dashboard_stats['avg_match_score'] = round(sum(match_scores) / len(match_scores), 1)
        
        # Recent Activities (same logic as HR Profile)
        recent_activities = []
        if hr_job_ids:
            recent_apps = list(Applied_EMP.find(
                {"job_id": {"$in": hr_job_ids}}
            ).sort("applied_at", -1).limit(10))
            
            for app in recent_apps:
                job = JOBS.find_one({"_id": app.get('job_id')}, {"Job_Profile": 1})
                time_ago = get_time_ago(app.get('applied_at', datetime.now()))
                
                activity_type = "new_application"
                if app.get('status') == 'interviewed':
                    activity_type = "interview_scheduled"
                elif app.get('status') == 'hired':
                    activity_type = "candidate_hired"
                
                recent_activities.append({
                    'icon': get_activity_icon(activity_type),
                    'color': get_activity_color(activity_type),
                    'title': get_activity_title(activity_type),
                    'description': f"{app.get('User_name', 'Unknown')} - {job.get('Job_Profile', 'Unknown Job') if job else 'Unknown Job'}",
                    'time_ago': time_ago,
                    'enhanced': app.get('semantic_similarity') is not None
                })
        
        # Top Candidates This Week
        top_candidates = []
        if hr_job_ids:
            top_apps = list(Applied_EMP.find({
                "job_id": {"$in": hr_job_ids},
                "applied_at": {"$gte": week_ago}
            }).sort("applied_at", -1).limit(50))  # Get more to sort by score
            
            # Calculate scores and sort
            candidates_with_scores = []
            for app in top_apps:
                user = IRS_USERS.find_one({"_id": app.get('user_id')}, {"Name": 1})
                job = JOBS.find_one({"_id": app.get('job_id')}, {"Job_Profile": 1})
                
                if user and job:
                    match_score = get_match_score(app)
                    
                    # Get individual scores
                    traditional_score = app.get('Matching_percentage', 0)
                    if isinstance(traditional_score, dict):
                        traditional_score = traditional_score.get('overall_score', 0)
                    try:
                        traditional_score = float(traditional_score)
                    except (ValueError, TypeError):
                        traditional_score = 0
                    
                    semantic_score = app.get('semantic_similarity')
                    if semantic_score is not None:
                        try:
                            semantic_score = float(semantic_score)
                        except (ValueError, TypeError):
                            semantic_score = None
                    
                    candidates_with_scores.append({
                        'id': str(app.get('user_id')),
                        'name': user.get('Name', 'Unknown'),
                        'job_title': job.get('Job_Profile', 'Unknown Job'),
                        'match_score': round(match_score, 1),
                        'traditional_score': round(traditional_score, 1),
                        'ai_enhanced': semantic_score is not None and semantic_score > 0,
                        'semantic_score': round(semantic_score * 100, 1) if semantic_score else None,
                    })
            
            # Sort by match score and take top 8
            candidates_with_scores.sort(key=lambda x: x['match_score'], reverse=True)
            top_candidates = candidates_with_scores[:8]
        
        # AI Insights
        ai_insights = []
        try:
            from sentence_transformer_ranker import get_semantic_ranker
            ranker = get_semantic_ranker()
            
            if ranker.is_available():
                if dashboard_stats['ai_enhanced_applications'] > 0:
                    ai_insights.append({
                        'title': 'AI Semantic Matching Active',
                        'description': f'{dashboard_stats["ai_enhanced_applications"]} applications analyzed with semantic similarity.',
                        'action': 'View Details',
                        'action_url': '/HR1/Company_Candidates'
                    })
                
                if dashboard_stats['avg_match_score'] > 75:
                    ai_insights.append({
                        'title': 'High-Quality Candidate Pool',
                        'description': f'Average match score: {dashboard_stats["avg_match_score"]}%',
                        'action': 'Review Top Candidates',
                        'action_url': '/HR1/Company_Candidates'
                    })
        except ImportError:
            pass
        
        logger.info(f"Dashboard loaded for HR {current_hr_user_id}: {dashboard_stats['total_jobs']} jobs, {dashboard_stats['total_applications']} applications")
        
        return render_template("hr_dashboard_enhanced.html",
                             current_date=current_date,
                             current_time=current_time,
                             dashboard_stats=dashboard_stats,
                             pipeline_stats=pipeline_stats,
                             recent_activities=recent_activities,
                             top_candidates=top_candidates,
                             ai_insights=ai_insights,
                             jobs=hr_jobs)
                             
    except Exception as e:
        logger.error(f"Error in HR_Dashboard_Enhanced: {e}")
        import traceback
        logger.error(traceback.format_exc())
        
        return render_template("hr_dashboard_enhanced.html",
                             current_date=datetime.now(),
                             current_time=datetime.now(),
                             dashboard_stats={
                                 'total_jobs': 0, 'new_jobs_this_week': 0, 'active_jobs': 0,
                                 'total_applications': 0, 'new_applications_today': 0,
                                 'new_applications_week': 0, 'interviews_scheduled': 0,
                                 'interviews_today': 0, 'avg_match_score': 0,
                                 'ai_enhanced_applications': 0
                             },
                             pipeline_stats={'applied': 0, 'screening': 0, 'interview': 0, 'offer': 0, 'hired': 0},
                             recent_activities=[],
                             top_candidates=[],
                             ai_insights=[],
                             jobs=[],
                             errorMsg="An error occurred loading the dashboard")

@app.route('/view_candidate_document/<document_id>')
def view_candidate_document(document_id):
    """View additional documents uploaded by candidates for HR to review"""
    try:
        # Check if HR is logged in
        if 'user_id' not in session or session.get('role') != 'employer':
            flash("Access denied. Employer login required.", "error")
            return redirect(url_for('loginpage'))
        
        try:
            doc_obj_id = ObjectId(document_id)
        except:
            flash("Invalid document ID", "error")
            return redirect(url_for('Company_Candidates'))
        
        # Get document from USER_DOCUMENTS collection
        document = mongo.db.USER_DOCUMENTS.find_one({"_id": doc_obj_id})
        
        if not document:
            flash("Document not found", "error")
            logger.error(f"Document not found: {document_id}")
            return redirect(url_for('Company_Candidates'))
        
        # ✅ Use FileName field (same as working applicant view_document)
        internal_filename = document.get('FileName')
        
        if not internal_filename:
            flash("Document file reference not found", "error")
            logger.error(f"No FileName found in document {document_id}")
            logger.error(f"Document fields: {list(document.keys())}")
            return redirect(url_for('Company_Candidates'))
        
        # ✅ Use DOCUMENT_UPLOAD_FOLDER variable (NOT app.config)
        # This matches the working applicant view_document route
        file_path = os.path.join(DOCUMENT_UPLOAD_FOLDER, internal_filename)
        
        logger.info(f"Attempting to serve document: {file_path}")
        
        # Check if file exists
        if not os.path.exists(file_path):
            flash("Document file not found on server", "error")
            logger.error(f"Document file not found: {file_path}")
            logger.error(f"DOCUMENT_UPLOAD_FOLDER: {DOCUMENT_UPLOAD_FOLDER}")
            logger.error(f"FileName: {internal_filename}")
            
            # List files in upload folder for debugging
            try:
                files_in_folder = os.listdir(DOCUMENT_UPLOAD_FOLDER)
                logger.error(f"Files in folder ({len(files_in_folder)} total): {files_in_folder[:10]}")
            except Exception as e:
                logger.error(f"Could not list folder: {e}")
            
            return redirect(url_for('Company_Candidates'))
        
        # Determine mimetype based on file extension
        file_ext = os.path.splitext(internal_filename)[1].lower()
        mimetype_map = {
            '.pdf': 'application/pdf',
            '.doc': 'application/msword',
            '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            '.txt': 'text/plain',
            '.jpg': 'image/jpeg',
            '.jpeg': 'image/jpeg',
            '.png': 'image/png',
        }
        mimetype = mimetype_map.get(file_ext, 'application/octet-stream')
        
        logger.info(f"Serving document with mimetype: {mimetype}")
        
        # Serve the file - exactly like the working applicant route
        return send_file(file_path, as_attachment=False, mimetype=mimetype)
        
    except Exception as e:
        logger.error(f"Error viewing document {document_id}: {e}")
        import traceback
        logger.error(traceback.format_exc())
        flash("Error loading document", "error")
        return redirect(url_for('Company_Candidates'))


@app.route('/download_candidate_document/<document_id>')
def download_candidate_document(document_id):
    """Download additional documents (force download instead of view)"""
    try:
        # Check if HR/Employer is logged in
        if 'user_id' not in session or session.get('role') != 'employer':
            return jsonify({"error": "Access denied"}), 403
        
        # Validate and fetch document
        try:
            doc_obj_id = ObjectId(document_id)
        except:
            return jsonify({"error": "Invalid document ID"}), 400
        
        document = mongo.db.USER_DOCUMENTS.find_one({"_id": doc_obj_id})
        
        if not document:
            return jsonify({"error": "Document not found"}), 404
        
        # Get filename
        internal_filename = document.get('InternalFilename') or document.get('FileName')
        original_filename = document.get('FileName') or document.get('DocumentName', 'document')
        
        if not internal_filename:
            return jsonify({"error": "File reference not found"}), 404
        
        # Build file path
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], internal_filename)
        
        if not os.path.exists(file_path):
            logger.error(f"Document file not found: {file_path}")
            return jsonify({"error": "File not found on server"}), 404
        
        # Force download with original filename
        return send_file(
            file_path,
            as_attachment=True,
            download_name=original_filename
        )
        
    except Exception as e:
        logger.error(f"Error downloading document {document_id}: {e}")
        return jsonify({"error": "Download failed"}), 500


# Add this helper function to fetch documents when loading candidates
def get_user_documents(user_id):
    """
    Helper function to get all documents for a specific user
    Call this when fetching candidate data
    """
    try:
        user_obj_id = ObjectId(user_id)
        documents = list(mongo.db.USER_DOCUMENTS.find(
            {"UserId": user_obj_id}
        ).sort("UploadedAt", -1))
        return documents
    except Exception as e:
        logger.error(f"Error fetching documents for user {user_id}: {e}")
        return []

@app.route('/view_candidate_resume/<resume_id>')
def view_candidate_resume(resume_id):
    """View uploaded resume for HR to review candidates"""
    try:
        # Check if HR is logged in
        if 'user_id' not in session or session.get('role') != 'employer':
            flash("Access denied. Employer login required.", "error")
            return redirect(url_for('loginpage'))
        
        # Get resume document
        resume = resumeFetchedData.find_one({"_id": ObjectId(resume_id)})
        
        if not resume:
            flash("Resume not found", "error")
            return redirect(url_for('Company_Candidates'))
        
        # Get the INTERNAL filename (with timestamp prefix)
        # Your upload saves as: timestamp + original_filename
        internal_filename = resume.get('InternalFilename')
        
        if not internal_filename:
            # Fallback to other possible fields
            internal_filename = resume.get('ResumeTitle') or resume.get('Filename')
        
        if not internal_filename:
            flash("Resume file reference not found", "error")
            logger.error(f"No filename found in resume document {resume_id}")
            logger.error(f"Resume document fields: {resume.keys()}")
            return redirect(url_for('Company_Candidates'))
        
        # Build file path - files are saved directly in UPLOAD_FOLDER
        # NOT in a 'uploaded_resumes' subdirectory based on your code
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], internal_filename)
        
        logger.info(f"Attempting to serve resume: {file_path}")
        
        # Check if file exists
        if not os.path.exists(file_path):
            flash("Resume file not found on server", "error")
            logger.error(f"Resume file not found: {file_path}")
            logger.error(f"UPLOAD_FOLDER: {app.config['UPLOAD_FOLDER']}")
            logger.error(f"Internal filename: {internal_filename}")
            
            # List files in upload folder for debugging
            try:
                files_in_folder = os.listdir(app.config['UPLOAD_FOLDER'])
                logger.error(f"Files in upload folder: {files_in_folder[:10]}")  # Show first 10
            except Exception as e:
                logger.error(f"Could not list upload folder: {e}")
            
            return redirect(url_for('Company_Candidates'))
        
        # Determine mimetype based on file extension
        file_ext = os.path.splitext(internal_filename)[1].lower()
        if file_ext == '.pdf':
            mimetype = 'application/pdf'
        elif file_ext in ['.docx', '.doc']:
            mimetype = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        else:
            mimetype = 'application/octet-stream'
        
        # Serve the file
        return send_file(
            file_path,
            as_attachment=False,  # Display in browser, not download
            mimetype=mimetype
        )
        
    except Exception as e:
        logger.error(f"Error viewing resume: {e}")
        import traceback
        logger.error(traceback.format_exc())
        flash("Error loading resume", "error")
        return redirect(url_for('Company_Candidates'))

@app.route('/hr_profile')
def hr_profile():
    """HR profile page with company details and stats"""
    try:
        if 'user_id' not in session or session.get('role') != 'employer':
            flash("Access denied. Employer login required.", "error")
            return redirect(url_for('loginpage'))
        
        user_id = session['user_id']
        
        # Get HR user profile information
        try:
            user_profile = IRS_USERS.find_one({"_id": ObjectId(user_id)})
            if not user_profile:
                flash("User profile not found", "error")
                return redirect(url_for('loginpage'))
        except Exception as e:
            logger.error(f"Error fetching HR profile for {user_id}: {e}")
            user_profile = {}
        
        # Calculate profile completion for HR
        try:
            completion_factors = {
                'basic_info': bool(user_profile.get('Name') and user_profile.get('Email')),
                'profile_picture': bool(user_profile.get('ProfilePicture')),
                'contact_info': bool(user_profile.get('PhoneNumber')),
                'job_details': bool(user_profile.get('JobTitle') and user_profile.get('CompanyName')),
                'company_info': bool(user_profile.get('Industry') and user_profile.get('CompanySize')),
                'location': bool(user_profile.get('Location')),
                'experience': bool(user_profile.get('HRExperience')),
                'linkedin': bool(user_profile.get('LinkedInProfile')),
                'company_details': bool(user_profile.get('CompanyDescription')),
                'hiring_volume': bool(user_profile.get('HiringVolume'))
            }
            
            completed = sum(completion_factors.values())
            profile_completion = round((completed / len(completion_factors)) * 100)
        except Exception as e:
            logger.error(f"Error calculating HR profile completion: {e}")
            profile_completion = 0
        
        # Get HR statistics
        try:
            current_date = datetime.now()
            week_ago = current_date - timedelta(days=7)
            month_ago = current_date - timedelta(days=30)
            
            # Get jobs created by this HR user
            hr_jobs = list(JOBS.find({"created_by": ObjectId(user_id)}).sort("CreatedAt", -1))
            hr_job_ids = [job['_id'] for job in hr_jobs]
            
            # Calculate stats
            hr_stats = {
                'total_jobs': len(hr_jobs),
                'total_applications': 0,
                'interviews_scheduled': 0,
                'candidates_hired': 0,
                'jobs_change': 0,
                'applications_change': 0
            }
            
            if hr_job_ids:
                # Applications for this HR user's jobs
                hr_stats['total_applications'] = Applied_EMP.count_documents({"job_id": {"$in": hr_job_ids}})
                hr_stats['interviews_scheduled'] = Applied_EMP.count_documents({
                    "job_id": {"$in": hr_job_ids},
                    "status": {"$in": ["interview_scheduled", "interviewed"]}
                })
                hr_stats['candidates_hired'] = Applied_EMP.count_documents({
                    "job_id": {"$in": hr_job_ids},
                    "status": "hired"
                })
                
                # Calculate changes
                hr_stats['applications_change'] = Applied_EMP.count_documents({
                    "job_id": {"$in": hr_job_ids},
                    "applied_at": {"$gte": week_ago}
                })
                
                hr_stats['jobs_change'] = JOBS.count_documents({
                    "created_by": ObjectId(user_id),
                    "CreatedAt": {"$gte": month_ago}
                })
            
            # Add application count to jobs
            for job in hr_jobs:
                job['application_count'] = Applied_EMP.count_documents({"job_id": job['_id']})
                
        except Exception as e:
            logger.error(f"Error calculating HR stats: {e}")
            hr_stats = {'total_jobs': 0, 'total_applications': 0, 'interviews_scheduled': 0, 'candidates_hired': 0}
            hr_jobs = []
        
        # Get recent activities
        try:
            recent_activities = []
            if hr_job_ids:
                recent_apps = list(Applied_EMP.find({
                    "job_id": {"$in": hr_job_ids}
                }).sort("applied_at", -1).limit(10))
                
                for app in recent_apps:
                    try:
                        job = JOBS.find_one({"_id": app.get('job_id')}, {"Job_Profile": 1})
                        time_ago = get_time_ago(app.get('applied_at', datetime.now()))
                        
                        activity_type = "new_application"
                        if app.get('status') == 'interviewed':
                            activity_type = "interview_completed"
                        elif app.get('status') == 'interview_scheduled':
                            activity_type = "interview_scheduled"
                        elif app.get('status') == 'hired':
                            activity_type = "candidate_hired"
                        
                        recent_activities.append({
                            'icon': get_activity_icon(activity_type),
                            'color': get_activity_color(activity_type),
                            'title': get_activity_title(activity_type),
                            'description': f"{app.get('User_name', 'Candidate')} applied to {job.get('Job_Profile', 'Job') if job else 'Job'}",
                            'time_ago': time_ago
                        })
                    except Exception as e:
                        logger.error(f"Error processing recent activity: {e}")
                        continue
        except Exception as e:
            logger.error(f"Error fetching recent activities: {e}")
            recent_activities = []
        
        logger.info(f"HR profile loaded for user {user_id}: {len(hr_jobs)} jobs, {hr_stats['total_applications']} applications")
        
        return render_template("hr_profile.html",
                             user_profile=user_profile,
                             profile_completion=profile_completion,
                             hr_stats=hr_stats,
                             hr_jobs=hr_jobs,
                             recent_activities=recent_activities)
                             
    except Exception as e:
        logger.error(f"Error in hr_profile: {e}")
        return render_template("hr_profile.html",
                             user_profile={},
                             profile_completion=0,
                             hr_stats={},
                             hr_jobs=[],
                             recent_activities=[],
                             errorMsg="An error occurred loading your profile")

@app.route('/update_hr_profile', methods=['POST'])
def update_hr_profile():
    """Update HR user profile information"""
    try:
        if 'user_id' not in session or session.get('role') != 'employer':
            flash("Access denied", "error")
            return redirect(url_for('loginpage'))
        
        user_id = session['user_id']
        
        # Get form data
        update_data = {
            'Name': request.form.get('full_name', '').strip(),
            'PhoneNumber': request.form.get('phone_number', '').strip(),
            'JobTitle': request.form.get('job_title', '').strip(),
            'Department': request.form.get('department', '').strip(),
            'HRExperience': request.form.get('hr_experience', '').strip(),
            'LinkedInProfile': request.form.get('linkedin_profile', '').strip(),
            'Location': request.form.get('location', '').strip(),
            'ProfileUpdatedAt': datetime.now()
        }
        
        # Remove empty values
        update_data = {k: v for k, v in update_data.items() if v}
        
        # Update user profile
        result = IRS_USERS.update_one(
            {"_id": ObjectId(user_id)},
            {"$set": update_data}
        )
        
        if result.modified_count > 0:
            logger.info(f"HR profile updated for user {user_id}")
            flash("Profile updated successfully!", "success")
        else:
            flash("No changes made to profile", "info")
            
        return redirect(url_for('hr_profile'))
        
    except Exception as e:
        logger.error(f"Error updating HR profile: {e}")
        flash("Error updating profile", "error")
        return redirect(url_for('hr_profile'))

@app.route('/update_company_profile', methods=['POST'])
def update_company_profile():
    """Update company profile information"""
    try:
        if 'user_id' not in session or session.get('role') != 'employer':
            flash("Access denied", "error")
            return redirect(url_for('loginpage'))
        
        user_id = session['user_id']
        
        # Get form data
        update_data = {
            'CompanyName': request.form.get('company_name', '').strip(),
            'Industry': request.form.get('industry', '').strip(),
            'CompanySize': request.form.get('company_size', '').strip(),
            'CompanyWebsite': request.form.get('company_website', '').strip(),
            'HiringVolume': request.form.get('hiring_volume', '').strip(),
            'CompanyAddress': request.form.get('company_address', '').strip(),
            'CompanyDescription': request.form.get('company_description', '').strip(),
            'CompanyUpdatedAt': datetime.now()
        }
        
        # Remove empty values
        update_data = {k: v for k, v in update_data.items() if v}
        
        # Update company profile
        result = IRS_USERS.update_one(
            {"_id": ObjectId(user_id)},
            {"$set": update_data}
        )
        
        if result.modified_count > 0:
            logger.info(f"Company profile updated for user {user_id}")
            flash("Company details updated successfully!", "success")
        else:
            flash("No changes made to company details", "info")
            
        return redirect(url_for('hr_profile'))
        
    except Exception as e:
        logger.error(f"Error updating company profile: {e}")
        flash("Error updating company details", "error")
        return redirect(url_for('hr_profile'))

@app.route('/upload_hr_profile_picture', methods=['POST'])
def upload_hr_profile_picture():
    """Upload and update HR profile picture"""
    try:
        if 'user_id' not in session or session.get('role') != 'employer':
            return jsonify({"success": False, "message": "Access denied"}), 403
        
        if 'profile_picture' not in request.files:
            return jsonify({"success": False, "message": "No file selected"}), 400
        
        file = request.files['profile_picture']
        
        if file.filename == '':
            return jsonify({"success": False, "message": "No file selected"}), 400
        
        if not allowed_profile_file(file.filename):
            return jsonify({"success": False, "message": "Invalid file type. Only PNG, JPG, JPEG, GIF allowed."}), 400
        
        user_id = session['user_id']
        
        # Generate unique filename
        file_extension = file.filename.rsplit('.', 1)[1].lower()
        filename = f"hr_profile_{user_id}_{uuid.uuid4().hex}.{file_extension}"
        filepath = os.path.join(PROFILE_UPLOAD_FOLDER, filename)
        
        try:
            # Save and resize image (same logic as applicant profile picture)
            file.save(filepath)
            
            # Resize image to standard size (400x400)
            with Image.open(filepath) as img:
                # Convert to RGB if necessary
                if img.mode in ('RGBA', 'P'):
                    img = img.convert('RGB')
                
                # Resize maintaining aspect ratio
                img.thumbnail((400, 400), Image.Resampling.LANCZOS)
                
                # Create square image with white background
                if img.size[0] != img.size[1]:
                    size = max(img.size)
                    background = Image.new('RGB', (size, size), (255, 255, 255))
                    background.paste(img, ((size - img.size[0]) // 2, (size - img.size[1]) // 2))
                    img = background
                
                img.save(filepath, 'JPEG', quality=90)
            
            # Remove old profile picture if exists
            old_user = IRS_USERS.find_one({"_id": ObjectId(user_id)}, {"ProfilePicture": 1})
            if old_user and old_user.get('ProfilePicture'):
                old_path = os.path.join('static', old_user['ProfilePicture'].lstrip('/'))
                if os.path.exists(old_path):
                    os.remove(old_path)
            
            # Update user profile
            profile_picture_url = f"/static/profile_pictures/{filename}"
            result = IRS_USERS.update_one(
                {"_id": ObjectId(user_id)},
                {"$set": {"ProfilePicture": profile_picture_url, "ProfileUpdatedAt": datetime.now()}}
            )
            
            if result.modified_count > 0:
                logger.info(f"HR profile picture updated for user {user_id}")
                return jsonify({
                    "success": True,
                    "message": "Profile picture updated successfully",
                    "image_url": profile_picture_url
                })
            else:
                return jsonify({"success": False, "message": "Failed to update profile picture"}), 500
                
        except Exception as e:
            logger.error(f"Error processing HR profile picture: {e}")
            # Clean up file if there was an error
            if os.path.exists(filepath):
                os.remove(filepath)
            return jsonify({"success": False, "message": "Error processing image"}), 500
        
    except Exception as e:
        logger.error(f"Error in upload_hr_profile_picture: {e}")
        return jsonify({"success": False, "message": "Server error"}), 500






@app.route('/HR_Homepage', methods=['GET', 'POST'])
def HR_Homepage():
    """Redirect to enhanced HR dashboard"""
    try:
        if 'user_id' not in session or session.get('role') != 'employer':
            flash("Access denied. Employer login required.", "error")
            return redirect(url_for('loginpage'))
        
        return redirect(url_for('HR_Dashboard_Enhanced'))
    except Exception as e:
        logger.error(f"Error in HR_Homepage: {e}")
        return redirect(url_for('HR_Dashboard_Enhanced'))
    


# API endpoints for real-time dashboard updates
@app.route('/api/dashboard_stats_hr')
def api_dashboard_stats_hr():
    """API endpoint for real-time HR dashboard statistics"""
    try:
        if 'user_id' not in session or session.get('role') != 'employer':
            return jsonify({"error": "Access denied"}), 403
        
        current_date = datetime.now()
        today_start = current_date.replace(hour=0, minute=0, second=0, microsecond=0)
        week_ago = current_date - timedelta(days=7)
        
        stats = {
            'total_jobs': JOBS.count_documents({}),
            'active_jobs': JOBS.count_documents({"Status": "Open"}),
            'total_applications': Applied_EMP.count_documents({}),
            'new_applications_today': Applied_EMP.count_documents({
                "applied_at": {"$gte": today_start}
            }),
            'new_applications_week': Applied_EMP.count_documents({
                "applied_at": {"$gte": week_ago}
            }),
            'interviews_scheduled': Applied_EMP.count_documents({
                "status": {"$in": ["interviewed", "interview_scheduled"]}
            }),
            'timestamp': current_date.isoformat()
        }
        
        return jsonify(stats)
        
    except Exception as e:
        logger.error(f"Error fetching HR dashboard stats: {e}")
        return jsonify({"error": "Failed to fetch stats"}), 500


@app.route('/api/dashboard_stats_applicant')
def api_dashboard_stats_applicant():
    """API endpoint for real-time applicant dashboard statistics"""
    try:
        if 'user_id' not in session or session.get('role') != 'applicant':
            return jsonify({"error": "Access denied"}), 403
        
        user_id = session['user_id']
        current_date = datetime.now()
        week_ago = current_date - timedelta(days=7)
        month_ago = current_date - timedelta(days=30)
        
        try:
            # Basic application statistics
            stats = {
                'total_applications': Applied_EMP.count_documents({"user_id": ObjectId(user_id)}),
                'pending_applications': Applied_EMP.count_documents({
                    "user_id": ObjectId(user_id),
                    "status": "pending"
                }),
                'applications_this_week': Applied_EMP.count_documents({
                    "user_id": ObjectId(user_id),
                    "applied_at": {"$gte": week_ago}
                }),
                'applications_this_month': Applied_EMP.count_documents({
                    "user_id": ObjectId(user_id),
                    "applied_at": {"$gte": month_ago}
                }),
                'interviews_scheduled': Applied_EMP.count_documents({
                    "user_id": ObjectId(user_id),
                    "status": {"$in": ["interviewed", "interview_scheduled"]}
                }),
                'applications_shortlisted': Applied_EMP.count_documents({
                    "user_id": ObjectId(user_id),
                    "status": "shortlisted"
                }),
                'applications_under_review': Applied_EMP.count_documents({
                    "user_id": ObjectId(user_id),
                    "status": "under_review"
                }),
                'applications_rejected': Applied_EMP.count_documents({
                    "user_id": ObjectId(user_id),
                    "status": {"$in": ["rejected", "not_selected"]}
                }),
                'timestamp': current_date.isoformat()
            }
            
            # Calculate average scores and AI metrics
            avg_pipeline = Applied_EMP.aggregate([
                {"$match": {"user_id": ObjectId(user_id)}},
                {"$group": {
                    "_id": None,
                    "avg_match_score": {"$avg": "$Matching_percentage"},
                    "avg_semantic_score": {"$avg": "$semantic_similarity"},
                    "ai_enhanced_count": {
                        "$sum": {
                            "$cond": [{"$ne": ["$enhanced_version", None]}, 1, 0]
                        }
                    }
                }}
            ])
            
            avg_results = list(avg_pipeline)
            if avg_results:
                result = avg_results[0]
                
                # Safe rounding with None checks
                avg_match = result.get('avg_match_score') or 0
                avg_semantic = result.get('avg_semantic_score') or 0
                
                
                stats.update({
                    'avg_match_score': round(float(avg_match), 1),
                    'avg_semantic_score': round(float(avg_semantic), 1), 
                    
                    'ai_enhanced_applications': result.get('ai_enhanced_count', 0)
                })
            
            # Calculate success rates and performance metrics
            if stats['total_applications'] > 0:
                stats['interview_success_rate'] = round(
                    (stats['interviews_scheduled'] / stats['total_applications']) * 100, 1
                )
                stats['shortlist_rate'] = round(
                    (stats['applications_shortlisted'] / stats['total_applications']) * 100, 1
                )
                stats['rejection_rate'] = round(
                    (stats['applications_rejected'] / stats['total_applications']) * 100, 1
                )
            else:
                stats.update({
                    'interview_success_rate': 0,
                    'shortlist_rate': 0,
                    'rejection_rate': 0
                })
            
            # Profile completion status
            try:
                user_profile = IRS_USERS.find_one({"_id": ObjectId(user_id)})
                user_resume = resumeFetchedData.find_one({"UserId": ObjectId(user_id)})
                
                completion_factors = {
                    'basic_info': bool(user_profile and user_profile.get('Name') and user_profile.get('Email')),
                    'resume_uploaded': bool(user_resume),
                    'skills_extracted': bool(user_resume and user_resume.get('SKILLS') and len(user_resume.get('SKILLS', [])) > 0),
                    'experience_data': bool(user_resume and (user_resume.get('YEARS OF EXPERIENCE') or user_resume.get('CALCULATED_EXPERIENCE_YEARS', 0) > 0)),
                    'education_data': bool(user_resume and user_resume.get('EDUCATION') and len(user_resume.get('EDUCATION', [])) > 0)
                }
                
                completed_factors = sum(completion_factors.values())
                stats['profile_completion'] = round((completed_factors / len(completion_factors)) * 100)
                stats['resume_count'] = resumeFetchedData.count_documents({"UserId": ObjectId(user_id)})
                
            except Exception as e:
                logger.error(f"Error calculating profile completion for user {user_id}: {e}")
                stats.update({
                    'profile_completion': 0,
                    'resume_count': 0
                })
            
            # Recent activity trends
            try:
                # Applications trend (this week vs last week)
                last_week_start = week_ago - timedelta(days=7)
                last_week_apps = Applied_EMP.count_documents({
                    "user_id": ObjectId(user_id),
                    "applied_at": {"$gte": last_week_start, "$lt": week_ago}
                })
                
                if last_week_apps > 0:
                    trend_percentage = round(
                        ((stats['applications_this_week'] - last_week_apps) / last_week_apps) * 100, 1
                    )
                    stats['application_trend'] = {
                        'percentage': trend_percentage,
                        'direction': 'up' if trend_percentage > 0 else 'down' if trend_percentage < 0 else 'stable'
                    }
                else:
                    stats['application_trend'] = {
                        'percentage': 100 if stats['applications_this_week'] > 0 else 0,
                        'direction': 'up' if stats['applications_this_week'] > 0 else 'stable'
                    }
                    
            except Exception as e:
                logger.error(f"Error calculating trends for user {user_id}: {e}")
                stats['application_trend'] = {'percentage': 0, 'direction': 'stable'}
            
            # Performance insights
            performance_level = 'beginner'
            if stats['total_applications'] > 10:
                if stats['avg_match_score'] > 80 and stats['interview_success_rate'] > 15:
                    performance_level = 'excellent'
                elif stats['avg_match_score'] > 70 and stats['interview_success_rate'] > 10:
                    performance_level = 'good'
                elif stats['avg_match_score'] > 60:
                    performance_level = 'average'
                else:
                    performance_level = 'needs_improvement'
            
            stats['performance_level'] = performance_level
            
            logger.info(f"Dashboard stats calculated for user {user_id}")
            return jsonify(stats)
            
        except Exception as e:
            logger.error(f"Database error calculating stats for user {user_id}: {e}")
            return jsonify({
                "error": "Failed to calculate statistics",
                "total_applications": 0,
                "pending_applications": 0,
                "timestamp": current_date.isoformat()
            }), 500
        
    except Exception as e:
        logger.error(f"Error in api_dashboard_stats_applicant: {e}")
        return jsonify({"error": "Server error"}), 500


# Additional helper endpoint for detailed application analytics
@app.route('/api/application_analytics/<user_id>')
def get_application_analytics(user_id):
    """Get detailed analytics for an applicant"""
    try:
        if 'user_id' not in session or session.get('role') != 'applicant':
            return jsonify({"error": "Access denied"}), 403
        
        # Verify user can only access their own data
        if session['user_id'] != user_id:
            return jsonify({"error": "Access denied"}), 403
        
        try:
            ObjectId(user_id)
        except:
            return jsonify({"error": "Invalid user ID"}), 400
        
        # Get application history with timeline
        applications = list(Applied_EMP.find(
            {"user_id": ObjectId(user_id)},
            {
                "applied_at": 1, "Matching_percentage": 1, "status": 1,
                "semantic_similarity": 1,
                "job_id": 1, "enhanced_version": 1
            }
        ).sort([("applied_at", 1)]))
        
        # Process applications for timeline and trends
        timeline_data = []
        monthly_stats = {}
        status_distribution = {}
        
        for app in applications:
            try:
                app_date = app.get('applied_at', datetime.now())
                month_key = app_date.strftime('%Y-%m')
                
                # Timeline data point
                timeline_data.append({
                    'date': app_date.isoformat(),
                    'match_score': app.get('Matching_percentage', 0),
                    'semantic_score': app.get('semantic_similarity', 0),
                    'status': app.get('status', 'pending'),
                    'ai_enhanced': bool(app.get('enhanced_version'))
                })
                
                # Monthly aggregation
                if month_key not in monthly_stats:
                    monthly_stats[month_key] = {
                        'count': 0,
                        'total_match_score': 0,
                        'ai_enhanced_count': 0
                    }
                
                monthly_stats[month_key]['count'] += 1
                monthly_stats[month_key]['total_match_score'] += app.get('Matching_percentage', 0)
                if app.get('enhanced_version'):
                    monthly_stats[month_key]['ai_enhanced_count'] += 1
                
                # Status distribution
                status = app.get('status', 'pending')
                status_distribution[status] = status_distribution.get(status, 0) + 1
                
            except Exception as e:
                logger.error(f"Error processing application in analytics: {e}")
                continue
        
        # Calculate monthly averages
        monthly_averages = {}
        for month, stats in monthly_stats.items():
            monthly_averages[month] = {
                'applications': stats['count'],
                'avg_match_score': round(stats['total_match_score'] / stats['count'], 1) if stats['count'] > 0 else 0,
                'ai_enhanced_percentage': round((stats['ai_enhanced_count'] / stats['count']) * 100, 1) if stats['count'] > 0 else 0
            }
        
        analytics_data = {
            'timeline': timeline_data,
            'monthly_summary': monthly_averages,
            'status_distribution': status_distribution,
            'total_applications': len(applications),
            'data_range': {
                'start': applications[0]['applied_at'].isoformat() if applications else None,
                'end': applications[-1]['applied_at'].isoformat() if applications else None
            },
            'timestamp': datetime.now().isoformat()
        }
        
        return jsonify(analytics_data)
        
    except Exception as e:
        logger.error(f"Error in get_application_analytics: {e}")
        return jsonify({"error": "Analytics calculation failed"}), 500




@app.route("/Company_Job_Posting")
def Company_Job_Posting():
    try:
        if 'user_id' not in session or session.get('role') != 'employer':
            flash("Access denied. Employer login required.", "error")
            return redirect(url_for('loginpage'))
        
        return render_template("Company_Job_Posting.html")
    except Exception as e:
        logger.error(f"Error in Company_Job_Posting: {e}")
        return render_template("Company_Job_Posting.html", errorMsg="An error occurred")


        

@app.route('/Company_Candidates')
def Company_Candidates():
    """Enhanced candidate listing with Predictive Job Fit Analysis"""
    try:
        if 'user_id' not in session or session.get('role') != 'employer':
            flash("Access denied. Employer login required.", "error")
            return redirect(url_for('loginpage'))
        
        current_hr_user_id = session['user_id']
        
        # Enhanced debugging
        logger.info(f"=== Company_Candidates Debug ===")
        logger.info(f"HR User ID: {current_hr_user_id}")
        
        # Get filter parameters
        selected_job_id = request.args.get('job_id', '')
        selected_status = request.args.get('status', 'all')
        
        logger.info(f"Filters - Job ID: {selected_job_id}, Status: {selected_status}")
        
        # Get HR's jobs with FALLBACK for missing created_by
        hr_jobs_query = {"created_by": ObjectId(current_hr_user_id)}
        hr_jobs = list(JOBS.find(hr_jobs_query, {"_id": 1, "Job_Profile": 1}))
        hr_job_ids = [job['_id'] for job in hr_jobs]
        
        logger.info(f"Jobs with created_by field: {len(hr_job_ids)}")
        
        # CRITICAL FIX: Fallback if no jobs have created_by
        if not hr_job_ids:
            logger.warning("No jobs found with created_by field!")
            
            # Check if ANY jobs exist
            total_jobs = JOBS.count_documents({})
            logger.info(f"Total jobs in database: {total_jobs}")
            
            if total_jobs > 0:
                logger.warning("Using ALL jobs as fallback (legacy data)")
                flash("⚠️ Showing all jobs. Run /HR1/quick_fix_all to fix ownership.", "warning")
                hr_jobs = list(JOBS.find({}, {"_id": 1, "Job_Profile": 1}))
                hr_job_ids = [job['_id'] for job in hr_jobs]
            else:
                logger.error("No jobs exist in database at all")
                flash("No jobs found. Please create a job first.", "info")
        
        if not hr_job_ids:
            logger.warning("No jobs exist in database")
            return render_empty_template()
        
        # Build application query
        applications_query = {"job_id": {"$in": hr_job_ids}}
        
        # Apply job filter if specified
        if selected_job_id:
            try:
                applications_query["job_id"] = ObjectId(selected_job_id)
                logger.info(f"Filtering by specific job: {selected_job_id}")
            except Exception as e:
                logger.error(f"Invalid job_id format: {selected_job_id} - {e}")
                flash("Invalid job ID format", "error")
        
        # Apply status filter
        if selected_status and selected_status != 'all':
            logger.info(f"Filtering by status: {selected_status}")
            applications_query["status"] = selected_status
        else:
            logger.info("Showing ALL statuses (no status filter)")
        
        logger.info(f"Final Query: {applications_query}")
        
        # Get applications
        applications = list(Applied_EMP.find(applications_query))
        logger.info(f"✅ Found {len(applications)} applications")
        
        # Get selected job title if filtering by job
        selected_job_title = ''
        if selected_job_id:
            try:
                job = JOBS.find_one({"_id": ObjectId(selected_job_id)}, {"Job_Profile": 1})
                if job:
                    selected_job_title = job.get('Job_Profile', '')
            except:
                pass
        
        # Process candidates
        candidates = []
        
        # Predictive fit counters
        excellent_fit_count = 0
        good_fit_count = 0
        moderate_fit_count = 0
        weak_fit_count = 0
        
        for app in applications:
            try:
                # Get user/resume data
                user = IRS_USERS.find_one({"_id": app.get('user_id')}, {"Name": 1})
                resume = resumeFetchedData.find_one({"_id": app.get('resume_id')})
                job = JOBS.find_one({"_id": app.get('job_id')}, {"Job_Profile": 1})
                
                if not user or not job:
                    continue
                
                # Get user documents
                user_documents = []
                try:
                    user_documents = list(mongo.db.USER_DOCUMENTS.find(
                        {"UserId": app.get('user_id')}
                    ).sort("UploadedAt", -1))
                except Exception as e:
                    logger.error(f"Error fetching documents for user {app.get('user_id')}: {e}")

                # Calculate combined match score using weighted formula
                match_score = get_match_score(app)

                # Get traditional score
                traditional_score = app.get('Matching_percentage', 0)
                if isinstance(traditional_score, dict):
                    traditional_score = traditional_score.get('overall_score', 0)
                try:
                    traditional_score = float(traditional_score)
                except:
                    traditional_score = 0

                # Get semantic score
                semantic_score = app.get('semantic_similarity')
                if semantic_score is not None:
                    try:
                        semantic_score = float(semantic_score)
                    except (ValueError, TypeError):
                        semantic_score = None
                
                # Extract component scores and feedback
                component_scores = None
                candidate_feedback = None
                detailed_match_info = None
                
                matching_data = app.get('Matching_percentage')
                if matching_data and isinstance(matching_data, dict):
                    component_scores = matching_data.get('component_scores', {})
                    candidate_feedback = matching_data.get('candidate_feedback', {})
                    detailed_match_info = matching_data.get('detailed_match', {})

                    # Fallback attempts for detailed_match_info
                    if not detailed_match_info:
                        detailed_match_info = app.get('detailed_match_info', {})
                    
                    if not candidate_feedback:
                        candidate_feedback = app.get('candidate_feedback', {})
                    
                    if not detailed_match_info and 'detailed_analysis' in matching_data:
                        detailed_match_info = matching_data.get('detailed_analysis', {})
                
                # Get status
                app_status = app.get('status', 'pending')
                if not app_status:
                    app_status = 'pending'
                
                # Determine predictive fit category
                if match_score >= 80:
                    excellent_fit_count += 1
                    fit_category = 'excellent'
                elif match_score >= 65:
                    good_fit_count += 1
                    fit_category = 'good'
                elif match_score >= 50:
                    moderate_fit_count += 1
                    fit_category = 'moderate'
                else:
                    weak_fit_count += 1
                    fit_category = 'weak'

                
                
                # Build candidate object
                candidate = {
                    'application_id': str(app['_id']),
                    'candidate_id': str(app.get('user_id')),
                    'user_id': str(app.get('user_id')),
                    'resume_id': str(app.get('resume_id', '')),
                    'job_id': str(app.get('job_id')),
                    'name': user.get('Name', 'Unknown'),
                    'job_title': job.get('Job_Profile', 'Unknown Position'),
                    'match_score': round(match_score, 1),
                    'traditional_score': round(traditional_score, 1),
                    'semantic_similarity': round(semantic_score * 100, 1) if semantic_score else None,
                    'combined_score': round(match_score, 1),
                    'fit_category': fit_category,  # NEW: Predictive fit category
                    'status': app_status,
                    'applied_date': app.get('applied_at', datetime.now()),
                    'ai_enhanced': semantic_score is not None and semantic_score > 0,
                    'skills': resume.get('SKILLS', []) if resume else [],
                    'experience': resume.get('WORKED AS', []) if resume else [],
                    'education': resume.get('EDUCATION', []) if resume else [],
                    'user_documents': user_documents,
                    'document_count': len(user_documents),
                    'component_scores': component_scores,
                    'candidate_feedback': candidate_feedback,
                    'detailed_match_info': detailed_match_info
                }
                
                candidates.append(candidate)
                
            except Exception as e:
                logger.error(f"Error processing application {app.get('_id')}: {e}")
                import traceback
                logger.error(traceback.format_exc())
                continue
        
        # Sort candidates by match_score (DESCENDING - highest first)
        candidates.sort(key=lambda x: x['match_score'], reverse=True)
        
        # Assign ranks (with tie handling)
        current_rank = 1
        for i, candidate in enumerate(candidates):
            if i > 0 and candidates[i]['match_score'] == candidates[i-1]['match_score']:
                candidate['rank'] = candidates[i-1]['rank']
            else:
                candidate['rank'] = current_rank
                current_rank += 1
        
        logger.info(f"✅ Successfully processed {len(candidates)} candidates")
        
        # Debug: Show top 5 candidates
        if candidates:
            logger.info("📊 Top 5 Candidates (by match_score):")
            for i, c in enumerate(candidates[:5], 1):
                logger.info(f"   #{i}: {c['name']} - Score: {c['match_score']}% - Fit: {c['fit_category']}")
        
        # Calculate statistics with predictive fit analysis
        stats = {
            'total_candidates': len(candidates),
            'excellent_fit': excellent_fit_count,  # NEW: 80-100%
            'good_fit': good_fit_count,            # NEW: 65-79%
            'moderate_fit': moderate_fit_count,    # NEW: 50-64%
            'weak_fit': weak_fit_count,            # NEW: Below 50%
            'under_review': sum(1 for c in candidates if c['status'] in ['under_review', 'pending']),
            'shortlisted': sum(1 for c in candidates if c['status'] == 'shortlisted'),
            'interview_pending': sum(1 for c in candidates if c['status'] == 'interview_pending'),
            'interviewed': sum(1 for c in candidates if c['status'] == 'interviewed')
        }
        
        # Log predictive fit distribution
        logger.info(f"📈 Predictive Fit Distribution:")
        logger.info(f"   🌟 Excellent (80-100%): {excellent_fit_count}")
        logger.info(f"   ✨ Good (65-79%): {good_fit_count}")
        logger.info(f"   💡 Moderate (50-64%): {moderate_fit_count}")
        logger.info(f"   📚 Weak (<50%): {weak_fit_count}")
        
        # Get all jobs for filter dropdown
        jobs = hr_jobs
        
        # Available statuses for filter
        available_statuses = [
            {'value': 'all', 'label': 'All Statuses'},
            {'value': 'pending', 'label': 'Pending'},
            {'value': 'under_review', 'label': 'Under Review'},
            {'value': 'shortlisted', 'label': 'Shortlisted'},
            {'value': 'interview_pending', 'label': 'Interview Pending'},
            {'value': 'interview_scheduled', 'label': 'Interview Scheduled'},
            {'value': 'interviewed', 'label': 'Interviewed'},
            {'value': 'not_selected', 'label': 'Not Selected'}
        ]
        
        return render_template('Company_Candidates.html',
                             candidates=candidates,
                             jobs=jobs,
                             stats=stats,
                             selected_job_id=selected_job_id,
                             selected_job_title=selected_job_title,
                             selected_status=selected_status,
                             available_statuses=available_statuses)
    except Exception as e:
        logger.error(f"Error in Company_Candidates: {e}")
        import traceback
        logger.error(traceback.format_exc())
        return render_empty_template(errorMsg="An error occurred loading candidates")


def render_empty_template(errorMsg=None):
    """Helper to render empty template"""
    available_statuses = [
        {'value': 'all', 'label': 'All Statuses'},
        {'value': 'pending', 'label': 'Pending'},
        {'value': 'under_review', 'label': 'Under Review'},
        {'value': 'shortlisted', 'label': 'Shortlisted'},
        {'value': 'interview_pending', 'label': 'Interview Pending'},
        {'value': 'interview_scheduled', 'label': 'Interview Scheduled'},
        {'value': 'interviewed', 'label': 'Interviewed'},
        {'value': 'not_selected', 'label': 'Not Selected'}
    ]
    
    return render_template('Company_Candidates.html',
                         candidates=[],
                         jobs=[],
                         stats={'total_candidates': 0, 'under_review': 0, 'shortlisted': 0, 
                               'interview_pending': 0, 'interviewed': 0},
                         selected_job_id='',
                         selected_job_title='',
                         selected_status='all',
                         available_statuses=available_statuses,
                         errorMsg=errorMsg)

@app.route('/debug_candidates')
def debug_candidates():
    """Temporary debug route - REMOVE after fixing"""
    try:
        if 'user_id' not in session:
            return "Please login first"
        
        current_hr_user_id = session['user_id']
        
        debug_info = {
            "current_user": {
                "user_id": current_hr_user_id,
                "role": session.get('role'),
                "name": session.get('user_name')
            }
        }
        
        # Check jobs with created_by
        jobs_with_created_by = list(JOBS.find(
            {"created_by": ObjectId(current_hr_user_id)},
            {"_id": 1, "Job_Profile": 1, "created_by": 1}
        ))
        debug_info["jobs_with_created_by"] = len(jobs_with_created_by)
        debug_info["job_details_with_created_by"] = [
            {
                "id": str(j['_id']), 
                "title": j.get('Job_Profile'),
                "created_by": str(j.get('created_by'))
            } for j in jobs_with_created_by
        ]
        
        # Check ALL jobs (without filter)
        all_jobs = list(JOBS.find({}, {"_id": 1, "Job_Profile": 1, "created_by": 1}))
        debug_info["total_jobs_in_db"] = len(all_jobs)
        debug_info["all_jobs"] = [
            {
                "id": str(j['_id']), 
                "title": j.get('Job_Profile'),
                "has_created_by": 'created_by' in j,
                "created_by": str(j.get('created_by')) if 'created_by' in j else None
            } for j in all_jobs
        ]
        
        # Check applications
        if jobs_with_created_by:
            job_ids = [j['_id'] for j in jobs_with_created_by]
            apps = list(Applied_EMP.find(
                {"job_id": {"$in": job_ids}},
                {"_id": 1, "job_id": 1, "user_id": 1, "status": 1, "Matching_percentage": 1}
            ))
            debug_info["applications_for_hr_jobs"] = len(apps)
            debug_info["application_samples"] = [
                {
                    "app_id": str(a['_id']),
                    "job_id": str(a.get('job_id')),
                    "user_id": str(a.get('user_id')),
                    "status": a.get('status'),
                    "score": a.get('Matching_percentage')
                } for a in apps[:5]
            ]
        else:
            debug_info["applications_for_hr_jobs"] = "N/A - no jobs with created_by"
        
        # Check ALL applications (without job filter)
        all_apps = list(Applied_EMP.find({}, {"_id": 1, "job_id": 1, "status": 1}).limit(10))
        debug_info["total_applications_in_db"] = Applied_EMP.count_documents({})
        debug_info["sample_applications"] = [
            {
                "app_id": str(a['_id']),
                "job_id": str(a.get('job_id')),
                "status": a.get('status')
            } for a in all_apps
        ]
        
        # Check unique statuses
        all_statuses = Applied_EMP.distinct("status")
        debug_info["unique_statuses_in_db"] = all_statuses
        
        # Check if jobs exist without created_by
        jobs_without_created_by = JOBS.count_documents({"created_by": {"$exists": False}})
        debug_info["jobs_without_created_by"] = jobs_without_created_by
        
        # Return formatted debug info
        import json
        return f"""
        <html>
        <head>
            <title>Debug Info</title>
            <style>
                body {{ font-family: monospace; padding: 20px; background: #1e1e1e; color: #d4d4d4; }}
                h2 {{ color: #4ec9b0; }}
                h3 {{ color: #dcdcaa; margin-top: 20px; }}
                pre {{ background: #252526; padding: 15px; border-radius: 5px; overflow-x: auto; }}
                .error {{ color: #f48771; }}
                .success {{ color: #4ec9b0; }}
                .warning {{ color: #dcdcaa; }}
            </style>
        </head>
        <body>
            <h2>🔍 Candidate Debug Information</h2>
            <pre>{json.dumps(debug_info, indent=2)}</pre>
            
            <h3 class="{'error' if debug_info['jobs_with_created_by'] == 0 else 'success'}">
                Diagnosis:
            </h3>
            <pre>
{'❌ PROBLEM FOUND: No jobs have created_by field!' if debug_info['jobs_with_created_by'] == 0 else '✅ Jobs have created_by field'}
{f'❌ PROBLEM: {jobs_without_created_by} jobs missing created_by field' if jobs_without_created_by > 0 else ''}
{'✅ Applications exist in database' if debug_info['total_applications_in_db'] > 0 else '⚠️ No applications in database at all'}
            </pre>
            
            <h3>🔧 Recommended Fixes:</h3>
            <pre>
1. If jobs missing created_by field:
   - Visit: /HR1/fix_existing_jobs_once
   
2. If no applications exist:
   - Apply to a job first as a candidate
   
3. If applications exist but not showing:
   - Check status values match between DB and code
   - Current statuses in DB: {all_statuses}
            </pre>
            
            <a href="/HR1/Company_Candidates" style="color: #4ec9b0;">← Back to Candidates Page</a>
        </body>
        </html>
        """
        
    except Exception as e:
        import traceback
        return f"""
        <html>
        <body style="font-family: monospace; padding: 20px; background: #1e1e1e; color: #f48771;">
            <h2>❌ Debug Error</h2>
            <pre>{str(e)}</pre>
            <pre>{traceback.format_exc()}</pre>
        </body>
        </html>
        """
                             
@app.route("/shortlist_candidate/<application_id>", methods=["POST"])
def shortlist_candidate(application_id):
    """Shortlist candidate WITHOUT automatically scheduling interview"""
    if 'user_id' not in session or session.get('role') != 'employer':
        return jsonify({"error": "Unauthorized"}), 403

    try:
        result = Applied_EMP.update_one(
            {"_id": ObjectId(application_id)},
            {"$set": {
                "status": "shortlisted",  # Keep as shortlisted
                "shortlisted_at": datetime.now(),
                "shortlisted_by": session.get('user_id')
            }}
        )
        
        if result.modified_count > 0:
            return jsonify({
                "success": True, 
                "message": "Candidate shortlisted successfully",
                "new_status": "shortlisted"
            })
        else:
            return jsonify({"error": "Application not found"}), 404
            
    except Exception as e:
        logger.error(f"Error shortlisting candidate {application_id}: {e}")
        return jsonify({"error": "Something went wrong"}), 500


@app.route('/export_candidates')
def export_candidates():
    """Export candidates to CSV with optional job filtering"""
    try:
        if 'user_id' not in session or session.get('role') != 'employer':
            flash("Access denied. Employer login required.", "error")
            return redirect(url_for('loginpage'))
        
        current_hr_user_id = session['user_id']
        
        # Get filter parameter
        selected_job_id = request.args.get('job_id', '')
        
        # Get HR's jobs
        hr_jobs_query = {"created_by": ObjectId(current_hr_user_id)}
        hr_jobs = list(JOBS.find(hr_jobs_query, {"_id": 1, "Job_Profile": 1}))
        hr_job_ids = [job['_id'] for job in hr_jobs]
        
        # Fallback if no jobs have created_by
        if not hr_job_ids:
            hr_jobs = list(JOBS.find({}, {"_id": 1, "Job_Profile": 1}))
            hr_job_ids = [job['_id'] for job in hr_jobs]
        
        if not hr_job_ids:
            flash("No jobs found.", "error")
            return redirect(url_for('Company_Candidates'))
        
        # Build query
        applications_query = {"job_id": {"$in": hr_job_ids}}
        
        if selected_job_id:
            try:
                applications_query["job_id"] = ObjectId(selected_job_id)
            except Exception as e:
                logger.error(f"Invalid job_id format: {selected_job_id}")
                flash("Invalid job ID format", "error")
                return redirect(url_for('Company_Candidates'))
        
        # Get applications
        applications = list(Applied_EMP.find(applications_query))
        
        # Process candidates for export
        export_data = []
        
        for app in applications:
            try:
                user = IRS_USERS.find_one({"_id": app.get('user_id')}, {"Name": 1, "Email": 1, "Phone": 1})
                resume = resumeFetchedData.find_one({"_id": app.get('resume_id')})
                job = JOBS.find_one({"_id": app.get('job_id')}, {"Job_Profile": 1})
                
                if not user or not job:
                    continue
                
                # Get scores
                traditional_score = app.get('Matching_percentage', 0)
                if isinstance(traditional_score, dict):
                    traditional_score = traditional_score.get('overall_score', 0)
                
                try:
                    traditional_score = float(traditional_score)
                except:
                    traditional_score = 0
                
                semantic_score = app.get('semantic_similarity')
                if semantic_score is not None:
                    try:
                        semantic_score = float(semantic_score) * 100
                    except:
                        semantic_score = None
                
                combined_score = app.get('combined_score')
                if combined_score is not None:
                    match_score = float(combined_score)
                elif semantic_score is not None and semantic_score > 0:
                    try:
                        from sentence_transformer_ranker import get_semantic_ranker
                        ranker = get_semantic_ranker()
                        if ranker.is_available():
                            match_score = ranker.calculate_combined_score(traditional_score, semantic_score)
                        else:
                            match_score = traditional_score
                    except:
                        match_score = traditional_score
                else:
                    match_score = traditional_score
                
                # Get status
                app_status = app.get('status', 'pending')
                if not app_status:
                    app_status = 'pending'
                
                # Extract skills, experience, education
                skills = ', '.join(resume.get('SKILLS', [])) if resume else ''
                experience = '; '.join(resume.get('WORKED AS', [])) if resume else ''
                education = '; '.join(resume.get('EDUCATION', [])) if resume else ''
                
                export_data.append({
                    'Rank': 0,  # Will be assigned after sorting
                    'Candidate Name': user.get('Name', 'Unknown'),
                    'Email': user.get('Email', 'N/A'),
                    
                    'Job Title': job.get('Job_Profile', 'Unknown'),
                    'Match Score (%)': round(match_score, 1),
                    'Traditional Score (%)': round(traditional_score, 1),
                    'Semantic Score (%)': round(semantic_score, 1) if semantic_score else 'N/A',
                    'AI Enhanced': 'Yes' if semantic_score else 'No',
                    'Status': app_status.replace('_', ' ').title(),
                    'Applied Date': app.get('applied_at', datetime.now()).strftime('%Y-%m-%d'),
                    'Skills': skills,
                    'Experience': experience,
                    'Education': education
                })
                
            except Exception as e:
                logger.error(f"Error processing application for export: {e}")
                continue
        
        # Sort by match score (descending)
        export_data.sort(key=lambda x: x['Match Score (%)'], reverse=True)
        
        # Assign ranks
        for rank, candidate in enumerate(export_data, start=1):
            candidate['Rank'] = rank
        
        # Create CSV
        if not export_data:
            flash("No candidates found to export.", "warning")
            return redirect(url_for('Company_Candidates'))
        
        # Generate CSV in memory
        import csv
        from io import StringIO
        
        output = StringIO()
        writer = csv.DictWriter(output, fieldnames=export_data[0].keys())
        writer.writeheader()
        writer.writerows(export_data)
        
        # Create response
        response = make_response(output.getvalue())
        response.headers['Content-Type'] = 'text/csv'
        
        # Set filename based on filter
        if selected_job_id:
            job = JOBS.find_one({"_id": ObjectId(selected_job_id)}, {"Job_Profile": 1})
            job_name = job.get('Job_Profile', 'job').replace(' ', '_') if job else 'job'
            filename = f"candidates_{job_name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv"
        else:
            filename = f"all_candidates_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv"
        
        response.headers['Content-Disposition'] = f'attachment; filename={filename}'
        
        logger.info(f"✅ Exported {len(export_data)} candidates to CSV")
        return response
        
    except Exception as e:
        logger.error(f"Error exporting candidates: {e}")
        import traceback
        logger.error(traceback.format_exc())
        flash("An error occurred during export.", "error")
        return redirect(url_for('Company_Candidates'))


@app.route('/test')
def test():
    try:
        # Test database connection
        mongo.db.command('ping')
        return jsonify({"status": "success", "message": "Connection Successful"})
    except Exception as e:
        logger.error(f"Test endpoint error: {e}")
        return jsonify({"status": "error", "message": str(e)}), 500




@app.route('/applicant_upload_resume')
def applicant_upload_resume_redirect():
    """Redirect old resume upload page to new profile page"""
    return redirect('/profile')





# def load_all_models():
#     global bert_tokenizer, bert_model, new_custom_nlp, original_nlp
    
#     # 1. Load BERT model (highest priority)
#     try:
#         logger.info("Loading BERT resume NER model...")
#         model_name = "yashpwr/resume-ner-bert-v2"
#         bert_tokenizer = AutoTokenizer.from_pretrained(model_name)
#         bert_model = AutoModelForTokenClassification.from_pretrained(model_name)
#         logger.info("BERT model loaded successfully")
        
#         # Simple test
#         if test_bert_model_basic():
#             logger.info("✅ BERT model test passed")
#         else:
#             logger.error("❌ BERT model test failed")
            
#     except Exception as e:
#         logger.error(f"Failed to load BERT model: {e}")
#         bert_tokenizer = None
#         bert_model = None
    
#     # 2. Load new custom trained model (second priority)
#     try:
#         logger.info("Loading new custom trained model...")
#         new_custom_nlp = spacy.load('./resume_ner_model_v2')
#         logger.info("New custom trained model loaded successfully")
#     except Exception as e:
#         logger.warning(f"Could not load new custom model: {e}")
#         new_custom_nlp = None
    
#     # 3. Load original model (third priority)
#     try:
#         logger.info("Loading original Resume Parser model...")
#         original_nlp = spacy.load('assets/ResumeModel/output/model-best')
#         logger.info("Original Resume Parser model loaded successfully")
#     except Exception as e:
#         logger.error(f"Failed to load original SpaCy model: {e}")
#         original_nlp = None
    
# Add these imports at the top of your file
# At the top of your app.py, add this import



@app.route("/uploadResume", methods=['POST'])
def uploadResume():
    """
    Resume upload handler using trained NER model
    Extracts: Skills, Certifications, Education, Experience
    """
    try:
        # ========== AUTHENTICATION & FILE VALIDATION ==========
        if 'user_id' not in session or 'user_name' not in session:
            flash("Please log in first", "error")
            return redirect(url_for('loginpage'))

        if 'resume' not in request.files:
            flash("No file selected", "error")
            return redirect('/profile')

        file = request.files['resume']

        if not file or file.filename == '':
            flash("No file selected", "error")
            return redirect('/profile')

        if not allowedExtension(file.filename):
            flash("Invalid file format. Only PDF and DOCX files are allowed.", "error")
            return redirect('/profile')

        filename = secure_filename(file.filename)
        if not filename:
            flash("Invalid filename", "error")
            return redirect('/profile')

        # ========== FILE HANDLING ==========
        file_extension = os.path.splitext(filename)[1]
        base_name = os.path.splitext(filename)[0]
        clean_filename = base_name + file_extension
        
        timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
        internal_filename = timestamp + filename
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], internal_filename)

        try:
            file.save(filepath)
            logger.info(f"Resume uploaded: {internal_filename} by user {session['user_id']}")
        except Exception as e:
            logger.error(f"Error saving file {internal_filename}: {e}")
            flash("Error saving file", "error")
            return redirect('/profile')

        # ========== TEXT EXTRACTION ==========
        try:
            if filepath.lower().endswith('.pdf'):
                import fitz
                with fitz.open(filepath) as doc:
                    text_of_resume = "".join([page.get_text() for page in doc])
            elif filepath.lower().endswith('.docx'):
                import docx
                doc = docx.Document(filepath)
                text_of_resume = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            else:
                raise ValueError("Unsupported file format")

            if not text_of_resume or len(text_of_resume.strip()) < 50:
                raise ValueError("Resume text is too short or empty")

            logger.info(f"Resume text extracted: {len(text_of_resume)} characters")

        except Exception as e:
            logger.error(f"Error extracting text from {internal_filename}: {e}")
            if os.path.exists(filepath):
                os.remove(filepath)
            flash("Error processing resume. Please ensure it's a valid document.", "error")
            return redirect('/profile')

     # ========== NER MODEL EXTRACTION ==========
        try:
            value_name = None
            value_skills = []
            value_certificate = []
            value_education = []
            value_workedAs = []
            value_experience = []
            value_linkedin = None
            calculated_experience = 0.0
            resume_data_annotated = ""
            
            # === USE NEW MODEL for SKILLS AND JOB TITLES yONLY ===
            if nlp_new:
                logger.info("Using NEW model for skills only...")
                doc_new = nlp_new(text_of_resume)
                
                seen_skills = set()
                seen_jobs = set()
                
                for ent in doc_new.ents:
                    text = ent.text.strip()
                    if len(text) < 2:
                        continue
                    
                    if ent.label_ in ['TECHNICAL_SKILL', 'SOFT_SKILL', 'SKILLS']:
                        if text.lower() not in seen_skills:
                            seen_skills.add(text.lower())
                            value_skills.append(text)

                    elif ent.label_ in ['JOB_TITLE', 'WORKED AS']:
                        if text.lower() not in seen_jobs:
                            seen_jobs.add(text.lower())
                            value_workedAs.append(text)
            
            # === USE OLD MODEL for Everything Else ===
            if nlp_old:
                logger.info("Using OLD model for experience, education, certs, jobs...")
                doc_old = nlp_old(text_of_resume)
                
                seen_skills_old = set()
                seen_certs = set()
                seen_education = set()
                seen_jobs = set()
                
                for ent in doc_old.ents:
                    text = ent.text.strip()
                    if len(text) < 2:
                        continue

                    
                    if ent.label_ in ['TECHNICAL_SKILL', 'SOFT_SKILL', 'SKILLS']:
                        if text.lower() not in seen_skills_old:
                            seen_skills_old.add(text.lower())
                            value_skills.append(text)  # Merge with new model results
                    
                    
                    elif ent.label_ == 'CERTIFICATION':
                        if text.lower() not in seen_certs:
                            seen_certs.add(text.lower())
                            value_certificate.append(text)
                    
                    elif ent.label_ in ['EDUCATION', 'DEGREE']:
                        if text.lower() not in seen_education:
                            seen_education.add(text.lower())
                            value_education.append(text)
                    
                    elif ent.label_ == 'WORKED AS':
                        if text.lower() not in seen_jobs:
                            seen_jobs.add(text.lower())
                            value_workedAs.append(text)
                    
                    elif ent.label_ == 'YEARS OF EXPERIENCE':
                        value_experience.append(text)
                    
                    
                    
              
                
                logger.info(f"OLD model results - Skills: {len(seen_skills_old)}, Exp: {value_experience}, Edu: {value_education}, Certs: {value_certificate}")

            # === FALLBACK: If old model didn't find experience, try new model ===
            if not value_experience and nlp_new:
                logger.info("Old model didn't find experience, trying new model...")
                for ent in doc_new.ents:
                    if ent.label_ in ['YEARS_OF_EXP', 'YEARS OF EXPERIENCE']:
                        value_experience = [ent.text.strip()]
                        logger.info(f"New model extracted experience: {ent.text}")
                        break
            
            # === CALCULATE EXPERIENCE YEARS ===
            if value_experience:
                for exp_text in value_experience:
                    exp_str = str(exp_text).lower()
                    
                    # Extract years
                    years_match = re.search(r'(\d+(?:\.\d+)?)\s*(?:year|yr|y)', exp_str)
                    if years_match:
                        calculated_experience += float(years_match.group(1))
                    
                    # Extract months and convert to years
                    months_match = re.search(r'(\d+(?:\.\d+)?)\s*(?:month|mon|m)', exp_str)
                    if months_match:
                        months = float(months_match.group(1))
                        calculated_experience += months / 12.0
                    
                    # If no unit specified, assume years
                    if not years_match and not months_match:
                        number_match = re.search(r'(\d+(?:\.\d+)?)', exp_str)
                        if number_match:
                            calculated_experience += float(number_match.group(1))
                
                calculated_experience = round(calculated_experience, 2)
            
            # Convert to None if empty
            value_skills = value_skills if value_skills else None
            value_certificate = value_certificate if value_certificate else None
            value_education = value_education if value_education else None
            value_workedAs = value_workedAs if value_workedAs else None
            
            extraction_method = "Dual Model (New+Old)"
            
            # ========== LOG RAW EXTRACTION RESULTS ==========
            logger.info(f"=== RAW EXTRACTION RESULTS (BEFORE CLEANING) ===")
            logger.info(f"METHOD: {extraction_method}")
            logger.info(f"NAME: {value_name}")
            logger.info(f"SKILLS: {len(value_skills or [])} found")
            if value_skills:
                for i, skill in enumerate(value_skills, 1):
                    logger.info(f"  [{i}] {skill}")
            
            logger.info(f"CERTIFICATIONS: {len(value_certificate or [])} found")
            if value_certificate:
                for i, cert in enumerate(value_certificate, 1):
                    logger.info(f"  [{i}] {cert}")

            logger.info(f"EDUCATION: {len(value_education or [])} found")
            if value_education:
                for i, edu in enumerate(value_education, 1):
                    logger.info(f"  [{i}] {edu}")

            logger.info(f"WORKED AS: {len(value_workedAs or [])} found")
            if value_workedAs:
                for i, job in enumerate(value_workedAs, 1):
                    logger.info(f"  [{i}] {job}")

            logger.info(f"EXPERIENCE: {value_experience} (calculated: {calculated_experience} years)")
            if value_linkedin:
                logger.info(f"LINKEDIN: {value_linkedin}")
            logger.info(f"=== END RAW EXTRACTION ===")
            
            # ========== CLEAN THE DATA ==========
            logger.info("Cleaning extracted data...")
            cleaned_data = clean_resume_data({
                'Name': value_name,
                'LINKEDIN LINK': value_linkedin,
                'SKILLS': value_skills,
                'CERTIFICATION': value_certificate,
                'EDUCATION': value_education,
                'WORKED AS': value_workedAs,
                'YEARS OF EXPERIENCE': value_experience,
                'CALCULATED_EXPERIENCE_YEARS': calculated_experience
            })

            # Use cleaned data for database storage
            value_name = cleaned_data['Name']
            value_linkedin = cleaned_data['LINKEDIN LINK']
            value_skills = cleaned_data['SKILLS']
            value_certificate = cleaned_data['CERTIFICATION']
            value_education = cleaned_data['EDUCATION']
            value_workedAs = cleaned_data['WORKED AS']
            value_experience = cleaned_data['YEARS OF EXPERIENCE']
            calculated_experience = cleaned_data['CALCULATED_EXPERIENCE_YEARS']

            # ========== LOG CLEANED RESULTS ==========
            logger.info(f"=== CLEANED EXTRACTION RESULTS ===")
            logger.info(f"NAME: {value_name}")
            logger.info(f"SKILLS: {len(value_skills or [])} found")
            if value_skills:
                for i, skill in enumerate(value_skills, 1):
                    logger.info(f"  [{i}] {skill}")
            
            logger.info(f"CERTIFICATIONS: {len(value_certificate or [])} found")
            if value_certificate:
                for i, cert in enumerate(value_certificate, 1):
                    logger.info(f"  [{i}] {cert}")

            logger.info(f"EDUCATION: {len(value_education or [])} found")
            if value_education:
                for i, edu in enumerate(value_education, 1):
                    logger.info(f"  [{i}] {edu}")

            logger.info(f"WORKED AS: {len(value_workedAs or [])} found")
            if value_workedAs:
                for i, job in enumerate(value_workedAs, 1):
                    logger.info(f"  [{i}] {job}")

            logger.info(f"EXPERIENCE: {value_experience} (calculated: {calculated_experience} years)")
            if value_linkedin:
                logger.info(f"LINKEDIN: {value_linkedin}")
            logger.info(f"=== END CLEANED RESULTS ===")
            
        except Exception as e:
            logger.error(f"Error in dual model extraction: {e}")
            import traceback
            logger.error(traceback.format_exc())
            
            value_name = None
            value_linkedin = None
            value_skills = None
            value_certificate = None
            value_education = None
            value_workedAs = None
            value_experience = None
            calculated_experience = 0.0 
            resume_data_annotated = text_of_resume[:1000] if text_of_resume else ""
            extraction_method = "Extraction Error"

        # ========== DATABASE STORAGE ==========
        try:
            with open(filepath, "rb") as f:
                file_data = f.read()
            
            resume_document = {
                "UserId": ObjectId(session['user_id']),
                "Name": value_name,
                "LINKEDIN LINK": value_linkedin,
                "SKILLS": value_skills,
                "CERTIFICATION": value_certificate,
                "EDUCATION": value_education,
                "WORKED AS": value_workedAs,
                "YEARS OF EXPERIENCE": value_experience,
                "CALCULATED_EXPERIENCE_YEARS": calculated_experience,
                "Appear": 0,
                "ResumeTitle": clean_filename,
                "InternalFilename": internal_filename,
                "ResumeData": text_of_resume,
                "FileData": file_data,
                "UploadedAt": datetime.now(),
                "ExtractionMethod": extraction_method
            }
            
            result = resumeFetchedData.insert_one(resume_document)
            logger.info(f"Resume stored for user {session['user_id']} using {extraction_method}")

            if result is None:
                flash("Problem in Resume Data Storage", "error")
                return redirect('/profile')
            else:
                skill_count = len(value_skills) if value_skills else 0
                cert_count = len(value_certificate) if value_certificate else 0
                edu_count = len(value_education) if value_education else 0
                
                success_msg = f"Resume processed! "
                if value_name:
                    success_msg += f"Name: {value_name}. "
                success_msg += f"Found {skill_count} skills"
                if cert_count > 0:
                    success_msg += f", {cert_count} certifications"
                if edu_count > 0:
                    success_msg += f", {edu_count} education entries"
                success_msg += "."
                
                flash(success_msg, "success")
                return redirect('/profile')

        except Exception as e:
            logger.error(f"Database error: {e}")
            import traceback
            logger.error(traceback.format_exc())
            
            if os.path.exists(filepath):
                os.remove(filepath)
            flash("Error storing resume data", "error")
            return redirect('/profile')

    except Exception as e:
        logger.error(f"Unexpected error in uploadResume: {e}")
        import traceback
        logger.error(traceback.format_exc())
        flash("An unexpected error occurred", "error")
        return redirect('/profile')



@app.route('/upload_generated_resume_pdf', methods=['POST'])
def upload_generated_resume_pdf():
    """
    Upload resume from builder - generates PDF directly, saves to disk, 
    then processes through extraction pipeline
    """
    try:
        # ========== AUTHENTICATION CHECK ==========
        if 'user_id' not in session or 'user_name' not in session:
            return jsonify({'success': False, 'message': 'Not authenticated'}), 401

        data = request.json
        
        # ========== VALIDATE DATA ==========
        if not data.get('fullName') or not data.get('email') or not data.get('phone'):
            return jsonify({'success': False, 'message': 'Missing required fields'}), 400
        
        logger.info(f"Building PDF resume from form data for user {session['user_id']}")
        
        # ========== GENERATE PDF ==========
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Frame
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_LEFT, TA_JUSTIFY
        from io import BytesIO
        
        # Create filename
        timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
        filename = f"Resume_{data.get('fullName', 'User').replace(' ', '_')}_{timestamp}.pdf"
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        
        # Create PDF document
        pdf_buffer = BytesIO()
        doc = SimpleDocTemplate(
            filepath,
            pagesize=letter,
            rightMargin=0.75*inch,
            leftMargin=0.75*inch,
            topMargin=0.5*inch,
            bottomMargin=0.5*inch
        )
        
        # Container for PDF elements
        story = []
        styles = getSampleStyleSheet()
        
        # Custom styles
        style_name = ParagraphStyle(
            'CustomName',
            parent=styles['Heading1'],
            fontSize=28,
            textColor=colors.HexColor('#1f2937'),
            spaceAfter=6,
            leading=32
        )
        
        style_title = ParagraphStyle(
            'CustomTitle',
            parent=styles['Normal'],
            fontSize=14,
            textColor=colors.HexColor('#4b5563'),
            spaceAfter=6
        )
        
        style_section_header = ParagraphStyle(
            'SectionHeader',
            parent=styles['Heading2'],
            fontSize=13,
            textColor=colors.HexColor('#1f2937'),
            spaceAfter=8,
            spaceBefore=12,
            borderWidth=1,
            borderColor=colors.HexColor('#333333'),
            borderPadding=4,
            leftIndent=0
        )
        
        style_body = ParagraphStyle(
            'CustomBody',
            parent=styles['Normal'],
            fontSize=11,
            textColor=colors.HexColor('#4b5563'),
            leading=16
        )
        
        style_small = ParagraphStyle(
            'SmallText',
            parent=styles['Normal'],
            fontSize=9,
            textColor=colors.HexColor('#6b7280'),
            leading=12
        )
        
        # ========== BUILD PDF CONTENT ==========
        
        # Header
        story.append(Paragraph(data.get('fullName', 'Your Name'), style_name))
        if data.get('professionalTitle'):
            story.append(Paragraph(data['professionalTitle'], style_title))
        if data.get('location'):
            story.append(Paragraph(data['location'], style_small))
        story.append(Spacer(1, 0.2*inch))
        
        # Create two-column layout using Table
        left_column = []
        right_column = []
        
        # LEFT COLUMN - Contact
        if data.get('email') or data.get('phone') or data.get('linkedin') or data.get('portfolio'):
            left_column.append(Paragraph('<b>CONTACT</b>', style_section_header))
            if data.get('email'):
                left_column.append(Paragraph(f"<b>Email:</b><br/>{data['email']}", style_small))
            if data.get('phone'):
                left_column.append(Paragraph(f"<b>Phone:</b><br/>{data['phone']}", style_small))
            if data.get('linkedin'):
                left_column.append(Paragraph(f"<b>LinkedIn:</b><br/>{data['linkedin']}", style_small))
            if data.get('portfolio'):
                left_column.append(Paragraph(f"<b>Portfolio:</b><br/>{data['portfolio']}", style_small))
            left_column.append(Spacer(1, 0.1*inch))
        
        # LEFT COLUMN - Skills
        top_skills = [s for s in data.get('topSkills', []) if s and s.strip()]
        if top_skills:
            left_column.append(Paragraph('<b>TOP SKILLS</b>', style_section_header))
            for skill in top_skills:
                left_column.append(Paragraph(f"• {skill}", style_small))
            left_column.append(Spacer(1, 0.1*inch))
        
        # LEFT COLUMN - Languages
        languages = [l for l in data.get('languages', []) if l.get('language', '').strip()]
        if languages:
            left_column.append(Paragraph('<b>LANGUAGES</b>', style_section_header))
            for lang in languages:
                lang_text = lang['language']
                if lang.get('proficiency'):
                    lang_text += f" ({lang['proficiency']})"
                left_column.append(Paragraph(lang_text, style_small))
        
        # RIGHT COLUMN - Summary
        if data.get('summary') and data['summary'].strip():
            right_column.append(Paragraph('<b>PROFESSIONAL SUMMARY</b>', style_section_header))
            right_column.append(Paragraph(data['summary'], style_body))
            right_column.append(Spacer(1, 0.15*inch))
        
        # RIGHT COLUMN - Experience
        experiences = [e for e in data.get('experiences', []) if e.get('company', '').strip()]
        if experiences:
            right_column.append(Paragraph('<b>EXPERIENCE</b>', style_section_header))
            for exp in experiences:
                right_column.append(Paragraph(f"<b>{exp.get('company', 'Company')}</b>", style_body))
                if exp.get('position'):
                    right_column.append(Paragraph(f"<i>{exp['position']}</i>", style_body))
                
                date_str = ""
                if exp.get('startDate'):
                    date_str = exp['startDate']
                if exp.get('endDate'):
                    date_str += f" - {exp['endDate']}"
                if exp.get('duration'):
                    date_str += f" ({exp['duration']})"
                if date_str:
                    right_column.append(Paragraph(date_str, style_small))
                
                if exp.get('location'):
                    right_column.append(Paragraph(exp['location'], style_small))
                if exp.get('description'):
                    right_column.append(Paragraph(exp['description'], style_body))
                right_column.append(Spacer(1, 0.1*inch))
        
        # RIGHT COLUMN - Education
        education = [e for e in data.get('education', []) if e.get('institution', '').strip()]
        if education:
            right_column.append(Paragraph('<b>EDUCATION</b>', style_section_header))
            for edu in education:
                right_column.append(Paragraph(f"<b>{edu.get('institution', 'Institution')}</b>", style_body))
                
                degree_parts = []
                if edu.get('degree'):
                    degree_parts.append(edu['degree'])
                if edu.get('field'):
                    degree_parts.append(edu['field'])
                if degree_parts:
                    right_column.append(Paragraph(' - '.join(degree_parts), style_body))
                
                if edu.get('graduationDate'):
                    right_column.append(Paragraph(edu['graduationDate'], style_small))
                if edu.get('location'):
                    right_column.append(Paragraph(edu['location'], style_small))
                right_column.append(Spacer(1, 0.1*inch))
        
        # RIGHT COLUMN - Certifications
        certifications = [c for c in data.get('certifications', []) if c.get('name', '').strip()]
        if certifications:
            right_column.append(Paragraph('<b>CERTIFICATIONS</b>', style_section_header))
            for cert in certifications:
                right_column.append(Paragraph(f"<b>{cert.get('name', 'Certification')}</b>", style_body))
                if cert.get('issuer'):
                    right_column.append(Paragraph(cert['issuer'], style_body))
                if cert.get('date'):
                    right_column.append(Paragraph(cert['date'], style_small))
                right_column.append(Spacer(1, 0.1*inch))
        
        # Create two-column table
        table_data = [[left_column, right_column]]
        col_table = Table(table_data, colWidths=[2.2*inch, 4.8*inch])
        col_table.setStyle(TableStyle([
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('LEFTPADDING', (0, 0), (-1, -1), 0),
            ('RIGHTPADDING', (0, 0), (-1, -1), 0),
        ]))
        
        story.append(col_table)
        
        # Build PDF
        doc.build(story)
        
        logger.info(f"Generated PDF resume saved to: {filepath}")
        
        # ========== EXTRACT TEXT FROM PDF ==========
        try:
            import fitz  # PyMuPDF
            text_parts = []
            
            with fitz.open(filepath) as pdf_doc:
                for page in pdf_doc:
                    text_parts.append(page.get_text())
            
            text_of_resume = "\n".join(text_parts)
            
            logger.info(f"📄 Extracted text length: {len(text_of_resume)} characters")
            logger.info(f"📄 Text preview (first 300 chars):\n{text_of_resume[:300]}")
            
        except Exception as e:
            logger.error(f"Error extracting text from generated PDF: {e}")
            import traceback
            logger.error(traceback.format_exc())
            if os.path.exists(filepath):
                os.remove(filepath)
            return jsonify({'success': False, 'message': 'Error processing PDF'}), 500

        # Validate extraction
        if not text_of_resume or len(text_of_resume.strip()) < 50:
            logger.warning(f"Generated PDF text too short for user {session['user_id']}: {len(text_of_resume)} chars")
            if os.path.exists(filepath):
                os.remove(filepath)
            return jsonify({'success': False, 'message': 'Resume content too short'}), 400
        
        try:
            # Initialize variables (FIXED - same as uploadResume)
            value_name = None
            value_skills = []  # ✅ Changed from None to []
            value_certificate = []  # ✅ Changed from None to []
            value_education = []  # ✅ Changed from None to []
            value_workedAs = []  # ✅ Changed from None to []
            value_experience = []
            value_linkedin = None
            calculated_experience = 0.0
            extraction_method = "Resume Builder (DOCX Generated)"
            
            # === USE NEW MODEL for SKILLS ONLY ===
            if nlp_new:
                logger.info("Using NEW model for skills only...")
                doc_new = nlp_new(text_of_resume)
                
                seen_skills = set()
                seen_jobs = set()
                
                for ent in doc_new.ents:
                    text = ent.text.strip()
                    if len(text) < 2:
                        continue
                    
                    if ent.label_ in ['TECHNICAL_SKILL', 'SOFT_SKILL', 'SKILLS']:
                        if text.lower() not in seen_skills:
                            seen_skills.add(text.lower())
                            value_skills.append(text)

                    elif ent.label_ in ['JOB_TITLE', 'WORKED AS']:
                        if text.lower() not in seen_jobs:
                            seen_jobs.add(text.lower())
                            value_workedAs.append(text)
            
            # === USE OLD MODEL for Everything Else ===
            if nlp_old:
                logger.info("Using OLD model for experience, education, certs...")
                doc_old = nlp_old(text_of_resume)
                
                seen_skills_old = set()
                seen_certs = set()
                seen_education = set()
                seen_jobs = set()
                
                for ent in doc_old.ents:
                    text = ent.text.strip()
                    if len(text) < 2:
                        continue

                    if ent.label_ in ['TECHNICAL_SKILL', 'SOFT_SKILL', 'SKILLS']:
                        if text.lower() not in seen_skills_old:
                            seen_skills_old.add(text.lower())
                            value_skills.append(text) 
                    
                    elif ent.label_ == 'CERTIFICATION':
                        if text.lower() not in seen_certs:
                            seen_certs.add(text.lower())
                            value_certificate.append(text)
                    
                    elif ent.label_ in ['EDUCATION', 'DEGREE']:
                        if text.lower() not in seen_education:
                            seen_education.add(text.lower())
                            value_education.append(text)
                    
                    elif ent.label_ == 'WORKED AS':
                        if text.lower() not in seen_jobs:
                            seen_jobs.add(text.lower())
                            value_workedAs.append(text)
                    
                    elif ent.label_ == 'YEARS OF EXPERIENCE':
                        value_experience.append(text)
            
            # === FALLBACK for experience ===
            if not value_experience and nlp_new:
                logger.info("Old model didn't find experience, trying new model...")
                for ent in doc_new.ents:
                    if ent.label_ in ['YEARS_OF_EXP', 'YEARS OF EXPERIENCE']:
                        value_experience = [ent.text.strip()]
                        break
            
            # === CALCULATE EXPERIENCE YEARS ===
            if value_experience:
                for exp_text in value_experience:
                    exp_str = str(exp_text).lower()
                    
                    years_match = re.search(r'(\d+(?:\.\d+)?)\s*(?:year|yr|y)', exp_str)
                    if years_match:
                        calculated_experience += float(years_match.group(1))
                    
                    months_match = re.search(r'(\d+(?:\.\d+)?)\s*(?:month|mon|m)', exp_str)
                    if months_match:
                        months = float(months_match.group(1))
                        calculated_experience += months / 12.0
                    
                    if not years_match and not months_match:
                        number_match = re.search(r'(\d+(?:\.\d+)?)', exp_str)
                        if number_match:
                            calculated_experience += float(number_match.group(1))
                
                calculated_experience = round(calculated_experience, 2)
            
            # Convert to None if empty
            value_skills = value_skills if value_skills else None
            value_certificate = value_certificate if value_certificate else None
            value_education = value_education if value_education else None
            value_workedAs = value_workedAs if value_workedAs else None
            
            # ========== CLEAN THE DATA (ADDED - Same as uploadResume) ==========
            logger.info("Cleaning extracted data...")
            cleaned_data = clean_resume_data({
                'Name': value_name,
                'LINKEDIN LINK': value_linkedin,
                'SKILLS': value_skills,
                'CERTIFICATION': value_certificate,
                'EDUCATION': value_education,
                'WORKED AS': value_workedAs,
                'YEARS OF EXPERIENCE': value_experience,
                'CALCULATED_EXPERIENCE_YEARS': calculated_experience
            })

            # Use cleaned data
            value_name = cleaned_data['Name']
            value_linkedin = cleaned_data['LINKEDIN LINK']
            value_skills = cleaned_data['SKILLS']
            value_certificate = cleaned_data['CERTIFICATION']
            value_education = cleaned_data['EDUCATION']
            value_workedAs = cleaned_data['WORKED AS']
            value_experience = cleaned_data['YEARS OF EXPERIENCE']
            calculated_experience = cleaned_data['CALCULATED_EXPERIENCE_YEARS']
            
            logger.info(f"✅ NER extraction completed - Skills: {len(value_skills or [])}, Certs: {len(value_certificate or [])}, Edu: {len(value_education or [])}")
            
        except Exception as e:
            logger.error(f"Error in NER extraction: {e}")
            import traceback
            logger.error(traceback.format_exc())
            extraction_method = "Resume Builder (DOCX Generated - Extraction Error)"

        try:
            with open(filepath, "rb") as f:
                file_data = f.read()
            
            resume_document = {
                "UserId": ObjectId(session['user_id']),
                "Name": data.get('fullName'),
                "LINKEDIN LINK": data.get('linkedin'),
                "SKILLS": value_skills,
                "CERTIFICATION": value_certificate,
                "EDUCATION": value_education,
                "WORKED AS": value_workedAs,
                "YEARS OF EXPERIENCE": value_experience if value_experience else None,
                "CALCULATED_EXPERIENCE_YEARS": calculated_experience,
                "Appear": 0,
                "ResumeTitle": filename.replace('.pdf', ''),
                "InternalFilename": filename,
                "ResumeData": text_of_resume,
                "FileData": file_data,
                "UploadedAt": datetime.now(),
                "ExtractionMethod": "Resume Builder (PDF Generated)"
            }
            
            result = resumeFetchedData.insert_one(resume_document)
            
            if result:
                skill_count = len(value_skills) if value_skills else 0
                cert_count = len(value_certificate) if value_certificate else 0
                edu_count = len(value_education) if value_education else 0
                
                return jsonify({
                    'success': True,
                    'message': f'PDF Resume created! Found {skill_count} skills, {cert_count} certifications, {edu_count} education entries.',
                    'resume_id': str(result.inserted_id)
                }), 200
        
        except Exception as e:
            logger.error(f"Database error: {e}")
            if os.path.exists(filepath):
                os.remove(filepath)
            return jsonify({'success': False, 'message': 'Error storing resume'}), 500
    
    except Exception as e:
        logger.error(f"Unexpected error: {e}")
        import traceback
        logger.error(traceback.format_exc())
        return jsonify({'success': False, 'message': 'An unexpected error occurred'}), 500


        
@app.route('/upload_generated_resume', methods=['POST'])
def upload_generated_resume():
    """
    Upload resume from builder - generates DOCX, saves to disk, 
    then processes through SAME extraction pipeline as manual uploads
    """
    try:
        # ========== AUTHENTICATION CHECK ==========
        if 'user_id' not in session or 'user_name' not in session:
            return jsonify({'success': False, 'message': 'Not authenticated'}), 401

        data = request.json
        
        # ========== VALIDATE DATA ==========
        if not data.get('fullName') or not data.get('email') or not data.get('phone'):
            return jsonify({'success': False, 'message': 'Missing required fields'}), 400
        
        logger.info(f"Building resume from form data for user {session['user_id']}")
        
        # ========== GENERATE DOCX (Same as before) ==========
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        
        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)

        # ========== HEADER SECTION ==========
        name_para = doc.add_paragraph()
        name_run = name_para.add_run(data.get('fullName', 'Your Name'))
        name_run.font.size = Pt(24)
        name_run.font.bold = True
        name_run.font.color.rgb = RGBColor(0, 0, 0)
        name_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Professional Title
        if data.get('professionalTitle'):
            title_para = doc.add_paragraph()
            title_run = title_para.add_run(data['professionalTitle'])
            title_run.font.size = Pt(12)
            title_run.font.color.rgb = RGBColor(100, 100, 100)
            title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            title_para.space_after = Pt(6)
        
        # Location
        if data.get('location'):
            location_para = doc.add_paragraph()
            location_run = location_para.add_run(data['location'])
            location_run.font.size = Pt(10)
            location_para.space_after = Pt(12)
        
        # ========== CREATE TWO COLUMN LAYOUT WITH TABLE ==========
        table = doc.add_table(rows=1, cols=2)
        table.allow_autofit = False
        table.columns[0].width = Inches(2.0)
        table.columns[1].width = Inches(4.5)
        
        left_cell = table.rows[0].cells[0]
        right_cell = table.rows[0].cells[1]
        
        # ========== LEFT COLUMN (SIDEBAR) ==========
        if data.get('email') or data.get('phone') or data.get('linkedin') or data.get('portfolio'):
            add_section_header(left_cell, "Contact")
            
            if data.get('email'):
                add_small_text(left_cell, data['email'])
            
            if data.get('phone'):
                add_small_text(left_cell, data['phone'])
            
            if data.get('linkedin'):
                add_small_text(left_cell, data['linkedin'])
                add_small_text(left_cell, "(LinkedIn)", italic=True)
            
            if data.get('portfolio'):
                add_small_text(left_cell, data['portfolio'])
                add_small_text(left_cell, "(Portfolio)", italic=True)
            
            left_cell.add_paragraph()
        
        # Top Skills Section
        top_skills = [s for s in data.get('topSkills', []) if s and s.strip()]
        if top_skills:
            add_section_header(left_cell, "Top Skills")
            for skill in top_skills:
                add_small_text(left_cell, skill)
            left_cell.add_paragraph()
        
        # Languages Section
        languages = [l for l in data.get('languages', []) if l.get('language', '').strip()]
        if languages:
            add_section_header(left_cell, "Languages")
            for lang in languages:
                add_small_text(left_cell, lang['language'])
                if lang.get('proficiency'):
                    add_small_text(left_cell, f"({lang['proficiency']})", italic=True)
            left_cell.add_paragraph()
        
        # ========== RIGHT COLUMN (MAIN CONTENT) ==========
        
        # Professional Summary Section
        if data.get('summary') and data['summary'].strip():
            add_main_section_header(right_cell, "Professional Summary")
            summary_para = right_cell.add_paragraph(data['summary'])
            summary_para_format = summary_para.paragraph_format
            summary_para_format.space_after = Pt(12)
            for run in summary_para.runs:
                run.font.size = Pt(10)
        
        # Experience Section
        experiences = [e for e in data.get('experiences', []) if e.get('company', '').strip()]
        if experiences:
            add_main_section_header(right_cell, "Experience")
            
            for exp in experiences:
                company_para = right_cell.add_paragraph()
                company_run = company_para.add_run(exp.get('company', 'Company Name'))
                company_run.font.size = Pt(11)
                company_run.font.bold = True
                company_para.space_after = Pt(2)
                
                if exp.get('position'):
                    position_para = right_cell.add_paragraph()
                    position_run = position_para.add_run(exp['position'])
                    position_run.font.size = Pt(10)
                    position_run.font.italic = True
                    position_para.space_after = Pt(2)
                
                date_parts = []
                if exp.get('startDate'):
                    date_parts.append(exp['startDate'])
                if exp.get('endDate'):
                    date_parts.append(exp['endDate'])
                
                date_str = ' - '.join(date_parts)
                if exp.get('duration'):
                    date_str += f" ({exp['duration']})"
                
                if date_str:
                    date_para = right_cell.add_paragraph()
                    date_run = date_para.add_run(date_str)
                    date_run.font.size = Pt(9)
                    date_run.font.color.rgb = RGBColor(100, 100, 100)
                    date_para.space_after = Pt(2)
                
                if exp.get('location'):
                    loc_para = right_cell.add_paragraph()
                    loc_run = loc_para.add_run(exp['location'])
                    loc_run.font.size = Pt(9)
                    loc_para.space_after = Pt(2)
                
                if exp.get('description'):
                    desc_para = right_cell.add_paragraph()
                    desc_run = desc_para.add_run(exp['description'])
                    desc_run.font.size = Pt(9)
                    desc_para.space_after = Pt(8)
                else:
                    right_cell.add_paragraph().space_after = Pt(8)
        
        # Education Section
        education = [e for e in data.get('education', []) if e.get('institution', '').strip()]
        if education:
            add_main_section_header(right_cell, "Education")
            
            for edu in education:
                inst_para = right_cell.add_paragraph()
                inst_run = inst_para.add_run(edu.get('institution', 'Institution Name'))
                inst_run.font.size = Pt(11)
                inst_run.font.bold = True
                inst_para.space_after = Pt(2)
                
                degree_parts = []
                if edu.get('degree'):
                    degree_parts.append(edu['degree'])
                if edu.get('field'):
                    degree_parts.append(edu['field'])
                
                if degree_parts:
                    degree_para = right_cell.add_paragraph()
                    degree_run = degree_para.add_run(' - '.join(degree_parts))
                    degree_run.font.size = Pt(10)
                    degree_para.space_after = Pt(2)
                
                if edu.get('graduationDate'):
                    grad_para = right_cell.add_paragraph()
                    grad_run = grad_para.add_run(edu['graduationDate'])
                    grad_run.font.size = Pt(9)
                    grad_run.font.color.rgb = RGBColor(100, 100, 100)
                    grad_para.space_after = Pt(2)
                
                if edu.get('location'):
                    loc_para = right_cell.add_paragraph()
                    loc_run = loc_para.add_run(edu['location'])
                    loc_run.font.size = Pt(9)
                    loc_para.space_after = Pt(8)
                else:
                    right_cell.add_paragraph().space_after = Pt(8)
        
        # Certifications Section
        certifications = [c for c in data.get('certifications', []) if c.get('name', '').strip()]
        if certifications:
            add_main_section_header(right_cell, "Certifications")
            
            for cert in certifications:
                cert_para = right_cell.add_paragraph()
                cert_run = cert_para.add_run(cert.get('name', 'Certification Name'))
                cert_run.font.size = Pt(11)
                cert_run.font.bold = True
                cert_para.space_after = Pt(2)
                
                if cert.get('issuer'):
                    issuer_para = right_cell.add_paragraph()
                    issuer_run = issuer_para.add_run(cert['issuer'])
                    issuer_run.font.size = Pt(10)
                    issuer_para.space_after = Pt(2)
                
                if cert.get('date'):
                    date_para = right_cell.add_paragraph()
                    date_run = date_para.add_run(cert['date'])
                    date_run.font.size = Pt(9)
                    date_run.font.color.rgb = RGBColor(100, 100, 100)
                    date_para.space_after = Pt(8)
                else:
                    right_cell.add_paragraph().space_after = Pt(8)
        
        # Additional Sections
        additional_sections = [s for s in data.get('additionalSections', []) if s.get('title', '').strip()]
        if additional_sections:
            for section in additional_sections:
                add_main_section_header(right_cell, section.get('title', 'Section'))
                content_para = right_cell.add_paragraph(section.get('content', ''))
                for run in content_para.runs:
                    run.font.size = Pt(10)
                content_para.space_after = Pt(12)
        
        # ========== SAVE DOCX TO DISK ==========
        timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
        filename = f"Resume_{data.get('fullName', 'User').replace(' ', '_')}_{timestamp}.docx"
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        
        doc.save(filepath)
        logger.info(f"Generated resume DOCX saved to: {filepath}")
        
        # ========== EXTRACT TEXT FROM GENERATED DOCX ==========
        # ========== EXTRACT TEXT FROM GENERATED DOCX ==========
        from docx import Document as DocxDocument
        try:
            docx_doc = DocxDocument(filepath)
            
            # Extract text from paragraphs AND tables
            text_parts = []
            
            # First, get all paragraphs outside tables
            for paragraph in docx_doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text.strip())
            
            # CRITICAL: Extract text from tables (where the two-column layout content is!)
            for table in docx_doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if paragraph.text.strip():
                                text_parts.append(paragraph.text.strip())
            
            text_of_resume = "\n".join(text_parts)
            
            # Log for debugging
            logger.info(f"📄 Extracted text length: {len(text_of_resume)} characters")
            logger.info(f"📄 Text preview (first 300 chars):\n{text_of_resume[:300]}")
            
        except Exception as e:
            logger.error(f"Error extracting text from generated DOCX: {e}")
            import traceback
            logger.error(traceback.format_exc())
            if os.path.exists(filepath):
                os.remove(filepath)
            return jsonify({'success': False, 'message': 'Error processing DOCX'}), 500

        # Validate extraction
        if not text_of_resume or len(text_of_resume.strip()) < 50:
            logger.warning(f"Generated resume text too short for user {session['user_id']}: {len(text_of_resume)} chars")
            logger.warning(f"Extracted text was: {text_of_resume}")
            if os.path.exists(filepath):
                os.remove(filepath)
            return jsonify({'success': False, 'message': 'Resume content too short'}), 400
        
        # ========== NER MODEL EXTRACTION (IDENTICAL TO uploadResume) ==========
        try:
            # Initialize variables (FIXED - same as uploadResume)
            value_name = None
            value_skills = []  # ✅ Changed from None to []
            value_certificate = []  # ✅ Changed from None to []
            value_education = []  # ✅ Changed from None to []
            value_workedAs = []  # ✅ Changed from None to []
            value_experience = []
            value_linkedin = None
            calculated_experience = 0.0
            extraction_method = "Resume Builder (DOCX Generated)"
            
            # === USE NEW MODEL for SKILLS ONLY ===
            if nlp_new:
                logger.info("Using NEW model for skills only...")
                doc_new = nlp_new(text_of_resume)
                
                seen_skills = set()
                seen_jobs = set()
                
                for ent in doc_new.ents:
                    text = ent.text.strip()
                    if len(text) < 2:
                        continue
                    
                    if ent.label_ in ['TECHNICAL_SKILL', 'SOFT_SKILL', 'SKILLS']:
                        if text.lower() not in seen_skills:
                            seen_skills.add(text.lower())
                            value_skills.append(text)

                    elif ent.label_ in ['JOB_TITLE', 'WORKED AS']:
                        if text.lower() not in seen_jobs:
                            seen_jobs.add(text.lower())
                            value_workedAs.append(text)
            
            # === USE OLD MODEL for Everything Else ===
            if nlp_old:
                logger.info("Using OLD model for experience, education, certs...")
                doc_old = nlp_old(text_of_resume)
                
                seen_skills_old = set()
                seen_certs = set()
                seen_education = set()
                seen_jobs = set()
                
                for ent in doc_old.ents:
                    text = ent.text.strip()
                    if len(text) < 2:
                        continue

                    if ent.label_ in ['TECHNICAL_SKILL', 'SOFT_SKILL', 'SKILLS']:
                        if text.lower() not in seen_skills_old:
                            seen_skills_old.add(text.lower())
                            value_skills.append(text) 
                    
                    elif ent.label_ == 'CERTIFICATION':
                        if text.lower() not in seen_certs:
                            seen_certs.add(text.lower())
                            value_certificate.append(text)
                    
                    elif ent.label_ in ['EDUCATION', 'DEGREE']:
                        if text.lower() not in seen_education:
                            seen_education.add(text.lower())
                            value_education.append(text)
                    
                    elif ent.label_ == 'WORKED AS':
                        if text.lower() not in seen_jobs:
                            seen_jobs.add(text.lower())
                            value_workedAs.append(text)
                    
                    elif ent.label_ == 'YEARS OF EXPERIENCE':
                        value_experience.append(text)
            
            # === FALLBACK for experience ===
            if not value_experience and nlp_new:
                logger.info("Old model didn't find experience, trying new model...")
                for ent in doc_new.ents:
                    if ent.label_ in ['YEARS_OF_EXP', 'YEARS OF EXPERIENCE']:
                        value_experience = [ent.text.strip()]
                        break
            
            # === CALCULATE EXPERIENCE YEARS ===
            if value_experience:
                for exp_text in value_experience:
                    exp_str = str(exp_text).lower()
                    
                    years_match = re.search(r'(\d+(?:\.\d+)?)\s*(?:year|yr|y)', exp_str)
                    if years_match:
                        calculated_experience += float(years_match.group(1))
                    
                    months_match = re.search(r'(\d+(?:\.\d+)?)\s*(?:month|mon|m)', exp_str)
                    if months_match:
                        months = float(months_match.group(1))
                        calculated_experience += months / 12.0
                    
                    if not years_match and not months_match:
                        number_match = re.search(r'(\d+(?:\.\d+)?)', exp_str)
                        if number_match:
                            calculated_experience += float(number_match.group(1))
                
                calculated_experience = round(calculated_experience, 2)
            
            # Convert to None if empty
            value_skills = value_skills if value_skills else None
            value_certificate = value_certificate if value_certificate else None
            value_education = value_education if value_education else None
            value_workedAs = value_workedAs if value_workedAs else None
            
            # ========== CLEAN THE DATA (ADDED - Same as uploadResume) ==========
            logger.info("Cleaning extracted data...")
            cleaned_data = clean_resume_data({
                'Name': value_name,
                'LINKEDIN LINK': value_linkedin,
                'SKILLS': value_skills,
                'CERTIFICATION': value_certificate,
                'EDUCATION': value_education,
                'WORKED AS': value_workedAs,
                'YEARS OF EXPERIENCE': value_experience,
                'CALCULATED_EXPERIENCE_YEARS': calculated_experience
            })

            # Use cleaned data
            value_name = cleaned_data['Name']
            value_linkedin = cleaned_data['LINKEDIN LINK']
            value_skills = cleaned_data['SKILLS']
            value_certificate = cleaned_data['CERTIFICATION']
            value_education = cleaned_data['EDUCATION']
            value_workedAs = cleaned_data['WORKED AS']
            value_experience = cleaned_data['YEARS OF EXPERIENCE']
            calculated_experience = cleaned_data['CALCULATED_EXPERIENCE_YEARS']
            
            logger.info(f"✅ NER extraction completed - Skills: {len(value_skills or [])}, Certs: {len(value_certificate or [])}, Edu: {len(value_education or [])}")
            
        except Exception as e:
            logger.error(f"Error in NER extraction: {e}")
            import traceback
            logger.error(traceback.format_exc())
            extraction_method = "Resume Builder (DOCX Generated - Extraction Error)"
        
        # ========== STORE IN DATABASE ==========
        try:
            with open(filepath, "rb") as f:
                file_data = f.read()
            
            resume_document = {
                "UserId": ObjectId(session['user_id']),
                "Name": data.get('fullName'),  # Use form data as backup
                "LINKEDIN LINK": data.get('linkedin'),
                "SKILLS": value_skills,
                "CERTIFICATION": value_certificate,
                "EDUCATION": value_education,
                "WORKED AS": value_workedAs,
                "YEARS OF EXPERIENCE": value_experience if value_experience else None,
                "CALCULATED_EXPERIENCE_YEARS": calculated_experience,
                "Appear": 0,
                "ResumeTitle": filename.replace('.docx', ''),
                "InternalFilename": filename,
                "ResumeData": text_of_resume,
                "FileData": file_data,
                "UploadedAt": datetime.now(),
                "ExtractionMethod": extraction_method
            }
            
            result = resumeFetchedData.insert_one(resume_document)
            logger.info(f"✅ Resume from builder stored in database for user {session['user_id']}")
            
            if result:
                skill_count = len(value_skills) if value_skills else 0
                cert_count = len(value_certificate) if value_certificate else 0
                edu_count = len(value_education) if value_education else 0
                
                return jsonify({
                    'success': True,
                    'message': f'Resume created and uploaded! Found {skill_count} skills, {cert_count} certifications, {edu_count} education entries.',
                    'resume_id': str(result.inserted_id),
                    'extraction_stats': {
                        'skills': skill_count,
                        'certifications': cert_count,
                        'education': edu_count,
                        'experience_years': calculated_experience
                    }
                }), 200
            else:
                os.remove(filepath)
                return jsonify({'success': False, 'message': 'Failed to store resume'}), 500
        
        except Exception as e:
            logger.error(f"Database error: {e}")
            import traceback
            logger.error(traceback.format_exc())
            if os.path.exists(filepath):
                os.remove(filepath)
            return jsonify({'success': False, 'message': 'Error storing resume in database'}), 500
    
    except Exception as e:
        logger.error(f"Unexpected error in upload_generated_resume: {e}")
        import traceback
        logger.error(traceback.format_exc())
        return jsonify({'success': False, 'message': 'An unexpected error occurred'}), 500

@app.route('/resume_builder')
def resume_builder():
    """Resume builder page"""
    if 'user_id' not in session or 'user_name' not in session:
        flash("Please log in first", "error")
        return redirect(url_for('loginpage'))
    
    return render_template('resume_builder.html')

@app.route('/generate_resume_docx', methods=['POST'])
def generate_resume_docx():
    """Generate DOCX resume from form data with improved design"""
    try:
        if 'user_id' not in session:
            return jsonify({'success': False, 'message': 'Not authenticated'}), 401

        data = request.json
        
        # Create document
        doc = Document()
        
        # Set document margins - slightly wider for better readability
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.6)
            section.bottom_margin = Inches(0.6)
            section.left_margin = Inches(0.7)
            section.right_margin = Inches(0.7)

        # ========== HEADER SECTION WITH MODERN STYLING ==========
        # Name - larger and bolder
        name_para = doc.add_paragraph()
        name_run = name_para.add_run(data.get('fullName', 'Your Name'))
        name_run.font.size = Pt(28)
        name_run.font.bold = True
        name_run.font.color.rgb = RGBColor(31, 41, 55)  # Dark gray
        name_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        name_para.space_after = Pt(4)
        
        # Professional Title with accent color
        if data.get('professionalTitle'):
            title_para = doc.add_paragraph()
            title_run = title_para.add_run(data['professionalTitle'])
            title_run.font.size = Pt(13)
            title_run.font.color.rgb = RGBColor(59, 130, 246)  # Blue accent
            title_run.font.bold = True
            title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            title_para.space_after = Pt(6)
        
        # Location with icon-like formatting
        if data.get('location'):
            location_para = doc.add_paragraph()
            location_run = location_para.add_run(f"📍 {data['location']}")
            location_run.font.size = Pt(10)
            location_run.font.color.rgb = RGBColor(107, 114, 128)
            location_para.space_after = Pt(8)
        
        # Horizontal line separator
        separator_para = doc.add_paragraph()
        separator_run = separator_para.add_run('─' * 85)
        separator_run.font.color.rgb = RGBColor(59, 130, 246)
        separator_para.space_after = Pt(12)
        
        # ========== CREATE TWO COLUMN LAYOUT ==========
        table = doc.add_table(rows=1, cols=2)
        table.allow_autofit = False
        table.columns[0].width = Inches(2.2)  # Slightly wider sidebar
        table.columns[1].width = Inches(4.3)  # Adjusted main content
        
        left_cell = table.rows[0].cells[0]
        right_cell = table.rows[0].cells[1]
        
        # Remove cell borders for cleaner look
        for cell in table.rows[0].cells:
            cell._element.get_or_add_tcPr().append(parse_xml(r'<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:top w:val="none"/><w:left w:val="none"/><w:bottom w:val="none"/><w:right w:val="none"/></w:tcBorders>'))
        
        # ========== LEFT COLUMN (SIDEBAR WITH BACKGROUND COLOR EFFECT) ==========
        # Contact Section with improved formatting
        if data.get('email') or data.get('phone') or data.get('linkedin') or data.get('portfolio'):
            add_section_header_improved(left_cell, "CONTACT")
            
            if data.get('email'):
                add_contact_item(left_cell, "✉", data['email'])
            
            if data.get('phone'):
                add_contact_item(left_cell, "📞", data['phone'])
            
            if data.get('linkedin'):
                add_contact_item(left_cell, "🔗", data['linkedin'], label="LinkedIn")
            
            if data.get('portfolio'):
                add_contact_item(left_cell, "🌐", data['portfolio'], label="Portfolio")
            
            left_cell.add_paragraph().space_after = Pt(8)
        
        # Top Skills Section with bullet styling
        top_skills = [s for s in data.get('topSkills', []) if s.strip()]
        if top_skills:
            add_section_header_improved(left_cell, "SKILLS")
            for skill in top_skills:
                add_skill_item(left_cell, skill)
            left_cell.add_paragraph().space_after = Pt(8)
        
        # Languages Section with proficiency indicators
        languages = [l for l in data.get('languages', []) if l.get('language', '').strip()]
        if languages:
            add_section_header_improved(left_cell, "LANGUAGES")
            for lang in languages:
                add_language_item(left_cell, lang['language'], lang.get('proficiency', ''))
            left_cell.add_paragraph().space_after = Pt(8)
        
        # ========== RIGHT COLUMN (MAIN CONTENT) ==========
        
        # Professional Summary with better formatting
        if data.get('summary') and data['summary'].strip():
            add_main_section_header_improved(right_cell, "PROFESSIONAL SUMMARY")
            summary_para = right_cell.add_paragraph()
            summary_run = summary_para.add_run(data['summary'])
            summary_run.font.size = Pt(10)
            summary_run.font.color.rgb = RGBColor(55, 65, 81)
            summary_para.paragraph_format.line_spacing = 1.3
            summary_para.space_after = Pt(14)
        
        # Experience Section with timeline design
        experiences = [e for e in data.get('experiences', []) if e.get('company', '').strip()]
        if experiences:
            add_main_section_header_improved(right_cell, "WORK EXPERIENCE")
            
            for idx, exp in enumerate(experiences):
                # Company name with emphasis
                company_para = right_cell.add_paragraph()
                company_run = company_para.add_run(exp.get('company', 'Company Name'))
                company_run.font.size = Pt(12)
                company_run.font.bold = True
                company_run.font.color.rgb = RGBColor(31, 41, 55)
                company_para.space_after = Pt(3)
                
                # Position with styling
                if exp.get('position'):
                    position_para = right_cell.add_paragraph()
                    position_run = position_para.add_run(exp['position'])
                    position_run.font.size = Pt(11)
                    position_run.font.italic = True
                    position_run.font.color.rgb = RGBColor(59, 130, 246)
                    position_para.space_after = Pt(3)
                
                # Date and duration in one line
                date_parts = []
                if exp.get('startDate'):
                    date_parts.append(exp['startDate'])
                if exp.get('endDate'):
                    date_parts.append(exp['endDate'])
                
                date_str = ' – '.join(date_parts)
                if exp.get('duration'):
                    date_str += f" • {exp['duration']}"
                
                if date_str:
                    date_para = right_cell.add_paragraph()
                    date_run = date_para.add_run(f"📅 {date_str}")
                    date_run.font.size = Pt(9)
                    date_run.font.color.rgb = RGBColor(107, 114, 128)
                    date_para.space_after = Pt(2)
                
                # Location
                if exp.get('location'):
                    loc_para = right_cell.add_paragraph()
                    loc_run = loc_para.add_run(f"📍 {exp['location']}")
                    loc_run.font.size = Pt(9)
                    loc_run.font.color.rgb = RGBColor(107, 114, 128)
                    loc_para.space_after = Pt(6)
                
                # Description with better spacing
                if exp.get('description'):
                    desc_lines = exp['description'].split('\n')
                    for line in desc_lines:
                        line = line.strip()
                        if line:
                            if line.startswith('-') or line.startswith('•'):
                                line = line[1:].strip()
                            desc_para = right_cell.add_paragraph()
                            desc_para.add_run('• ').font.color.rgb = RGBColor(59, 130, 246)
                            desc_run = desc_para.add_run(line)
                            desc_run.font.size = Pt(10)
                            desc_run.font.color.rgb = RGBColor(75, 85, 99)
                            desc_para.paragraph_format.left_indent = Inches(0.15)
                            desc_para.space_after = Pt(4)
                    
                    right_cell.add_paragraph().space_after = Pt(10)
                else:
                    right_cell.add_paragraph().space_after = Pt(10)
        
        # Education Section with improved layout
        education = [e for e in data.get('education', []) if e.get('institution', '').strip()]
        if education:
            add_main_section_header_improved(right_cell, "EDUCATION")
            
            for edu in education:
                # Institution name
                inst_para = right_cell.add_paragraph()
                inst_run = inst_para.add_run(edu.get('institution', 'Institution Name'))
                inst_run.font.size = Pt(12)
                inst_run.font.bold = True
                inst_run.font.color.rgb = RGBColor(31, 41, 55)
                inst_para.space_after = Pt(3)
                
                # Degree and field
                degree_parts = []
                if edu.get('degree'):
                    degree_parts.append(edu['degree'])
                if edu.get('field'):
                    degree_parts.append(edu['field'])
                
                if degree_parts:
                    degree_para = right_cell.add_paragraph()
                    degree_run = degree_para.add_run(' – '.join(degree_parts))
                    degree_run.font.size = Pt(10)
                    degree_run.font.color.rgb = RGBColor(75, 85, 99)
                    degree_para.space_after = Pt(3)
                
                # Graduation date with icon
                if edu.get('graduationDate'):
                    grad_para = right_cell.add_paragraph()
                    grad_run = grad_para.add_run(f"🎓 {edu['graduationDate']}")
                    grad_run.font.size = Pt(9)
                    grad_run.font.color.rgb = RGBColor(107, 114, 128)
                    grad_para.space_after = Pt(2)
                
                # Location
                if edu.get('location'):
                    loc_para = right_cell.add_paragraph()
                    loc_run = loc_para.add_run(f"📍 {edu['location']}")
                    loc_run.font.size = Pt(9)
                    loc_run.font.color.rgb = RGBColor(107, 114, 128)
                    loc_para.space_after = Pt(10)
                else:
                    right_cell.add_paragraph().space_after = Pt(10)
        
        # Certifications Section
        certifications = [c for c in data.get('certifications', []) if c.get('name', '').strip()]
        if certifications:
            add_main_section_header_improved(right_cell, "CERTIFICATIONS")
            
            for cert in certifications:
                # Certification name
                cert_para = right_cell.add_paragraph()
                cert_run = cert_para.add_run(f"🏆 {cert.get('name', 'Certification Name')}")
                cert_run.font.size = Pt(11)
                cert_run.font.bold = True
                cert_run.font.color.rgb = RGBColor(31, 41, 55)
                cert_para.space_after = Pt(3)
                
                # Issuer
                if cert.get('issuer'):
                    issuer_para = right_cell.add_paragraph()
                    issuer_run = issuer_para.add_run(cert['issuer'])
                    issuer_run.font.size = Pt(10)
                    issuer_run.font.color.rgb = RGBColor(75, 85, 99)
                    issuer_para.paragraph_format.left_indent = Inches(0.15)
                    issuer_para.space_after = Pt(2)
                
                # Date
                if cert.get('date'):
                    date_para = right_cell.add_paragraph()
                    date_run = date_para.add_run(f"Issued: {cert['date']}")
                    date_run.font.size = Pt(9)
                    date_run.font.color.rgb = RGBColor(107, 114, 128)
                    date_para.paragraph_format.left_indent = Inches(0.15)
                    date_para.space_after = Pt(10)
                else:
                    right_cell.add_paragraph().space_after = Pt(10)
        
        # Additional Sections
        additional_sections = [s for s in data.get('additionalSections', []) if s.get('title', '').strip()]
        if additional_sections:
            for section in additional_sections:
                add_main_section_header_improved(right_cell, section.get('title', 'SECTION').upper())
                
                content = section.get('content', '').strip()
                if content:
                    lines = content.split('\n')
                    for line in lines:
                        trimmed = line.strip()
                        if not trimmed:
                            continue
                        
                        if trimmed.startswith('-') or trimmed.startswith('*') or trimmed.startswith('•'):
                            text = trimmed[1:].strip()
                            bullet_para = right_cell.add_paragraph()
                            bullet_para.add_run('• ').font.color.rgb = RGBColor(59, 130, 246)
                            bullet_run = bullet_para.add_run(text)
                            bullet_run.font.size = Pt(10)
                            bullet_run.font.color.rgb = RGBColor(75, 85, 99)
                            bullet_para.paragraph_format.left_indent = Inches(0.15)
                            bullet_para.space_after = Pt(4)
                        else:
                            content_para = right_cell.add_paragraph(trimmed)
                            for run in content_para.runs:
                                run.font.size = Pt(10)
                                run.font.color.rgb = RGBColor(75, 85, 99)
                            content_para.space_after = Pt(6)
                    
                    right_cell.add_paragraph().space_after = Pt(10)

        # Save document to BytesIO
        doc_io = BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        
        # Generate filename
        filename = f"Resume_{data.get('fullName', 'User').replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.docx"
        
        return send_file(
            doc_io,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=filename
        )
    
    except Exception as e:
        logger.error(f"Error generating resume DOCX: {e}")
        import traceback
        logger.error(traceback.format_exc())
        return jsonify({'success': False, 'message': str(e)}), 500

def add_section_header_improved(cell, text):
    """Add a modern sidebar section header with accent color"""
    para = cell.add_paragraph()
    run = para.add_run(text)
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.color.rgb = RGBColor(59, 130, 246)  # Blue accent
    para.space_after = Pt(6)
    
    # Add subtle underline
    underline_para = cell.add_paragraph()
    underline_run = underline_para.add_run('━' * 20)
    underline_run.font.size = Pt(8)
    underline_run.font.color.rgb = RGBColor(59, 130, 246)
    underline_para.space_after = Pt(8)


def add_main_section_header_improved(cell, text):
    """Add a main section header with modern styling"""
    para = cell.add_paragraph()
    run = para.add_run(text)
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(31, 41, 55)
    para.space_after = Pt(4)
    
    # Add colored line below
    line_para = cell.add_paragraph()
    line_run = line_para.add_run('━' * 60)
    line_run.font.size = Pt(10)
    line_run.font.color.rgb = RGBColor(59, 130, 246)
    line_para.space_after = Pt(10)


def add_contact_item(cell, icon, text, label=None):
    """Add a contact item with icon and optional label"""
    para = cell.add_paragraph()
    
    # Icon with accent color
    icon_run = para.add_run(f"{icon} ")
    icon_run.font.size = Pt(9)
    
    # Label if provided
    if label:
        label_run = para.add_run(f"{label}\n")
        label_run.font.size = Pt(8)
        label_run.font.bold = True
        label_run.font.color.rgb = RGBColor(107, 114, 128)
    
    # Main text
    text_run = para.add_run(text)
    text_run.font.size = Pt(8)
    text_run.font.color.rgb = RGBColor(55, 65, 81)
    para.space_after = Pt(6)


def add_skill_item(cell, skill):
    """Add a skill item with bullet point"""
    para = cell.add_paragraph()
    bullet_run = para.add_run('▪ ')
    bullet_run.font.size = Pt(9)
    bullet_run.font.color.rgb = RGBColor(59, 130, 246)
    
    skill_run = para.add_run(skill)
    skill_run.font.size = Pt(9)
    skill_run.font.color.rgb = RGBColor(55, 65, 81)
    para.paragraph_format.left_indent = Inches(0.1)
    para.space_after = Pt(4)


def add_language_item(cell, language, proficiency):
    """Add a language item with proficiency level"""
    para = cell.add_paragraph()
    
    lang_run = para.add_run(language)
    lang_run.font.size = Pt(9)
    lang_run.font.bold = True
    lang_run.font.color.rgb = RGBColor(31, 41, 55)
    
    if proficiency:
        prof_run = para.add_run(f"\n  {proficiency}")
        prof_run.font.size = Pt(8)
        prof_run.font.italic = True
        prof_run.font.color.rgb = RGBColor(107, 114, 128)
    
    para.space_after = Pt(6)


@app.route('/download_resume_template')
def download_resume_template():
    """Download blank resume template with improved design"""
    try:
        doc = Document()
        
        # Set margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.6)
            section.bottom_margin = Inches(0.6)
            section.left_margin = Inches(0.7)
            section.right_margin = Inches(0.7)
        
        # ========== HEADER SECTION ==========
        # Add instruction note at top
        instruction_para = doc.add_paragraph()
        instruction_run = instruction_para.add_run('📋 RESUME TEMPLATE - Replace all [bracketed text] with your information')
        instruction_run.font.size = Pt(9)
        instruction_run.font.italic = True
        instruction_run.font.color.rgb = RGBColor(107, 114, 128)
        instruction_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        instruction_para.space_after = Pt(12)
        
        # Name
        name_para = doc.add_paragraph()
        name_run = name_para.add_run('[YOUR FULL NAME]')
        name_run.font.size = Pt(28)
        name_run.font.bold = True
        name_run.font.color.rgb = RGBColor(31, 41, 55)
        name_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        name_para.space_after = Pt(4)
        
        # Professional Title
        title_para = doc.add_paragraph()
        title_run = title_para.add_run('[Your Professional Title | Key Skills | Expertise Area]')
        title_run.font.size = Pt(13)
        title_run.font.color.rgb = RGBColor(59, 130, 246)
        title_run.font.bold = True
        title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        title_para.space_after = Pt(6)
        
        # Location
        location_para = doc.add_paragraph()
        location_run = location_para.add_run('📍 [City, Country/State]')
        location_run.font.size = Pt(10)
        location_run.font.color.rgb = RGBColor(107, 114, 128)
        location_para.space_after = Pt(8)
        
        # Separator
        separator_para = doc.add_paragraph()
        separator_run = separator_para.add_run('─' * 85)
        separator_run.font.color.rgb = RGBColor(59, 130, 246)
        separator_para.space_after = Pt(12)
        
        # ========== TWO COLUMN LAYOUT ==========
        table = doc.add_table(rows=1, cols=2)
        table.allow_autofit = False
        table.columns[0].width = Inches(2.2)
        table.columns[1].width = Inches(4.3)
        
        left_cell = table.rows[0].cells[0]
        right_cell = table.rows[0].cells[1]
        
        # Remove borders
        for cell in table.rows[0].cells:
            cell._element.get_or_add_tcPr().append(parse_xml(r'<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:top w:val="none"/><w:left w:val="none"/><w:bottom w:val="none"/><w:right w:val="none"/></w:tcBorders>'))
        
        # ========== LEFT SIDEBAR ==========
        # Contact Section
        add_section_header_improved(left_cell, "CONTACT")
        add_contact_item(left_cell, "✉", "[your.email@example.com]")
        add_contact_item(left_cell, "📞", "[+1 (123) 456-7890]")
        add_contact_item(left_cell, "🔗", "[linkedin.com/in/yourprofile]", label="LinkedIn")
        add_contact_item(left_cell, "🌐", "[yourportfolio.com]", label="Portfolio")
        left_cell.add_paragraph().space_after = Pt(8)
        
        # Skills Section
        add_section_header_improved(left_cell, "SKILLS")
        add_skill_item(left_cell, "[Skill 1 - e.g., Python]")
        add_skill_item(left_cell, "[Skill 2 - e.g., Project Management]")
        add_skill_item(left_cell, "[Skill 3 - e.g., Data Analysis]")
        add_skill_item(left_cell, "[Skill 4 - e.g., JavaScript]")
        add_skill_item(left_cell, "[Skill 5 - e.g., AWS]")
        left_cell.add_paragraph().space_after = Pt(8)
        
        # Languages Section
        add_section_header_improved(left_cell, "LANGUAGES")
        add_language_item(left_cell, "[Language 1 - e.g., English]", "[Native/Fluent/Professional]")
        add_language_item(left_cell, "[Language 2 - e.g., Spanish]", "[Conversational/Basic]")
        left_cell.add_paragraph().space_after = Pt(8)
        
        # ========== RIGHT COLUMN ==========
        # Professional Summary
        add_main_section_header_improved(right_cell, "PROFESSIONAL SUMMARY")
        summary_para = right_cell.add_paragraph()
        summary_run = summary_para.add_run('[Write 2-3 sentences highlighting your professional background, key achievements, and career objectives. Focus on what makes you unique and the value you bring to potential employers.]')
        summary_run.font.size = Pt(10)
        summary_run.font.color.rgb = RGBColor(55, 65, 81)
        summary_para.paragraph_format.line_spacing = 1.3
        summary_para.space_after = Pt(14)
        
        # Experience Section
        add_main_section_header_improved(right_cell, "WORK EXPERIENCE")
        
        # Experience 1
        company_para = right_cell.add_paragraph()
        company_run = company_para.add_run('[Company Name]')
        company_run.font.size = Pt(12)
        company_run.font.bold = True
        company_run.font.color.rgb = RGBColor(31, 41, 55)
        company_para.space_after = Pt(3)
        
        position_para = right_cell.add_paragraph()
        position_run = position_para.add_run('[Job Title/Position]')
        position_run.font.size = Pt(11)
        position_run.font.italic = True
        position_run.font.color.rgb = RGBColor(59, 130, 246)
        position_para.space_after = Pt(3)
        
        date_para = right_cell.add_paragraph()
        date_run = date_para.add_run('📅 [January 2020] – [December 2023] • [3 years 11 months]')
        date_run.font.size = Pt(9)
        date_run.font.color.rgb = RGBColor(107, 114, 128)
        date_para.space_after = Pt(2)
        
        loc_para = right_cell.add_paragraph()
        loc_run = loc_para.add_run('📍 [City, Country]')
        loc_run.font.size = Pt(9)
        loc_run.font.color.rgb = RGBColor(107, 114, 128)
        loc_para.space_after = Pt(6)
        
        # Job responsibilities
        responsibilities = [
            '[Achievement or responsibility using action verb - e.g., Led team of 5 developers]',
            '[Quantified achievement - e.g., Increased efficiency by 30%]',
            '[Key project or initiative - e.g., Implemented new CRM system]',
            '[Additional responsibility or achievement]'
        ]
        for resp in responsibilities:
            desc_para = right_cell.add_paragraph()
            desc_para.add_run('• ').font.color.rgb = RGBColor(59, 130, 246)
            desc_run = desc_para.add_run(resp)
            desc_run.font.size = Pt(10)
            desc_run.font.color.rgb = RGBColor(75, 85, 99)
            desc_para.paragraph_format.left_indent = Inches(0.15)
            desc_para.space_after = Pt(4)
        
        right_cell.add_paragraph().space_after = Pt(10)
        
        # Experience 2 (template)
        company_para2 = right_cell.add_paragraph()
        company_run2 = company_para2.add_run('[Previous Company Name]')
        company_run2.font.size = Pt(12)
        company_run2.font.bold = True
        company_run2.font.color.rgb = RGBColor(31, 41, 55)
        company_para2.space_after = Pt(3)
        
        position_para2 = right_cell.add_paragraph()
        position_run2 = position_para2.add_run('[Previous Job Title]')
        position_run2.font.size = Pt(11)
        position_run2.font.italic = True
        position_run2.font.color.rgb = RGBColor(59, 130, 246)
        position_para2.space_after = Pt(3)
        
        date_para2 = right_cell.add_paragraph()
        date_run2 = date_para2.add_run('📅 [Start Date] – [End Date] • [Duration]')
        date_run2.font.size = Pt(9)
        date_run2.font.color.rgb = RGBColor(107, 114, 128)
        date_para2.space_after = Pt(2)
        
        loc_para2 = right_cell.add_paragraph()
        loc_run2 = loc_para2.add_run('📍 [City, Country]')
        loc_run2.font.size = Pt(9)
        loc_run2.font.color.rgb = RGBColor(107, 114, 128)
        loc_para2.space_after = Pt(6)
        
        desc_para2 = right_cell.add_paragraph()
        desc_para2.add_run('• ').font.color.rgb = RGBColor(59, 130, 246)
        desc_run2 = desc_para2.add_run('[Add your key achievements and responsibilities here]')
        desc_run2.font.size = Pt(10)
        desc_run2.font.color.rgb = RGBColor(75, 85, 99)
        desc_para2.paragraph_format.left_indent = Inches(0.15)
        desc_para2.space_after = Pt(10)
        
        # Education Section
        add_main_section_header_improved(right_cell, "EDUCATION")
        
        inst_para = right_cell.add_paragraph()
        inst_run = inst_para.add_run('[University/Institution Name]')
        inst_run.font.size = Pt(12)
        inst_run.font.bold = True
        inst_run.font.color.rgb = RGBColor(31, 41, 55)
        inst_para.space_after = Pt(3)
        
        degree_para = right_cell.add_paragraph()
        degree_run = degree_para.add_run('[Degree Type - e.g., Bachelor of Science] – [Field of Study - e.g., Computer Science]')
        degree_run.font.size = Pt(10)
        degree_run.font.color.rgb = RGBColor(75, 85, 99)
        degree_para.space_after = Pt(3)
        
        grad_para = right_cell.add_paragraph()
        grad_run = grad_para.add_run('🎓 [Graduation Date - e.g., May 2020]')
        grad_run.font.size = Pt(9)
        grad_run.font.color.rgb = RGBColor(107, 114, 128)
        grad_para.space_after = Pt(2)
        
        edu_loc_para = right_cell.add_paragraph()
        edu_loc_run = edu_loc_para.add_run('📍 [City, Country]')
        edu_loc_run.font.size = Pt(9)
        edu_loc_run.font.color.rgb = RGBColor(107, 114, 128)
        edu_loc_para.space_after = Pt(10)
        
        # Certifications Section
        add_main_section_header_improved(right_cell, "CERTIFICATIONS")
        
        cert_para = right_cell.add_paragraph()
        cert_run = cert_para.add_run('🏆 [Certification Name - e.g., AWS Certified Solutions Architect]')
        cert_run.font.size = Pt(11)
        cert_run.font.bold = True
        cert_run.font.color.rgb = RGBColor(31, 41, 55)
        cert_para.space_after = Pt(3)
        
        issuer_para = right_cell.add_paragraph()
        issuer_run = issuer_para.add_run('[Issuing Organization - e.g., Amazon Web Services]')
        issuer_run.font.size = Pt(10)
        issuer_run.font.color.rgb = RGBColor(75, 85, 99)
        issuer_para.paragraph_format.left_indent = Inches(0.15)
        issuer_para.space_after = Pt(2)
        
        cert_date_para = right_cell.add_paragraph()
        cert_date_run = cert_date_para.add_run('Issued: [Month Year - e.g., June 2023]')
        cert_date_run.font.size = Pt(9)
        cert_date_run.font.color.rgb = RGBColor(107, 114, 128)
        cert_date_para.paragraph_format.left_indent = Inches(0.15)
        cert_date_para.space_after = Pt(10)
        
        # Additional Section Example
        add_main_section_header_improved(right_cell, "PROJECTS / AWARDS (OPTIONAL)")
        
        project_para = right_cell.add_paragraph()
        project_para.add_run('• ').font.color.rgb = RGBColor(59, 130, 246)
        project_run = project_para.add_run('[Project or Award Name] - [Brief description]')
        project_run.font.size = Pt(10)
        project_run.font.color.rgb = RGBColor(75, 85, 99)
        project_para.paragraph_format.left_indent = Inches(0.15)
        project_para.space_after = Pt(4)
        
        # Save to BytesIO
        doc_io = BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        
        return send_file(
            doc_io,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name='Resume_Template_Professional.docx'
        )
    
    except Exception as e:
        logger.error(f"Error generating template: {e}")
        flash("Error generating template", "error")
        return redirect('/resume_builder')

# Helper functions
def add_section_header(cell, text):
    """Add a sidebar section header"""
    para = cell.add_paragraph()
    run = para.add_run(text.upper())
    run.font.size = Pt(9)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    para.space_after = Pt(6)

def add_main_section_header(cell, text):
    """Add a main section header with underline"""
    para = cell.add_paragraph()
    run = para.add_run(text.upper())
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    para.space_after = Pt(8)
    
    # Add line below
    line_para = cell.add_paragraph()
    line_para.add_run('_' * 60)
    line_para.space_after = Pt(8)

def add_small_text(cell, text, italic=False):
    """Add small text to cell"""
    para = cell.add_paragraph()
    run = para.add_run(text)
    run.font.size = Pt(8)
    if italic:
        run.font.italic = True
    para.space_after = Pt(2)





















    
    
@app.route("/deleteResume/<resume_id>", methods=["POST"])
def deleteResume(resume_id):
    try:
        if 'user_id' not in session:
            flash("Please log in first", "error")
            return redirect(url_for('loginpage'))

        # Validate resume_id
        try:
            ObjectId(resume_id)
        except:
            flash("Invalid resume ID", "error")
            return redirect('/profile')

        try:
            # Find resume by ID and user ID (security check)
            resume = resumeFetchedData.find_one({
                "_id": ObjectId(resume_id), 
                "UserId": ObjectId(session['user_id'])
            })
            
            if not resume:
                flash("Resume not found", "error")
                return redirect('/profile')

            # Delete file from filesystem
            if resume.get("ResumeTitle"):
                filepath = os.path.join(app.config['UPLOAD_FOLDER'], resume["ResumeTitle"])
                if os.path.exists(filepath):
                    try:
                        os.remove(filepath)
                        logger.info(f"Deleted resume file: {filepath}")
                    except Exception as e:
                        logger.error(f"Error deleting file {filepath}: {e}")

            # Delete from database
            resumeFetchedData.delete_one({"_id": ObjectId(resume_id)})
            logger.info(f"Deleted resume {resume_id} for user {session['user_id']}")

            # Delete related applications
            try:
                Applied_EMP.delete_many({"resume_id": ObjectId(resume_id)})
                logger.info(f"Deleted applications for resume {resume_id}")
            except Exception as e:
                logger.error(f"Error deleting applications for resume {resume_id}: {e}")

            flash("Resume deleted successfully!", "success")
            
        except Exception as e:
            logger.error(f"Error deleting resume {resume_id}: {e}")
            flash("Error occurred while deleting resume", "error")

        return redirect('/profile')

    except Exception as e:
        logger.error(f"Unexpected error in deleteResume: {e}")
        flash("An unexpected error occurred", "error")
        return redirect('/profile')

@app.route('/HR1/get_user_resumes')
def get_user_resumes():
    try:
        if 'user_id' not in session:
            logger.warning("Session missing in /HR1/get_user_resumes")
            return jsonify({"error": "Not logged in"}), 401

        user_id = session['user_id']
        logger.info(f"Fetching resumes for user: {user_id}")

        try:
            resumes = resumeFetchedData.find({'UserId': ObjectId(user_id)})
            result = []
            
            for resume in resumes:
                try:
                    result.append({
                        '_id': str(resume['_id']),
                        'ResumeTitle': resume.get('ResumeTitle', 'Untitled'),
                        'UploadedAt': resume.get('UploadedAt', ''),
                        'Name': resume.get('Name', 'Unknown')
                    })
                except Exception as e:
                    logger.error(f"Error processing resume {resume.get('_id')}: {e}")
                    continue

            logger.info(f"Returning {len(result)} resumes for user {user_id}")
            return jsonify(result)
            
        except Exception as e:
            logger.error(f"Database error fetching resumes for user {user_id}: {e}")
            return jsonify({"error": "Database error"}), 500

    except Exception as e:
        logger.error(f"Unexpected error in get_user_resumes: {e}")
        return jsonify({"error": "Server error"}), 500

@app.route('/viewdetails', methods=['POST', 'GET'])
def viewdetails():
    try:
        employee_id = request.form.get('employee_id', '').strip()
        
        if not employee_id:
            return jsonify({'error': 'Employee ID is required'})

        try:
            ObjectId(employee_id)
        except:
            return jsonify({'error': 'Invalid employee ID format'})

        try:
            result = resumeFetchedData.find_one({"UserId": ObjectId(employee_id)})
            
            if not result:
                return jsonify({'error': 'Employee not found'})

            response_data = {
                'name': result.get('Name'),
                'linkedin_link': result.get('LINKEDIN LINK'),
                'skills': result.get('SKILLS', []),
                'certificate': result.get('CERTIFICATION', [])
            }
            
            logger.info(f"Retrieved details for employee {employee_id}")
            return jsonify(response_data)
            
        except Exception as e:
            logger.error(f"Database error fetching employee details {employee_id}: {e}")
            return jsonify({'error': 'Database error'})

    except Exception as e:
        logger.error(f"Error in viewdetails: {e}")
        return jsonify({'error': 'Server error'})

@app.route("/empSearch", methods=['POST'])
def empSearch():
    try:
        if 'user_id' not in session or session.get('role') != 'employer':
            return render_template("Company_Job_Posting.html", errorMsg="Access denied")

        category = request.form.get('category', '').strip()
        
        if not category:
            return render_template("Company_Job_Posting.html", errorMsg="Category is required")

        logger.info(f"Searching employees for category: {category}")

        try:
            # Find jobs matching the category
            job_cursor = JOBS.find({"Job_Profile": {"$regex": category, "$options": "i"}}, {"_id": 1})
            job_ids = [job['_id'] for job in job_cursor]

            if not job_ids:
                return render_template("Company_Job_Posting.html", 
                                     errorMsg="No jobs found for this category", 
                                     len=0, data={})

            # Find applied employees for these jobs
            TopEmployeers = Applied_EMP.find(
                {"job_id": {"$in": job_ids}},
                {"user_id": 1, "Matching_percentage": 1}
            ).sort([("Matching_percentage", -1)])

            selectedResumes = {}
            cnt = 0

            for application in TopEmployeers:
                try:
                    user_data = IRS_USERS.find_one(
                        {"_id": ObjectId(application['user_id'])},
                        {"Name": 1, "Email": 1, "_id": 1}
                    )
                    
                    if user_data:
                        selectedResumes[cnt] = {
                            "Name": user_data.get('Name', 'Unknown'),
                            "Email": user_data.get('Email', 'Unknown'),
                            "_id": str(user_data['_id']),
                            "Matching_percentage": application.get('Matching_percentage', 0)
                        }
                        cnt += 1
                        
                except Exception as e:
                    logger.error(f"Error processing application {application.get('user_id')}: {e}")
                    continue

            logger.info(f"Found {len(selectedResumes)} candidates for category {category}")
            return render_template("Company_Job_Posting.html", 
                                 len=len(selectedResumes), 
                                 data=selectedResumes)

        except Exception as e:
            logger.error(f"Database error in employee search: {e}")
            return render_template("Company_Job_Posting.html", 
                                 errorMsg="Database error occurred")

    except Exception as e:
        logger.error(f"Error in empSearch: {e}")
        return render_template("Company_Job_Posting.html", 
                             errorMsg="An unexpected error occurred")






# @app.route("/feedback", methods=["POST"])
# def add_feedback():
#     """Collect recruiter feedback for learning"""
#     # from enhanced_ranking_system import enhanced_ranking_system
    
#     try:
#         data = request.get_json()
        
#         recruiter_id = session.get('user_id')
#         candidate_id = data.get('candidate_id')
#         feedback_type = data.get('feedback_type')  # 'ranking', 'interview', 'hire'
#         rating = float(data.get('rating', 0))
#         comments = data.get('comments', '')
        
#         if not all([recruiter_id, candidate_id, feedback_type, rating]):
#             return jsonify({"success": False, "message": "Missing required fields"}), 400
        
#         if enhanced_ranking_system:
#             # Record feedback
#             enhanced_ranking_system.feedback_system.record_feedback(
#                 candidate_id, recruiter_id, feedback_type, rating, comments
#             )
            
#             # Update weights if enough feedback
#             job_id = data.get('job_id')
#             if job_id:
#                 new_weights = enhanced_ranking_system.feedback_system.adjust_ranking_weights(
#                     job_id, recruiter_id
#                 )
#                 if new_weights:
#                     return jsonify({
#                         "success": True,
#                         "message": "Feedback recorded and AI weights updated",
#                         "new_weights": new_weights
#                     })
            
#             return jsonify({
#                 "success": True,
#                 "message": "Feedback recorded successfully"
#             })
#         else:
#             return jsonify({"success": False, "message": "Enhanced system not available"}), 503
        
#     except Exception as e:
#         logger.error(f"Feedback recording failed: {e}")
#         return jsonify({"success": False, "message": str(e)}), 500

@app.route("/rank_candidates", methods=["POST"])
def rank_candidates():
    """AI-powered candidate ranking using Sentence Transformers"""
    from sentence_transformer_ranker import get_semantic_ranker
    
    try:
        data = request.get_json()
        
        job_id = data.get('job_id')
        resume_ids = data.get('resume_ids', [])
        recruiter_id = session.get('user_id')
        
        if not job_id or not resume_ids:
            return jsonify({"success": False, "message": "Missing job_id or resume_ids"}), 400
        
        # Get semantic ranker instance
        ranker = get_semantic_ranker()
        
        if not ranker.is_available():
            return jsonify({"success": False, "message": "Semantic ranking system not available"}), 503
        
        # Prepare MongoDB collections
        mongo_collections = {
            'JOBS': JOBS,
            'Applied_EMP': Applied_EMP,
            'resumeFetchedData': resumeFetchedData
        }
        
        # Get job details
        job = JOBS.find_one({"_id": ObjectId(job_id)})
        if not job:
            return jsonify({"success": False, "message": "Job not found"}), 404
        
        # Extract job requirements once
        job_text = ranker.extract_job_requirements(job)
        
        # Process each resume
        ranked_candidates = []
        for resume_id in resume_ids:
            try:
                # Get resume data
                resume = resumeFetchedData.find_one({"_id": ObjectId(resume_id)})
                if not resume:
                    logger.warning(f"Resume {resume_id} not found")
                    continue
                
                # Get application data for traditional score
                application = Applied_EMP.find_one({
                    "job_id": ObjectId(job_id),
                    "resume_id": ObjectId(resume_id)
                })
                
                # Extract resume content
                resume_text = ranker.extract_resume_content(resume)
                
                # Calculate semantic similarity
                semantic_score = ranker.calculate_semantic_similarity(job_text, resume_text)
                
                # Get traditional score from application
                traditional_score = 0
                if application:
                    traditional_score = application.get('Matching_percentage', 0)
                    if isinstance(traditional_score, dict):
                        traditional_score = traditional_score.get('overall_score', 0)
                
                # Calculate combined score
                combined_score = ranker.calculate_combined_score(
                    traditional_score, 
                    semantic_score * 100  # Convert 0-1 to 0-100
                )
                
                ranked_candidates.append({
                    "resume_id": resume_id,
                    "candidate_name": resume.get('Name', 'Unknown'),
                    "traditional_score": round(float(traditional_score), 2),
                    "semantic_score": round(semantic_score * 100, 2),  # Convert to percentage
                    "combined_score": round(combined_score, 2),
                    "skills": resume.get('SKILLS', [])[:5],  # Top 5 skills
                    "experience": resume.get('WORKED AS', [])[:3],  # Top 3 positions
                    "years_experience": resume.get('YEARS OF EXPERIENCE', []),
                })
                
            except Exception as e:
                logger.error(f"Error processing resume {resume_id}: {e}")
                continue
        
        # Sort by combined score (descending)
        ranked_candidates.sort(key=lambda x: x['combined_score'], reverse=True)
        
        # Add rank numbers
        for i, candidate in enumerate(ranked_candidates):
            candidate['rank'] = i + 1
        
        return jsonify({
            "success": True,
            "total_candidates": len(ranked_candidates),
            "ranked_candidates": ranked_candidates,
            "ai_features_used": ["semantic_similarity", "sentence_transformers", "combined_scoring"],
            "processing_metadata": {
                "timestamp": datetime.now().isoformat(),
                "recruiter_id": recruiter_id,
                "job_id": job_id,
                "model_used": ranker.model_name
            }
        })
        
    except Exception as e:
        logger.error(f"Batch ranking failed: {e}")
        import traceback
        logger.error(traceback.format_exc())
        return jsonify({"success": False, "message": str(e)}), 500





def get_clean_dashboard_stats(user_id=None, role=None):
    """Get dashboard statistics with only legitimate metrics"""
    try:
        stats = {}
        
        if role == 'applicant' and user_id:
            apps = list(Applied_EMP.find({"user_id": ObjectId(user_id)}))
            
            # LEGITIMATE metrics only
            stats = {
                'total_applications': len(apps),
                'avg_match_score': round(sum(app.get('Matching_percentage', 0) for app in apps) / len(apps), 1) if apps else 0,
                'avg_semantic_score': round(sum(app.get('semantic_similarity', 0) for app in apps) / len(apps), 1) if apps else 0,
                'interviews_scheduled': len([app for app in apps if app.get('status') in ['interviewed', 'interview_scheduled']]),
                
                'ai_enhanced_applications': len([app for app in apps if app.get('enhanced_version')])
            }
            
        elif role == 'employer':
            all_apps = list(Applied_EMP.find({}))
            
            stats = {
                'total_applications': len(all_apps),
                'avg_match_score': round(sum(app.get('Matching_percentage', 0) for app in all_apps) / len(all_apps), 1) if all_apps else 0,
                'semantic_analysis_count': len([app for app in all_apps if app.get('semantic_similarity') is not None]),
                
                'total_jobs': JOBS.count_documents({}),
                'active_jobs': JOBS.count_documents({"Status": "Open"})
            }
        
        # Add legitimate system performance metrics
        # stats['legitimate_ai_features'] = list(LegitimateAIFeatures.FEATURES.keys())
        stats['system_version'] = '3.0_legitimate'
        stats['timestamp'] = datetime.now().isoformat()
        
        return stats
        
    except Exception as e:
        logger.error(f"Error calculating clean dashboard stats: {e}")
        return {'error': 'Unable to calculate statistics'}

print("✅ All fake AI components removed")















# Add these routes to your app.py file
@app.route('/interviews')
def applicant_interviews():
    """FIXED: Applicant interview dashboard with proper date handling"""
    try:
        if 'user_id' not in session or session.get('role') != 'applicant':
            flash("Access denied. Applicant login required.", "error")
            return redirect(url_for('loginpage'))
        
        user_id = session['user_id']
        current_date = datetime.now()
        
        # FIXED QUERY - Include interview_pending and shortlisted candidates
        interview_applications = list(Applied_EMP.find({
            "user_id": ObjectId(user_id),
            "$or": [
                {"status": {"$in": ["interview_scheduled", "interviewed", "interview_pending", "shortlisted", "under_review"]}},
                {"interview_date": {"$exists": True}}
            ]
        }).sort([("applied_at", -1)]))
        
        logger.info(f"Found {len(interview_applications)} interview applications for user {user_id}")
        
        # Process interviews data
        upcoming_interviews = []
        past_interviews = []
        pending_interviews = []
        
        for app in interview_applications:
            try:
                # Get job details
                job = JOBS.find_one({"_id": app.get('job_id')}, {
                    "Job_Profile": 1, "CompanyName": 1, "Job_Description": 1
                })
                
                if not job:
                    continue
                
                # FIXED: Proper date handling for template
                interview_date_obj = app.get('interview_date')
                interview_date_display = None
                
                if interview_date_obj:
                    if isinstance(interview_date_obj, datetime):
                        interview_date_display = interview_date_obj
                    elif isinstance(interview_date_obj, str):
                        try:
                            interview_date_display = datetime.fromisoformat(interview_date_obj.replace('Z', '+00:00'))
                        except:
                            interview_date_display = None
                
                interview_data = {
                    'application_id': str(app['_id']),
                    'job_id': str(app.get('job_id')),
                    'job_title': job.get('Job_Profile', 'Unknown Job'),
                    'company_name': job.get('CompanyName', 'Company'),
                    'job_description': job.get('Job_Description', '')[:200] + "...",
                    'applied_date': app.get('applied_at'),
                    'match_score': round(get_match_score(app), 1),
                    'status': app.get('status'),
                    'interview_date': interview_date_display,  # Pass as datetime object
                    'interview_time': app.get('interview_time'),
                    'interview_type': app.get('interview_type', 'Panel Interview'),
                    'interview_location': app.get('interview_location'),
                    'meeting_link': app.get('meeting_link'),
                    'interviewer_name': app.get('interviewer_name'),
                    'interviewer_email': app.get('interviewer_email'),
                    'notes': app.get('interview_notes', ''),
                    'feedback_received': app.get('interview_feedback'),
                    'ai_enhanced': bool(app.get('enhanced_version'))
                }
                
                # IMPROVED CATEGORIZATION LOGIC
                status = app.get('status', '')
                
                logger.info(f"Processing application {app.get('_id')}: status={status}, has_date={bool(interview_date_display)}")
                
                # Check if it's pending (shortlisted or interview_pending without date)
                if status in ['shortlisted', 'interview_pending'] and not interview_date_display:
                    pending_interviews.append(interview_data)
                    logger.info(f"Added to pending: {interview_data['job_title']}")
                elif status == 'interview_scheduled' and interview_date_display:
                    # Check if upcoming or past
                    if interview_date_display > current_date:
                        upcoming_interviews.append(interview_data)
                        logger.info(f"Added to upcoming: {interview_data['job_title']}")
                    else:
                        past_interviews.append(interview_data)
                        logger.info(f"Added to past: {interview_data['job_title']}")
                elif status == 'interviewed':
                    past_interviews.append(interview_data)
                    logger.info(f"Added to past (interviewed): {interview_data['job_title']}")
                elif status == 'under_review':
                    # Under review candidates might be pending interview
                    pending_interviews.append(interview_data)
                    logger.info(f"Added to pending (under review): {interview_data['job_title']}")
                else:
                    # Fallback - if status suggests interview process, add to pending
                    if status in ['shortlisted', 'interview_pending']:
                        pending_interviews.append(interview_data)
                        logger.info(f"Added to pending (fallback): {interview_data['job_title']}")
                    
            except Exception as e:
                logger.error(f"Error processing interview application {app.get('_id')}: {e}")
                continue
        
        # Sort interviews
        upcoming_interviews.sort(key=lambda x: x.get('interview_date', datetime.now()))
        past_interviews.sort(key=lambda x: x.get('interview_date', datetime.now()), reverse=True)
        
        # Statistics
        stats = {
            'total_interviews': len(interview_applications),
            'upcoming_count': len(upcoming_interviews),
            'completed_count': len(past_interviews),
            'pending_count': len(pending_interviews),
            'success_rate': 0
        }
        
        # Calculate success rate
        if len(past_interviews) > 0:
            successful_interviews = Applied_EMP.count_documents({
                "user_id": ObjectId(user_id),
                "status": {"$in": ["offer_made", "hired", "shortlisted_final"]}
            })
            stats['success_rate'] = round((successful_interviews / len(past_interviews)) * 100, 1)
        
        logger.info(f"Applicant {user_id} interviews: pending={len(pending_interviews)}, upcoming={len(upcoming_interviews)}, past={len(past_interviews)}")
        
        return render_template("applicant_interviews.html",
                             upcoming_interviews=upcoming_interviews,
                             past_interviews=past_interviews,
                             pending_interviews=pending_interviews,
                             stats=stats,
                             current_date=current_date)
                             
    except Exception as e:
        logger.error(f"Error in applicant_interviews: {e}")
        import traceback
        logger.error(f"Full traceback: {traceback.format_exc()}")
        return render_template("applicant_interviews.html",
                             upcoming_interviews=[],
                             past_interviews=[],
                             pending_interviews=[],
                             stats={},
                             current_date=datetime.now(),
                             errorMsg="An error occurred loading interviews")


# In hr_interviews route, fix the query to properly include interview_pending candidates

@app.route('/HR1/interviews')
def hr_interviews():
    """FIXED: HR interview management with detailed debugging"""
    try:
        if 'user_id' not in session or session.get('role') != 'employer':
            flash("Access denied. Employer login required.", "error")
            return redirect(url_for('loginpage'))
        
        current_hr_user_id = session['user_id']
        current_date = datetime.now()
        
        # Get filter parameters
        job_filter = request.args.get('job_id')
        status_filter = request.args.get('status', 'all')
        date_filter = request.args.get('date_range', 'all')
        
        # STEP 1: Get HR user's jobs with debugging
        hr_jobs = list(JOBS.find({"created_by": ObjectId(current_hr_user_id)}, {"_id": 1, "Job_Profile": 1}))
        hr_job_ids = [job['_id'] for job in hr_jobs]
        
        logger.info(f"DEBUG: HR User {current_hr_user_id} has {len(hr_jobs)} jobs")
        for job in hr_jobs:
            logger.info(f"DEBUG: Job {job['_id']}: {job.get('Job_Profile', 'No title')}")
        
        if not hr_job_ids:
            logger.warning(f"No jobs found for HR user {current_hr_user_id}")
            return render_template("hr_interviews.html",
                                 interviews=[],
                                 jobs=[],
                                 stats={},
                                 current_date=current_date,
                                 filters={'job_id': job_filter, 'status': status_filter, 'date_range': date_filter})
        
        # STEP 2: Check applications for these jobs
        all_apps_for_jobs = list(Applied_EMP.find({"job_id": {"$in": hr_job_ids}}, {"status": 1, "job_id": 1, "User_name": 1}))
        logger.info(f"DEBUG: Found {len(all_apps_for_jobs)} total applications for HR user's jobs")
        
        # Log status breakdown
        status_counts = {}
        for app in all_apps_for_jobs:
            status = app.get('status', 'unknown')
            status_counts[status] = status_counts.get(status, 0) + 1
        logger.info(f"DEBUG: Application status breakdown: {status_counts}")
        
        # STEP 3: Build query with broader criteria
        query = {"job_id": {"$in": hr_job_ids}}
        
        if status_filter == 'all':
            # Include MORE statuses to catch candidates
            interview_statuses = [
                "interview_scheduled", "interviewed", "interview_pending", 
                "shortlisted", "under_review"  # Added under_review
            ]
            query["$or"] = [
                {"status": {"$in": interview_statuses}},
                {"interview_date": {"$exists": True}}
            ]
            logger.info(f"DEBUG: Using broad status filter: {interview_statuses}")
        else:
            query["status"] = status_filter
            logger.info(f"DEBUG: Filtering by specific status: {status_filter}")
        
        # Apply job filter
        if job_filter:
            try:
                if ObjectId(job_filter) in hr_job_ids:
                    query["job_id"] = ObjectId(job_filter)
                    logger.info(f"DEBUG: Filtering by job: {job_filter}")
                else:
                    job_filter = None
                    logger.warning(f"Job filter {job_filter} not in HR user's jobs")
            except:
                job_filter = None
                logger.error(f"Invalid job_filter format: {job_filter}")
        
        # Date filtering (simplified for debugging)
        if date_filter != 'all':
            logger.info(f"DEBUG: Date filter applied: {date_filter}")
            # Keep your existing date logic here
        
        logger.info(f"DEBUG: Final query: {query}")
        
        # STEP 4: Execute query
        interview_applications = list(Applied_EMP.find(query).sort([("applied_at", -1)]))
        logger.info(f"DEBUG: Query returned {len(interview_applications)} applications")
        
        # Log what we found
        for app in interview_applications:
            logger.info(f"DEBUG: Found application - Status: {app.get('status')}, Job: {app.get('job_id')}, User: {app.get('User_name', 'Unknown')}")
        
        # STEP 5: Process interviews
        interviews = []
        for app in interview_applications:
            try:
                candidate = IRS_USERS.find_one({"_id": app.get('user_id')}, {"Name": 1, "Email": 1})
                job = JOBS.find_one({"_id": app.get('job_id')}, {"Job_Profile": 1, "CompanyName": 1})
                
                if not candidate:
                    logger.warning(f"Candidate not found for user_id: {app.get('user_id')}")
                    continue
                if not job:
                    logger.warning(f"Job not found for job_id: {app.get('job_id')}")
                    continue
                
                # FIXED: Safe handling of None values for numeric fields
                def safe_round(value, default=0, digits=1):
                    """Safely round a value, handling None and non-numeric values"""
                    try:
                        if value is None:
                            return default
                        return round(float(value), digits)
                    except (ValueError, TypeError):
                        return default
                
                interview_data = {
                    'application_id': str(app['_id']),
                    'candidate_id': str(app.get('user_id')),
                    'candidate_name': candidate.get('Name', 'Unknown'),
                    'candidate_email': candidate.get('Email', ''),
                    'job_id': str(app.get('job_id')),
                    'job_title': job.get('Job_Profile', 'Unknown Job'),
                    'company_name': job.get('CompanyName', 'Company'),
                    # FIXED: Safe rounding with None handling
                    'match_score': safe_round(app.get('Matching_percentage'), 0),
                    'semantic_score': safe_round(app.get('semantic_similarity'), 0),
                    'status': app.get('status'),
                    'applied_date': app.get('applied_at'),
                    'interview_date': app.get('interview_date'),
                    'interview_time': app.get('interview_time'),
                    'interview_type': app.get('interview_type', 'Panel Interview'),
                    'interview_location': app.get('interview_location'),
                    'meeting_link': app.get('meeting_link'),
                    'interviewer_name': app.get('interviewer_name'),
                    'interviewer_email': app.get('interviewer_email'),
                    'notes': app.get('interview_notes', ''),
                    'feedback_submitted': bool(app.get('interview_feedback')),
                    'ai_enhanced': bool(app.get('enhanced_version'))
                }
                
                interviews.append(interview_data)
                logger.info(f"DEBUG: Successfully processed interview for {candidate.get('Name')} - {app.get('status')}")
                
            except Exception as e:
                logger.error(f"Error processing application {app.get('_id')}: {e}")
                # Log the actual data to see what's causing issues
                logger.error(f"Application data: Matching_percentage={app.get('Matching_percentage')}, semantic_similarity={app.get('semantic_similarity')}")
                continue
        
        # Get jobs for dropdown
        jobs = list(JOBS.find({"created_by": ObjectId(current_hr_user_id)}, {"_id": 1, "Job_Profile": 1, "Status": 1}))
        
        # Statistics
        stats = {
            'total_interviews': len(interviews),
            'today_interviews': 0,
            'pending_feedback': len([i for i in interviews if not i.get('feedback_submitted')]),
            'completed_interviews': len([i for i in interviews if i.get('status') == 'interviewed']),
            'pending_schedule': len([i for i in interviews if i.get('status') == 'interview_pending'])
        }
        
        logger.info(f"DEBUG: Final results - {len(interviews)} interviews processed, {stats['pending_schedule']} pending schedule")
        
        return render_template("hr_interviews.html",
                             interviews=interviews,
                             jobs=jobs,
                             stats=stats,
                             current_date=current_date,
                             filters={
                                 'job_id': job_filter,
                                 'status': status_filter,
                                 'date_range': date_filter
                             })
                             
    except Exception as e:
        logger.error(f"Error in hr_interviews: {e}")
        import traceback
        logger.error(f"Full traceback: {traceback.format_exc()}")
        return render_template("hr_interviews.html",
                             interviews=[],
                             jobs=[],
                             stats={},
                             current_date=datetime.now(),
                             filters={},
                             errorMsg="An error occurred loading interviews")


@app.route("/update_candidate_status", methods=["POST"])
def update_candidate_status():
    """Update candidate status with notifications"""
    try:
        if 'user_id' not in session or session.get('role') != 'employer':
            return jsonify({"success": False, "message": "Access denied"}), 403
        
        data = request.get_json()
        application_id = data.get('application_id')
        new_status = data.get('status')
        
        # Get application details
        application = Applied_EMP.find_one({"_id": ObjectId(application_id)})
        if not application:
            return jsonify({"success": False, "message": "Application not found"})
        
        old_status = application.get('status', 'pending')
        
        # Update status
        update_data = {
            "status": new_status,
            "status_updated_at": datetime.now(),
            "status_updated_by": session.get('user_id')
        }
        
        result = Applied_EMP.update_one(
            {"_id": ObjectId(application_id)},
            {"$set": update_data}
        )
        
        if result.modified_count > 0:
            # Get job title for notification
            job = JOBS.find_one({"_id": application['job_id']}, {"Job_Profile": 1})
            job_title = job.get('Job_Profile', 'Unknown Job') if job else 'Unknown Job'
            
            # CREATE STATUS UPDATE NOTIFICATION
            create_status_update_notification(
                applicant_id=application['user_id'],
                job_title=job_title,
                old_status=old_status,
                new_status=new_status,
                job_id=application['job_id']
            )
            
            return jsonify({"success": True, "message": f"Status updated to {new_status}"})
        else:
            return jsonify({"success": False, "message": "No changes made"})
            
    except Exception as e:
        logger.error(f"Error updating candidate status: {e}")
        return jsonify({"success": False, "message": "Error updating status"})


@app.route('/submit_interview_feedback', methods=['POST'])
def submit_interview_feedback():
    """Submit interview feedback (HR only)"""
    try:
        if 'user_id' not in session or session.get('role') != 'employer':
            return jsonify({"success": False, "message": "Access denied"}), 403
        
        data = request.get_json()
        application_id = data.get('application_id')
        feedback = data.get('feedback', {})
        
        if not application_id:
            return jsonify({"success": False, "message": "Application ID required"}), 400
        
        # Verify application exists
        application = Applied_EMP.find_one({"_id": ObjectId(application_id)})
        if not application:
            return jsonify({"success": False, "message": "Application not found"}), 404
        
        # Structure feedback data
        feedback_data = {
            "interview_feedback": {
                "overall_rating": feedback.get('overall_rating'),
                "technical_skills": feedback.get('technical_skills'),
                "communication": feedback.get('communication'),
                "cultural_fit": feedback.get('cultural_fit'),
                "experience_relevance": feedback.get('experience_relevance'),
                "strengths": feedback.get('strengths', ''),
                "areas_for_improvement": feedback.get('areas_for_improvement', ''),
                "recommendation": feedback.get('recommendation', ''),
                "additional_notes": feedback.get('additional_notes', ''),
                "submitted_by": session.get('user_id'),
                "submitted_at": datetime.now()
            },
            "status": "interviewed",
            "feedback_submitted_at": datetime.now()
        }
        
        result = Applied_EMP.update_one(
            {"_id": ObjectId(application_id)},
            {"$set": feedback_data}
        )
        
        if result.modified_count > 0:
            logger.info(f"Feedback submitted for application {application_id}")
            return jsonify({"success": True, "message": "Feedback submitted successfully"})
        else:
            return jsonify({"success": False, "message": "Failed to submit feedback"}), 500
            
    except Exception as e:
        logger.error(f"Error submitting feedback: {e}")
        return jsonify({"success": False, "message": f"Error submitting feedback: {str(e)}"}), 500


@app.route('/get_interview_details/<application_id>')
def get_interview_details(application_id):
    """Get detailed interview information"""
    try:
        if 'user_id' not in session:
            return jsonify({"error": "Access denied"}), 403
        
        try:
            application = Applied_EMP.find_one({"_id": ObjectId(application_id)})
        except:
            return jsonify({"error": "Invalid application ID"}), 400
            
        if not application:
            return jsonify({"error": "Interview not found"}), 404
        
        # Check permission (candidate can only see their own, HR can see all)
        user_role = session.get('role')
        if user_role == 'applicant' and str(application.get('user_id')) != session.get('user_id'):
            return jsonify({"error": "Access denied"}), 403
        
        # Get related data
        candidate = IRS_USERS.find_one({"_id": application.get('user_id')}, {"Name": 1, "Email": 1})
        job = JOBS.find_one({"_id": application.get('job_id')}, {"Job_Profile": 1, "CompanyName": 1})
        
        interview_details = {
            'application_id': str(application['_id']),
            'candidate_name': candidate.get('Name', 'Unknown') if candidate else 'Unknown',
            'candidate_email': candidate.get('Email', '') if candidate else '',
            'job_title': job.get('Job_Profile', 'Unknown Job') if job else 'Unknown Job',
            'company_name': job.get('CompanyName', 'Company') if job else 'Company',
            'interview_date': application.get('interview_date').isoformat() if application.get('interview_date') else None,
            'interview_time': application.get('interview_time'),
            'interview_type': application.get('interview_type'),
            'interview_location': application.get('interview_location'),
            'meeting_link': application.get('meeting_link'),
            'interviewer_name': application.get('interviewer_name'),
            'interviewer_email': application.get('interviewer_email'),
            'notes': application.get('interview_notes', ''),
            'feedback': application.get('interview_feedback'),
            'status': application.get('status')
        }
        
        return jsonify(interview_details)
        
    except Exception as e:
        logger.error(f"Error getting interview details: {e}")
        return jsonify({"error": f"Server error: {str(e)}"}), 500
    



# Add these missing routes to your app.py file

@app.route('/schedule_interview', methods=['POST'])
def schedule_interview():
    """Schedule interview with notifications"""
    try:
        if 'user_id' not in session or session.get('role') != 'employer':
            return jsonify({"success": False, "message": "Access denied"}), 403
        
        data = request.get_json()
        application_id = data.get('application_id')
        
        # Get application details
        application = Applied_EMP.find_one({"_id": ObjectId(application_id)})
        if not application:
            return jsonify({"success": False, "message": "Application not found"})
        
        # Prepare interview data
        interview_date = datetime.fromisoformat(data.get('interview_date')) if data.get('interview_date') else None
        
        interview_data = {
            "status": "interview_scheduled",
            "interview_date": interview_date,
            "interview_time": data.get('interview_time'),
            "interview_type": data.get('interview_type'),
            "interview_location": data.get('interview_location'),
            "interviewer_name": data.get('interviewer_name'),
            "interview_scheduled_at": datetime.now(),
            "interview_scheduled_by": session.get('user_id')
        }
        
        # Update application
        result = Applied_EMP.update_one(
            {"_id": ObjectId(application_id)},
            {"$set": interview_data}
        )
        
        if result.modified_count > 0:
            # Get job details for notification
            job = JOBS.find_one({"_id": application['job_id']}, {"Job_Profile": 1})
            job_title = job.get('Job_Profile', 'Unknown Job') if job else 'Unknown Job'
            
            # CREATE INTERVIEW NOTIFICATIONS
            create_interview_notification(
                applicant_id=application['user_id'],
                hr_id=session['user_id'],
                job_title=job_title,
                interview_date=interview_date,
                interview_type=data.get('interview_type', 'Interview')
            )
            
            return jsonify({"success": True, "message": "Interview scheduled successfully"})
        else:
            return jsonify({"success": False, "message": "Failed to schedule interview"})
            
    except Exception as e:
        logger.error(f"Error scheduling interview: {e}")
        return jsonify({"success": False, "message": "Error scheduling interview"})

@app.route('/update_interview', methods=['POST'])
def update_interview():
    """Update interview details"""
    try:
        if 'user_id' not in session or session.get('role') != 'employer':
            return jsonify({"success": False, "message": "Access denied"}), 403
        
        data = request.get_json()
        application_id = data.get('application_id')
        
        if not application_id:
            return jsonify({"success": False, "message": "Application ID required"}), 400
        
        # Verify application exists and belongs to this HR user's job
        application = Applied_EMP.find_one({"_id": ObjectId(application_id)})
        if not application:
            return jsonify({"success": False, "message": "Application not found"}), 404
        
        # Verify the job belongs to this HR user
        job = JOBS.find_one({
            "_id": application.get('job_id'),
            "created_by": ObjectId(session['user_id'])
        })
        if not job:
            return jsonify({"success": False, "message": "Access denied"}), 403
        
        # Prepare update data
        update_data = {
            "status": data.get('status'),
            "interview_date": datetime.fromisoformat(data.get('interview_date')) if data.get('interview_date') else None,
            "interview_time": data.get('interview_time'),
            "interview_type": data.get('interview_type'),
            "interview_location": data.get('interview_location'),
            "interviewer_name": data.get('interviewer_name'),
            "interviewer_email": data.get('interviewer_email'),
            "meeting_link": data.get('meeting_link'),
            "interview_notes": data.get('interview_notes'),
            "interview_updated_at": datetime.now(),
            "interview_updated_by": session.get('user_id')
        }
        
        # Remove None values
        update_data = {k: v for k, v in update_data.items() if v is not None}
        
        # Update application
        result = Applied_EMP.update_one(
            {"_id": ObjectId(application_id)},
            {"$set": update_data}
        )
        
        if result.modified_count > 0:
            logger.info(f"Interview updated for application {application_id}")
            return jsonify({"success": True, "message": "Interview updated successfully"})
        else:
            return jsonify({"success": False, "message": "No changes made"}), 200
            
    except Exception as e:
        logger.error(f"Error updating interview: {e}")
        return jsonify({"success": False, "message": f"Error updating interview: {str(e)}"}), 500


# Also add a route to move candidates from shortlisted to interview_pending
@app.route("/move_to_interview/<application_id>", methods=["POST"])
def move_to_interview(application_id):
    """Move shortlisted candidate to interview_pending status"""
    if 'user_id' not in session or session.get('role') != 'employer':
        return jsonify({"error": "Unauthorized"}), 403

    try:
        # Verify the application belongs to this HR user's job
        application = Applied_EMP.find_one({"_id": ObjectId(application_id)})
        if not application:
            return jsonify({"error": "Application not found"}), 404
        
        job = JOBS.find_one({
            "_id": application.get('job_id'),
            "created_by": ObjectId(session['user_id'])
        })
        if not job:
            return jsonify({"error": "Access denied"}), 403
        
        result = Applied_EMP.update_one(
            {"_id": ObjectId(application_id), "status": "shortlisted"},
            {"$set": {
                "status": "interview_pending",
                "moved_to_interview_at": datetime.now(),
                "moved_by": session.get('user_id')
            }}
        )
        
        if result.modified_count > 0:
            return jsonify({
                "success": True, 
                "message": "Candidate moved to interview phase",
                "new_status": "interview_pending"
            })
        else:
            return jsonify({"error": "Candidate not found or not shortlisted"}), 404
            
    except Exception as e:
        logger.error(f"Error moving candidate to interview: {e}")
        return jsonify({"error": "Something went wrong"}), 500












@app.route('/debug_create_notifications')
def debug_create_notifications():
    """Debug route to create test notifications - REMOVE AFTER TESTING"""
    try:
        if 'user_id' not in session:
            return "Please log in first"
        
        user_id = session['user_id']
        
        # Create test notifications
        test_notifications = [
            {
                "user_id": ObjectId(user_id),
                "title": "Test Application Notification",
                "message": "This is a test notification for application updates.",
                "type": "application_submitted",
                "related_id": None,
                "action_url": "/HR1/show_job",
                "is_read": False,
                "created_at": datetime.now(),
                "is_active": True
            },
            {
                "user_id": ObjectId(user_id),
                "title": "Test Interview Notification", 
                "message": "This is a test notification for interview scheduling.",
                "type": "interview_scheduled",
                "related_id": None,
                "action_url": "/interviews",
                "is_read": False,
                "created_at": datetime.now() - timedelta(hours=2),
                "is_active": True
            },
            {
                "user_id": ObjectId(user_id),
                "title": "Test Status Update",
                "message": "This is a test notification for status changes.",
                "type": "status_update", 
                "related_id": None,
                "action_url": None,
                "is_read": True,
                "created_at": datetime.now() - timedelta(days=1),
                "is_active": True
            }
        ]
        
        # Insert test notifications
        result = NOTIFICATIONS.insert_many(test_notifications)
        
        return f"Created {len(result.inserted_ids)} test notifications. Now visit /notifications"
        
    except Exception as e:
        return f"Error creating test notifications: {e}"

        
        # Add this route to mark all notifications as read
@app.route('/mark_all_notifications_read', methods=['POST'])
def api_mark_all_notifications_read():
    """API endpoint to mark all notifications as read"""
    try:
        if 'user_id' not in session:
            return jsonify({"success": False, "message": "Not logged in"}), 401
        
        count = mark_all_notifications_read(session['user_id'])
        
        return jsonify({"success": True, "count": count})
            
    except Exception as e:
        logger.error(f"Error marking all notifications as read: {e}")
        return jsonify({"success": False, "message": "Server error"}), 500
    

@app.route('/notifications')
def notifications_page():
    """Notifications page for both applicants and HR"""
    try:
        if 'user_id' not in session:
            flash("Please log in first", "error")
            return redirect(url_for('loginpage'))
        
        user_id = session['user_id']
        user_role = session.get('role')
        
        # Get all notifications for user
        notifications = get_user_notifications(user_id, limit=50)
        
        # Group notifications by date
        grouped_notifications = {}
        for notification in notifications:
            date_key = notification['created_at'].strftime('%Y-%m-%d')
            date_label = notification['created_at'].strftime('%B %d, %Y')
            
            # Check if it's today or yesterday
            today = datetime.now().date()
            notif_date = notification['created_at'].date()
            
            if notif_date == today:
                date_label = "Today"
            elif notif_date == today - timedelta(days=1):
                date_label = "Yesterday"
            
            if date_key not in grouped_notifications:
                grouped_notifications[date_key] = {
                    'label': date_label,
                    'notifications': []
                }
            
            # Add time ago and icon
            notification['time_ago'] = get_time_ago(notification['created_at'])
            notification['icon'] = get_notification_icon(notification['type'])
            notification['color'] = get_notification_color(notification['type'])
            
            grouped_notifications[date_key]['notifications'].append(notification)
        
        # Sort grouped notifications by date (newest first)
        sorted_groups = sorted(grouped_notifications.items(), reverse=True)
        
        # Mark all as read when user visits the page
        mark_all_notifications_read(user_id)
        
        return render_template("notifications.html",
                             grouped_notifications=sorted_groups,
                             user_role=user_role)
                             
    except Exception as e:
        logger.error(f"Error in notifications page: {e}")
        return render_template("notifications.html",
                             grouped_notifications=[],
                             user_role=session.get('role'),
                             errorMsg="Error loading notifications")



# Add a new API endpoint for getting semantic rankings (add this as a new route)
@app.route('/api/semantic_ranking/<job_id>')
def api_semantic_ranking(job_id):
    """API endpoint to get semantic ranking for a specific job"""
    try:
        if 'user_id' not in session or session.get('role') != 'employer':
            return jsonify({"error": "Access denied"}), 403
        
        # Verify job belongs to current HR user
        job = JOBS.find_one({
            "_id": ObjectId(job_id),
            "created_by": ObjectId(session['user_id'])
        })
        
        if not job:
            return jsonify({"error": "Job not found or access denied"}), 404
        
        mongo_collections = {
            'JOBS': JOBS,
            'Applied_EMP': Applied_EMP,
            'resumeFetchedData': resumeFetchedData
        }
        
        ranked_candidates = get_enhanced_candidate_ranking(job_id, mongo_collections)
        
        # Format for API response
        response_data = {
            "job_id": job_id,
            "job_title": job.get('Job_Profile', 'Unknown Job'),
            "total_candidates": len(ranked_candidates),
            "semantic_enhanced": len([c for c in ranked_candidates if c['semantic_score'] > 0]),
            "candidates": ranked_candidates,
            "ranking_method": "semantic_enhanced" if any(c['semantic_score'] > 0 for c in ranked_candidates) else "traditional",
            "timestamp": datetime.now().isoformat()
        }
        
        return jsonify(response_data)
        
    except Exception as e:
        logger.error(f"Error in semantic ranking API: {e}")
        return jsonify({"error": "Failed to get semantic ranking"}), 500


# Add a batch update endpoint for existing applications
@app.route('/api/update_semantic_scores', methods=['POST'])
def api_update_semantic_scores():
    """Batch update existing applications with semantic scores"""
    try:
        if 'user_id' not in session or session.get('role') != 'employer':
            return jsonify({"error": "Access denied"}), 403
        
        data = request.get_json()
        job_id = data.get('job_id') if data else None
        
        mongo_collections = {
            'JOBS': JOBS,
            'Applied_EMP': Applied_EMP,
            'resumeFetchedData': resumeFetchedData
        }
        
        ranker = get_semantic_ranker()
        if not ranker.is_available():
            return jsonify({"error": "Semantic ranking not available"}), 503
        
        updated_count = ranker.batch_update_applications_with_semantic_scores(
            mongo_collections, job_id
        )
        
        return jsonify({
            "success": True,
            "updated_count": updated_count,
            "job_id": job_id,
            "message": f"Updated {updated_count} applications with semantic scores"
        })
        
    except Exception as e:
        logger.error(f"Error updating semantic scores: {e}")
        return jsonify({"error": "Failed to update semantic scores"}), 500












@app.route('/api/notifications/unread_count')
def api_unread_count():
    """API endpoint to get unread notification count"""
    try:
        if 'user_id' not in session:
            return jsonify({"count": 0})
        
        count = get_unread_count(session['user_id'])
        return jsonify({"count": count})
        
    except Exception as e:
        logger.error(f"Error getting unread count: {e}")
        return jsonify({"count": 0})

@app.route('/notification/<notification_id>')
def notification_detail(notification_id):
    """View notification detail and mark as read"""
    try:
        if 'user_id' not in session:
            flash("Please log in first", "error")
            return redirect(url_for('loginpage'))
        
        user_id = session['user_id']
        
        # Get notification
        notification = NOTIFICATIONS.find_one({
            "_id": ObjectId(notification_id),
            "user_id": ObjectId(user_id)
        })
        
        if not notification:
            flash("Notification not found", "error")
            return redirect(url_for('notifications_page'))
        
        # Mark as read
        mark_notification_read(notification_id, user_id)
        
        # If notification has an action URL, redirect there
        if notification.get('action_url'):
            return redirect(notification['action_url'])
        else:
            # Otherwise, go to notifications page
            flash("Notification viewed", "info")
            return redirect(url_for('notifications_page'))
            
    except Exception as e:
        logger.error(f"Error viewing notification detail: {e}")
        flash("Error viewing notification", "error")
        return redirect(url_for('notifications_page'))

@app.route('/mark_notification_read/<notification_id>', methods=['POST'])
def api_mark_notification_read(notification_id):
    """API endpoint to mark notification as read"""
    try:
        if 'user_id' not in session:
            return jsonify({"success": False, "message": "Not logged in"}), 401
        
        success = mark_notification_read(notification_id, session['user_id'])
        
        if success:
            return jsonify({"success": True})
        else:
            return jsonify({"success": False, "message": "Notification not found"}), 404
            
    except Exception as e:
        logger.error(f"Error marking notification as read: {e}")
        return jsonify({"success": False, "message": "Server error"}), 500

def get_notification_icon(notification_type):
    """Get icon for notification type"""
    icons = {
        'application_submitted': '📄',
        'new_application': '👤',
        'interview_scheduled': '📅',
        'status_update': '🔄',
        'job_update': '💼',
        'system': '⚙️',
        'hired': '🎉',
        'rejected': '📋'
    }
    return icons.get(notification_type, '📢')

def get_notification_color(notification_type):
    """Get color for notification type"""
    colors = {
        'application_submitted': '#10b981',
        'new_application': '#3b82f6',
        'interview_scheduled': '#f59e0b',
        'status_update': '#6b7280',
        'job_update': '#8b5cf6',
        'system': '#6b7280',
        'hired': '#10b981',
        'rejected': '#ef4444'
    }
    return colors.get(notification_type, '#6b7280')




# Add this route to your app.py - My Applications Page

@app.route('/my_applications')
def my_applications():
    """My Applications page for applicants to track their job applications"""
    try:
        if 'user_id' not in session or session.get('role') != 'applicant':
            flash("Access denied. Applicant login required.", "error")
            return redirect(url_for('loginpage'))
        
        user_id = session['user_id']
        
        # Get filter parameters
        status_filter = request.args.get('status', 'all')
        sort_by = request.args.get('sort', 'recent')
        
        # Get all applications for this user
        try:
            applications_query = {"user_id": ObjectId(user_id)}
            
            # Apply status filter
            if status_filter != 'all':
                if status_filter == 'active':
                    applications_query["status"] = {"$in": ["pending", "under_review", "shortlisted", "interview_pending"]}
                elif status_filter == 'interview':
                    applications_query["status"] = {"$in": ["interview_scheduled", "interviewed"]}
                elif status_filter == 'completed':
                    applications_query["status"] = {"$in": ["hired", "not_selected", "rejected"]}
                else:
                    applications_query["status"] = status_filter
            
            # Get applications
            applications = list(Applied_EMP.find(applications_query))
            
            # Sort applications
            if sort_by == 'recent':
                applications.sort(key=lambda x: x.get('applied_at', datetime.min), reverse=True)
            elif sort_by == 'oldest':
                applications.sort(key=lambda x: x.get('applied_at', datetime.min))
            elif sort_by == 'score':
                applications.sort(key=lambda x: x.get('Matching_percentage', 0), reverse=True)
            elif sort_by == 'company':
                pass
            
        except Exception as e:
            logger.error(f"Error fetching applications for user {user_id}: {e}")
            applications = []
        
        # Process applications with job details
        processed_applications = []
        for app in applications:
            try:
                # Get job details
                job = JOBS.find_one({"_id": app.get('job_id')}, {
                    "Job_Profile": 1, "CompanyName": 1, "Job_Description": 1, 
                    "Salary": 1, "Location": 1, "Status": 1, "LastDate": 1,
                    "CreatedAt": 1
                })
                
                if not job:
                    continue
                
                # Get resume details
                resume = resumeFetchedData.find_one({"_id": app.get('resume_id')}, {
                    "ResumeTitle": 1, "Name": 1
                })
                
                # Calculate days since application
                applied_date = app.get('applied_at', datetime.now())
                days_ago = (datetime.now() - applied_date).days
                
                # Determine status info
                status = app.get('status', 'pending')
                status_info = get_application_status_info(status)
                
                # Calculate match score safely
                # Calculate match score safely
                match_score = get_match_score(app)  # ✅ Fixed

                # Get component scores from application
                component_scores = app.get('component_scores') or \
                      (app.get('Matching_percentage', {}).get('component_scores'))
    
                app['component_scores'] = component_scores  # Ensure it's at root level
                app['match_score'] = app.get('combined_score') or \
                                    app.get('Matching_percentage', {}).get('overall_score', 0)

                try:
                    match_score = float(match_score)
                    if match_score <= 1.0:
                        match_score *= 100
                except (ValueError, TypeError):
                    match_score = 0.0

                # Process component scores
                processed_component_scores = {}
                for component, score in component_scores.items():
                    try:
                        score_val = float(score) if score is not None else 0.0
                        if score_val <= 1.0:
                            score_val *= 100
                        processed_component_scores[component] = round(score_val, 1)
                    except (ValueError, TypeError):
                        processed_component_scores[component] = 0.0
                
                # Check if job is still active
                job_active = job.get('Status', '').lower() == 'open'
                job_deadline = job.get('LastDate')
                deadline_passed = False
                if job_deadline and isinstance(job_deadline, datetime):
                    deadline_passed = job_deadline < datetime.now()
                
                # Determine next action for applicant
                next_action = get_next_action_for_applicant(status, job_active, deadline_passed)
                
                processed_app = {
                    'application_id': str(app.get('_id')),
                    'job_id': str(app.get('job_id')),
                    'job_title': job.get('Job_Profile', 'Unknown Job'),
                    'company_name': job.get('CompanyName', 'Unknown Company'),
                    'job_description': job.get('Job_Description', '')[:150] + "..." if job.get('Job_Description') else '',
                    'salary': job.get('Salary'),
                    'location': job.get('Location'),
                    'applied_date': applied_date,
                    'days_ago': days_ago,
                    'status': status,
                    'status_info': status_info,
                    'match_score': round(match_score, 1),
                    'component_scores': processed_component_scores,  # NEW: Add component scores
                    'resume_title': resume.get('ResumeTitle', 'Resume') if resume else 'Unknown Resume',
                    'job_active': job_active,
                    'deadline_passed': deadline_passed,
                    'job_deadline': job_deadline,
                    'next_action': next_action,
                    'interview_date': app.get('interview_date'),
                    'interview_time': app.get('interview_time'),
                    'semantic_score': app.get('semantic_similarity', 0),
                    'ai_enhanced': bool(app.get('enhanced_version')),
                    'can_withdraw': status in ['pending', 'under_review'] and job_active,
                    'candidate_feedback': app.get('Matching_percentage', {}).get('candidate_feedback') if isinstance(app.get('Matching_percentage'), dict) else None
                }
                
                processed_applications.append(processed_app)
                
            except Exception as e:
                logger.error(f"Error processing application {app.get('_id')}: {e}")
                continue
        
        # Sort by company name if requested
        if sort_by == 'company':
            processed_applications.sort(key=lambda x: x['company_name'])
        
        # Calculate statistics
        stats = calculate_application_stats(processed_applications)
        
        # Group applications by status
        grouped_applications = {
            'active': [app for app in processed_applications if app['status'] in ['pending', 'under_review', 'shortlisted', 'interview_pending']],
            'interview': [app for app in processed_applications if app['status'] in ['interview_scheduled', 'interviewed']],
            'completed': [app for app in processed_applications if app['status'] in ['hired', 'not_selected', 'rejected']],
            'all': processed_applications
        }
        
        logger.info(f"My Applications loaded for user {user_id}: {len(processed_applications)} applications")
        
        return render_template("my_applications.html",
                             applications=processed_applications,
                             grouped_applications=grouped_applications,
                             stats=stats,
                             status_filter=status_filter,
                             sort_by=sort_by,
                             available_statuses=get_available_status_filters())
                             
    except Exception as e:
        logger.error(f"Error in my_applications: {e}")
        return render_template("my_applications.html",
                             applications=[],
                             grouped_applications={'active': [], 'interview': [], 'completed': [], 'all': []},
                             stats={},
                             status_filter='all',
                             sort_by='recent',
                             available_statuses=[],
                             errorMsg="An error occurred loading your applications")

@app.route('/application_status/<job_id>')
def application_status(job_id):
    """View application status for a specific job"""
    try:
        if 'user_id' not in session:
            flash("Please log in first", "error")
            return redirect(url_for('loginpage'))
        
        user_id = session['user_id']
        
        logger.info(f"Loading application status for user={user_id}, job={job_id}")
        
        # Get the job details - using mongo.db.JOBS
        job = mongo.db.JOBS.find_one({"_id": ObjectId(job_id)})
        if not job:
            logger.error(f"Job not found: {job_id}")
            flash("Job not found", "error")
            return redirect(url_for('notifications_page'))
        
        logger.info(f"Job found: {job.get('Job_Profile', 'Unknown')}")
        
        # Find the application in Applied_EMP collection
        application = mongo.db.Applied_EMP.find_one({
            "user_id": ObjectId(user_id),
            "job_id": ObjectId(job_id)
        })
        
        if not application:
            logger.error(f"Application not found for user={user_id}, job={job_id}")
            # Try to find ANY application for this user to debug
            any_app = mongo.db.Applied_EMP.find_one({"user_id": ObjectId(user_id)})
            if any_app:
                logger.info(f"Found other applications for user. Sample job_id: {any_app.get('job_id')}")
            else:
                logger.info(f"No applications found for user {user_id}")
            flash("Application not found", "error")
            return redirect(url_for('notifications_page'))
        
        logger.info(f"Application found with status: {application.get('status', 'unknown')}")
        
        # Get status information
        current_status = application.get('status', 'pending')
        status_info = get_application_status_info(current_status)
        
        # Get company name from job (matching your schema)
        company_name = job.get('CompanyName', 'Company')
        
        # Get resume details
        resume = mongo.db.resumeFetchedData.find_one(
            {"_id": application.get('resume_id')}, 
            {"ResumeTitle": 1, "Name": 1}
        )
        
        logger.info(f"Resume found: {resume.get('ResumeTitle', 'N/A') if resume else 'Not found'}")
        
        # Get interview details if available
        interview = None
        if current_status in ['interview_scheduled', 'interviewed']:
            # Check if interview data is in the application itself
            if application.get('interview_date'):
                interview = {
                    'interview_date': application.get('interview_date'),
                    'interview_time': application.get('interview_time'),
                    'interview_type': application.get('interview_type'),
                    'notes': application.get('interview_notes')
                }
                logger.info(f"Interview data found in application")
            else:
                # Try INTERVIEWS collection as fallback
                try:
                    interview = mongo.db.INTERVIEWS.find_one({
                        "applicant_id": ObjectId(user_id),
                        "job_id": ObjectId(job_id)
                    })
                    if interview:
                        logger.info(f"Interview data found in INTERVIEWS collection")
                except Exception as e:
                    logger.warning(f"Could not check INTERVIEWS collection: {e}")
        
        # Format application date
        applied_date = application.get('applied_at', datetime.now())
        if isinstance(applied_date, datetime):
            applied_date_str = applied_date.strftime('%B %d, %Y')
        else:
            applied_date_str = str(applied_date)
        
        # # Calculate match score
        # match_score = application.get('Matching_percentage', 0)
        # if isinstance(match_score, dict):
        #     match_score = match_score.get('overall_score', 0)
        # try:
        #     match_score = float(match_score)
        #     if match_score <= 1.0:
        #         match_score *= 100
        # except (ValueError, TypeError):
        #     match_score = 0.0
        
        # logger.info(f"Match score calculated: {match_score}%")
        # Calculate match score
        match_score = get_match_score(application)

        logger.info(f"Match score calculated: {match_score}%")
        
        # Get job details matching your schema
        job_info = {
            'title': job.get('Job_Profile', 'Unknown Job'),
            'job_type': job.get('Job_Type', 'Not specified'),
            'location': job.get('Location', 'Not specified'),
            'salary': job.get('Salary', 'Not disclosed'),
            'description': job.get('Job_Description', ''),
            'status': job.get('Status', 'Unknown')
        }
        
        logger.info(f"Successfully prepared application status page for job: {job_info['title']}")
        
        return render_template('application_status.html',
                             job=job_info,
                             application=application,
                             status_info=status_info,
                             company_name=company_name,
                             interview=interview,
                             applied_date=applied_date_str,
                             match_score=round(match_score, 1),
                             resume_title=resume.get('ResumeTitle', 'Resume') if resume else 'Unknown Resume')
                             
    except Exception as e:
        logger.error(f"Error in application_status: {e}")
        import traceback
        logger.error(traceback.format_exc())
        flash("Error loading application status", "error")
        return redirect(url_for('notifications_page'))


def get_application_status_info(status):
    """Get user-friendly status information"""
    status_map = {
        'pending': {
            'label': 'Application Submitted',
            'description': 'Your application is being reviewed',
            'color': '#6b7280',
            'icon': 'clock',
            'progress': 20
        },
        'under_review': {
            'label': 'Under Review',
            'description': 'HR team is reviewing your application',
            'color': '#3b82f6',
            'icon': 'eye',
            'progress': 40
        },
        'shortlisted': {
            'label': 'Shortlisted',
            'description': 'You\'ve been selected for the next round',
            'color': '#10b981',
            'icon': 'check-circle',
            'progress': 60
        },
        'interview_pending': {
            'label': 'Interview Pending',
            'description': 'Interview scheduling in progress',
            'color': '#f59e0b',
            'icon': 'calendar-event',
            'progress': 70
        },
        'interview_scheduled': {
            'label': 'Interview Scheduled',
            'description': 'Your interview has been scheduled',
            'color': '#f59e0b',
            'icon': 'calendar-check',
            'progress': 80
        },
        'interviewed': {
            'label': 'Interviewed',
            'description': 'Interview completed, awaiting decision',
            'color': '#8b5cf6',
            'icon': 'chat-dots',
            'progress': 90
        },
        'hired': {
            'label': 'Hired',
            'description': 'Congratulations! You got the job',
            'color': '#10b981',
            'icon': 'trophy',
            'progress': 100
        },
        'not_selected': {
            'label': 'Not Selected',
            'description': 'Application was not successful this time',
            'color': '#ef4444',
            'icon': 'x-circle',
            'progress': 100
        },
        'rejected': {
            'label': 'Not Selected',
            'description': 'Application was not successful this time',
            'color': '#ef4444',
            'icon': 'x-circle',
            'progress': 100
        }
    }
    
    return status_map.get(status, {
        'label': status.replace('_', ' ').title(),
        'description': 'Status update',
        'color': '#6b7280',
        'icon': 'info-circle',
        'progress': 50
    })

def get_next_action_for_applicant(status, job_active, deadline_passed):
    """Determine what the applicant should do next"""
    if status == 'hired':
        return {'action': 'celebrate', 'message': 'Congratulations on your new job!', 'button': None}
    elif status in ['not_selected', 'rejected']:
        return {'action': 'continue_search', 'message': 'Keep applying to similar positions', 'button': 'Find Similar Jobs'}
    elif status == 'interview_scheduled':
        return {'action': 'prepare', 'message': 'Prepare for your upcoming interview', 'button': 'Interview Tips'}
    elif status == 'interviewed':
        return {'action': 'wait', 'message': 'Follow up politely if no response in 1-2 weeks', 'button': None}
    elif status in ['pending', 'under_review', 'shortlisted', 'interview_pending']:
        if not job_active:
            return {'action': 'job_closed', 'message': 'This position has been closed', 'button': None}
        elif deadline_passed:
            return {'action': 'deadline_passed', 'message': 'Application deadline has passed', 'button': None}
        else:
            return {'action': 'wait', 'message': 'Sit tight, they\'ll be in touch soon', 'button': None}
    else:
        return {'action': 'unknown', 'message': 'Check back for updates', 'button': None}

def calculate_application_stats(applications):
    """Calculate application statistics"""
    if not applications:
        return {
            'total': 0, 'active': 0, 'interview_stage': 0, 'completed': 0,
            'success_rate': 0, 'avg_match_score': 0, 'response_rate': 0
        }
    
    total = len(applications)
    active = len([app for app in applications if app['status'] in ['pending', 'under_review', 'shortlisted', 'interview_pending']])
    interview_stage = len([app for app in applications if app['status'] in ['interview_scheduled', 'interviewed']])
    completed = len([app for app in applications if app['status'] in ['hired', 'not_selected', 'rejected']])
    hired = len([app for app in applications if app['status'] == 'hired'])
    
    # Calculate rates
    success_rate = (hired / completed * 100) if completed > 0 else 0
    avg_match_score = sum(app['match_score'] for app in applications) / total if total > 0 else 0
    response_rate = ((total - len([app for app in applications if app['status'] == 'pending'])) / total * 100) if total > 0 else 0
    
    return {
        'total': total,
        'active': active,
        'interview_stage': interview_stage,
        'completed': completed,
        'hired': hired,
        'success_rate': round(success_rate, 1),
        'avg_match_score': round(avg_match_score, 1),
        'response_rate': round(response_rate, 1)
    }

def get_available_status_filters():
    """Get available status filter options"""
    return [
        {'value': 'all', 'label': 'All Applications', 'count_key': 'total'},
        {'value': 'active', 'label': 'Active', 'count_key': 'active'},
        {'value': 'interview', 'label': 'Interview Stage', 'count_key': 'interview_stage'},
        {'value': 'completed', 'label': 'Completed', 'count_key': 'completed'},
        {'value': 'pending', 'label': 'Pending Review', 'count_key': None},
        {'value': 'shortlisted', 'label': 'Shortlisted', 'count_key': None},
        {'value': 'hired', 'label': 'Hired', 'count_key': 'hired'},
    ]

@app.route('/withdraw_application/<application_id>', methods=['POST'])
def withdraw_application(application_id):
    """Allow applicant to withdraw their application"""
    try:
        if 'user_id' not in session or session.get('role') != 'applicant':
            return jsonify({"success": False, "message": "Access denied"}), 403
        
        user_id = session['user_id']
        
        # Verify this application belongs to the current user
        application = Applied_EMP.find_one({
            "_id": ObjectId(application_id),
            "user_id": ObjectId(user_id)
        })
        
        if not application:
            return jsonify({"success": False, "message": "Application not found"}), 404
        
        # Check if withdrawal is allowed
        current_status = application.get('status', 'pending')
        if current_status not in ['pending', 'under_review']:
            return jsonify({"success": False, "message": "Cannot withdraw application at this stage"}), 400
        
        # Update application status to withdrawn
        result = Applied_EMP.update_one(
            {"_id": ObjectId(application_id)},
            {"$set": {
                "status": "withdrawn",
                "withdrawn_at": datetime.now(),
                "withdrawn_by": ObjectId(user_id)
            }}
        )
        
        if result.modified_count > 0:
            # Get job details for notification
            job = JOBS.find_one({"_id": application['job_id']}, {"Job_Profile": 1})
            job_title = job.get('Job_Profile', 'Unknown Job') if job else 'Unknown Job'
            
            logger.info(f"Application {application_id} withdrawn by user {user_id}")
            return jsonify({
                "success": True, 
                "message": f"Application for {job_title} has been withdrawn"
            })
        else:
            return jsonify({"success": False, "message": "Failed to withdraw application"}), 500
            
    except Exception as e:
        logger.error(f"Error withdrawing application {application_id}: {e}")
        return jsonify({"success": False, "message": "Error withdrawing application"}), 500

@app.route('/reapply_job/<job_id>', methods=['GET'])
def reapply_job(job_id):
    """Redirect to job application page for reapplying"""
    try:
        if 'user_id' not in session or session.get('role') != 'applicant':
            flash("Access denied. Applicant login required.", "error")
            return redirect(url_for('loginpage'))
        
        # Verify job exists and is still open
        job = JOBS.find_one({"_id": ObjectId(job_id)}, {"Status": 1, "Job_Profile": 1})
        
        if not job:
            flash("Job not found", "error")
            return redirect(url_for('my_applications'))
        
        if job.get('Status', '').lower() != 'open':
            flash("This job is no longer accepting applications", "warning")
            return redirect(url_for('my_applications'))
        
        # Redirect to job details page
        return redirect(f"/HR1/show_job?job_id={job_id}")
        
    except Exception as e:
        logger.error(f"Error in reapply_job: {e}")
        flash("Error accessing job", "error")
        return redirect(url_for('my_applications'))












def _generate_next_steps(avg_score, interview_rate, weekly_rate):
    """Generate specific next steps based on performance metrics"""
    steps = []
    
    if avg_score < 60:
        steps.append("Optimize resume keywords to match job requirements more closely")
        steps.append("Target roles that better align with your current skill set")
    elif avg_score >= 80:
        steps.append("Apply for senior-level positions with 10-15% higher requirements")
        steps.append("Focus on 3-5 premium opportunities per week")
    
    if interview_rate < 5:
        steps.append("Personalize cover letters to show specific company knowledge")
        steps.append("Apply within 24-48 hours of job posting for higher visibility")
    
    if weekly_rate > 15:
        steps.append("Reduce application volume to 5-8 per week for better targeting")
    elif weekly_rate < 3:
        steps.append("Increase application frequency to 5-7 per week")
    
    if not steps:
        steps = ["Continue current strategy - performance metrics are strong"]
    
    return steps

def _assess_performance_level(avg_score, interview_rate, shortlist_rate):
    """Assess overall performance level"""
    if avg_score >= 80 and interview_rate >= 15:
        return "excellent"
    elif avg_score >= 70 and interview_rate >= 10:
        return "strong"
    elif avg_score >= 60 and interview_rate >= 5:
        return "good"
    elif avg_score >= 50:
        return "developing"
    else:
        return "needs_optimization"





# @app.route('/api/applicant_insights')
# def api_applicant_insights():
#     """Enhanced API endpoint with legitimate predictive insights"""
#     if 'user_id' not in session or session.get('role') != 'applicant':
#         return jsonify({"error": "Access denied"}), 403
    
#     insights = get_legitimate_applicant_predictive_insights(session['user_id'])
#     return jsonify(insights)

# @app.route('/api/hr_insights/<job_id>')
# def api_hr_insights(job_id):
#     """Enhanced API endpoint with legitimate predictive insights"""
#     if 'user_id' not in session or session.get('role') != 'employer':
#         return jsonify({"error": "Access denied"}), 403
    
#     insights = get_legitimate_hr_predictive_insights(job_id)
#     return jsonify(insights)


@app.route('/debug_interviews')
def debug_interviews():
    """Temporary debug route - remove after fixing"""
    if 'user_id' not in session:
        return "Not logged in"
    
    user_id = session['user_id']
    
    # Get all applications for this user
    applications = list(Applied_EMP.find({"user_id": ObjectId(user_id)}))
    
    debug_info = []
    for app in applications:
        job = JOBS.find_one({"_id": app.get('job_id')}, {"Job_Profile": 1})
        debug_info.append({
            'job_title': job.get('Job_Profile', 'Unknown') if job else 'Unknown',
            'status': app.get('status'),
            'has_interview_date': bool(app.get('interview_date')),
            'application_id': str(app.get('_id'))
        })
    
    return f"<pre>{debug_info}</pre>"








# Email configuration - add these to your app.py
EMAIL_CONFIG = {
    'SMTP_SERVER': 'smtp.gmail.com',  # Change to your email provider
    'SMTP_PORT': 587,
    'EMAIL_ADDRESS': 'oueiaorbe0501@gmail.com',  # Your app's email
    'EMAIL_PASSWORD': 'mgqu ovkm carn zvxj',  # App password
    'FROM_NAME': 'HireArchy Team'
}

def send_verification_email(email, name, verification_token, user_type):
    """Send email verification link"""
    try:
        verification_url = f"{os.environ.get('BASE_URL', 'http://hirearchy.site')}/verify-email/{verification_token}"
        
        if user_type == 'applicant':
            subject = "Verify Your Job Seeker Account - HireArchy"
            body = f"""
Hi {name},

Welcome to HireArchy! Please verify your email address to activate your job seeker account.

Click the link below to verify your email:
{verification_url}

This link will expire in 24 hours.

Best regards,
The HireArchy Team
"""
        else:  # HR/employer
            subject = "Verify Your HR Account - HireArchy"
            body = f"""
Hi {name},

Thank you for registering as an HR professional on HireArchy!

Please verify your email address to activate your recruiter account:
{verification_url}

This link will expire in 24 hours.

Best regards,
The HireArchy Team
"""

        # Create email message
        msg = MIMEMultipart()
        msg['From'] = f"{EMAIL_CONFIG['FROM_NAME']} <{EMAIL_CONFIG['EMAIL_ADDRESS']}>"
        msg['To'] = email
        msg['Subject'] = subject
        
        msg.attach(MIMEText(body, 'plain'))

        # Send email
        server = smtplib.SMTP(EMAIL_CONFIG['SMTP_SERVER'], EMAIL_CONFIG['SMTP_PORT'])
        server.starttls()
        server.login(EMAIL_CONFIG['EMAIL_ADDRESS'], EMAIL_CONFIG['EMAIL_PASSWORD'])
        text = msg.as_string()
        server.sendmail(EMAIL_CONFIG['EMAIL_ADDRESS'], email, text)
        server.quit()
        
        logger.info(f"Verification email sent to {email}")
        return True
        
    except Exception as e:
        logger.error(f"Failed to send verification email to {email}: {e}")
        return False
       


def validate_enhanced_applicant_data(form_data):
    """Validate applicant registration data"""
    errors = []
    
    full_name = form_data.get('full_name', '').strip()
    email = form_data.get('email', '').strip()
    password = form_data.get('password', '')
    confirm_password = form_data.get('confirm_password', '')
    
    # Required field validation
    if not full_name or len(full_name) < 2:
        errors.append("Full name must be at least 2 characters long")
    elif len(full_name) > 100:
        errors.append("Full name too long (max 100 characters)")
    
    if not validate_email(email):
        errors.append("Invalid email format")
    
    if password != confirm_password:
        errors.append("Passwords do not match")
    
    is_valid_password, password_msg = validate_password(password)
    if not is_valid_password:
        errors.append(password_msg)
    
    # Optional field validation
    phone = form_data.get('phone_number', '').strip()
    if phone and len(phone) < 10:
        errors.append("Phone number too short")
    
    linkedin = form_data.get('linkedin_profile', '').strip()
    if linkedin and 'linkedin.com' not in linkedin.lower():
        errors.append("Please enter a valid LinkedIn URL")
    
    return errors

def validate_enhanced_hr_data(form_data):
    """Validate HR registration data"""
    errors = []
    
    # Required fields
    full_name = form_data.get('full_name', '').strip()
    work_email = form_data.get('work_email', '').strip()
    job_title = form_data.get('job_title', '').strip()
    company_name = form_data.get('company_name', '').strip()
    company_size = form_data.get('company_size', '').strip()
    industry = form_data.get('industry', '').strip()
    phone_number = form_data.get('phone_number', '').strip()
    password = form_data.get('password', '')
    confirm_password = form_data.get('confirm_password', '')
    
    # Validate required fields
    if not full_name or len(full_name) < 2:
        errors.append("Full name must be at least 2 characters long")
    elif len(full_name) > 100:
        errors.append("Full name too long")
    
    if not validate_email(work_email):
        errors.append("Invalid work email format")
    
    if not job_title or len(job_title) < 2:
        errors.append("Job title is required")
    elif len(job_title) > 100:
        errors.append("Job title too long")
    
    if not company_name or len(company_name) < 2:
        errors.append("Company name is required")
    elif len(company_name) > 100:
        errors.append("Company name too long")
    
    if not company_size:
        errors.append("Company size is required")
    
    if not industry:
        errors.append("Industry is required")
    
    if not phone_number or len(phone_number) < 10:
        errors.append("Valid phone number is required")
    
    if password != confirm_password:
        errors.append("Passwords do not match")
    
    # Enhanced password validation for HR
    if len(password) < 8:
        errors.append("Password must be at least 8 characters long")
    if not re.search(r'[A-Za-z]', password):
        errors.append("Password must contain at least one letter")
    if not re.search(r'\d', password):
        errors.append("Password must contain at least one number")
    if not re.search(r'[!@#$%^&*(),.?":{}|<>]', password):
        errors.append("Password must contain at least one special character")
    
    # Optional field validation
    company_website = form_data.get('company_website', '').strip()
    if company_website and not (company_website.startswith('http://') or company_website.startswith('https://')):
        errors.append("Company website must start with http:// or https://")
    
    return errors

@app.route('/register_applicant')
def register_applicant():
    return render_template('register_applicant.html')

@app.route('/register_employer')  
def register_employer():
    """Redirect to existing signup with employer role"""
    return render_template('register_employer.html')

@app.route('/signup_applicant', methods=['POST'])
def signup_applicant():
    """Enhanced applicant registration with email verification"""
    try:
        # Validate input
        validation_errors = validate_enhanced_applicant_data(request.form)
        if validation_errors:
            for error in validation_errors:
                flash(error, 'error')
            return render_template('register_applicant.html')
        
        full_name = request.form.get('full_name', '').strip()
        email = request.form.get('email', '').strip()
        password = request.form.get('password', '')
        # phone_number = request.form.get('phone_number', '').strip()
        # location = request.form.get('location', '').strip()
        # current_title = request.form.get('current_title', '').strip()
        # experience_level = request.form.get('experience_level', '').strip()
        # industry_preference = request.form.get('industry_preference', '').strip()
        # linkedin_profile = request.form.get('linkedin_profile', '').strip()
        
        # Check if user already exists
        existing_user = IRS_USERS.find_one({'Email': email})
        if existing_user:
            flash('Email already registered', 'error')
            return render_template('register_applicant.html')
        
        # Generate verification token
        verification_token = secrets.token_urlsafe(32)
        
        # Create user document
        hashpass = generate_password_hash(password)
        user_data = {
            'Name': full_name,
            'Email': email,
            'Password': hashpass,
            'Role': 'applicant',
            'CreatedAt': datetime.now(),
            'EmailVerified': False,
            'VerificationToken': verification_token,
            'TokenExpiry': datetime.now() + timedelta(hours=24)
        }
        
        # Insert user
        result = IRS_USERS.insert_one(user_data)
        if result.inserted_id:
            # Send verification email
            email_sent = send_verification_email(email, full_name, verification_token, 'applicant')
            
            if email_sent:
                flash('Registration successful! Please check your email to verify your account.', 'success')
                logger.info(f"New applicant registered: {email}")
                return render_template('email_verification_sent.html', email=email, user_type='applicant')
            else:
                flash('Registration successful, but verification email failed to send. Please contact support.', 'warning')
                return redirect(url_for('loginpage'))
        else:
            flash('Registration failed. Please try again.', 'error')
            return render_template('register_applicant.html')
            
    except Exception as e:
        logger.error(f"Error in applicant registration: {e}")
        flash('Registration failed due to server error.', 'error')
        return render_template('register_applicant.html')

@app.route('/signup_hr', methods=['POST'])
def signup_hr():
    """Enhanced HR registration with email verification"""
    try:
        # Validate input
        validation_errors = validate_enhanced_hr_data(request.form)
        if validation_errors:
            for error in validation_errors:
                flash(error, 'error')
            return render_template('register_employer.html')
        
        full_name = request.form.get('full_name', '').strip()
        work_email = request.form.get('work_email', '').strip()
        job_title = request.form.get('job_title', '').strip()
        company_name = request.form.get('company_name', '').strip()
        company_website = request.form.get('company_website', '').strip()
        password = request.form.get('password', '')
        # company_size = request.form.get('company_size', '').strip()
        # industry = request.form.get('industry', '').strip()
        # phone_number = request.form.get('phone_number', '').strip()
       
       
        # hiring_volume = request.form.get('hiring_volume', '').strip()
        
        # Check if user already exists
        existing_user = IRS_USERS.find_one({'Email': work_email})
        if existing_user:
            flash('Email already registered', 'error')
            return render_template('register_employer.html')
        
        # Generate verification token
        verification_token = secrets.token_urlsafe(32)
        
        # Create user document
        hashpass = generate_password_hash(password)
        user_data = {
            'Name': full_name,
            'Email': work_email,
            'Password': hashpass,
            'Role': 'employer',
            'CreatedAt': datetime.now(),
            'EmailVerified': False,
            'VerificationToken': verification_token,
            'TokenExpiry': datetime.now() + timedelta(hours=24),
            # Enhanced HR fields
            'JobTitle': job_title,
            'CompanyName': company_name,
            # 'CompanySize': company_size
        }
        
        # Insert user
        result = IRS_USERS.insert_one(user_data)
        if result.inserted_id:
            # Send verification email
            email_sent = send_verification_email(work_email, full_name, verification_token, 'employer')
            
            if email_sent:
                flash('Registration successful! Please check your email to verify your account.', 'success')
                logger.info(f"New HR user registered: {work_email} at {company_name}")
                return render_template('email_verification_sent.html', email=work_email, user_type='employer')
            else:
                flash('Registration successful, but verification email failed to send. Please contact support.', 'warning')
                return redirect(url_for('loginpage'))
        else:
            flash('Registration failed. Please try again.', 'error')
            return render_template('register_employer.html')
            
    except Exception as e:
        logger.error(f"Error in HR registration: {e}")
        flash('Registration failed due to server error.', 'error')
        return render_template('register_employer.html')

@app.route('/verify-email/<token>')
def verify_email(token):
    """Email verification endpoint"""
    try:
        # Find user with this token
        user = IRS_USERS.find_one({
            'VerificationToken': token,
            'EmailVerified': False,
            'TokenExpiry': {'$gt': datetime.now()}
        })
        
        if not user:
            flash('Invalid or expired verification link.', 'error')
            return render_template('verification_failed.html')
        
        # Verify the user
        IRS_USERS.update_one(
            {'_id': user['_id']},
            {
                '$set': {'EmailVerified': True},
                '$unset': {'VerificationToken': '', 'TokenExpiry': ''}
            }
        )
        
        logger.info(f"Email verified for user: {user['Email']}")
        flash('Email verified successfully! You can now log in.', 'success')
        return render_template('verification_success.html', user_type=user['Role'])
        
    except Exception as e:
        logger.error(f"Error in email verification: {e}")
        flash('Verification failed due to server error.', 'error')
        return render_template('verification_failed.html')


@app.route('/resend-verification', methods=['GET', 'POST'])
def resend_verification():
    """Resend verification email - FIXED to support both GET and POST"""
    try:
        if request.method == 'POST':
            email = request.form.get('email', '').strip()
        else:
            # Support GET request with query parameter
            email = request.args.get('email', '').strip()
        
        if not email:
            flash('Email address is required.', 'error')
            return redirect(url_for('loginpage'))
        
        # Find unverified user
        user = IRS_USERS.find_one({
            'Email': email,
            'EmailVerified': False
        })
        
        if not user:
            flash('No unverified account found with this email.', 'error')
            return redirect(url_for('loginpage'))
        
        # Generate new token
        verification_token = secrets.token_urlsafe(32)
        
        # Update user with new token
        IRS_USERS.update_one(
            {'_id': user['_id']},
            {
                '$set': {
                    'VerificationToken': verification_token,
                    'TokenExpiry': datetime.now() + timedelta(hours=24)
                }
            }
        )
        
        # Send new verification email
        email_sent = send_verification_email(email, user['Name'], verification_token, user['Role'])
        
        if email_sent:
            flash('Verification email resent! Please check your inbox.', 'success')
            logger.info(f"Verification email resent to: {email}")
        else:
            flash('Failed to send verification email. Please try again later.', 'error')
            logger.error(f"Failed to send verification email to: {email}")
        
        return redirect(url_for('loginpage'))
        
    except Exception as e:
        logger.error(f"Error resending verification: {e}")
        import traceback
        logger.error(traceback.format_exc())
        flash('Failed to resend verification email.', 'error')
        return redirect(url_for('loginpage'))


# ✅ OPTIONAL: Add a dedicated resend page
@app.route('/resend-verification-page')
def resend_verification_page():
    """Page to resend verification email"""
    email = request.args.get('email', '')
    return render_template('resend_verification.html', email=email)

def is_personal_email_domain(email):
    """Check if email domain is personal"""
    personal_domains = {
        'gmail.com', 'yahoo.com', 'hotmail.com', 'outlook.com', 'aol.com',
        'icloud.com', 'protonmail.com', 'mail.com', 'yandex.com'
    }
    domain = email.split('@')[1].lower()
    return domain in personal_domains










# Error handlers
@app.errorhandler(404)
def not_found_error(error):
    logger.warning(f"404 error: {request.url}")
    return render_template('error.html', error_code=404, error_message="Page not found"), 404

@app.errorhandler(500)
def internal_error(error):
    logger.error(f"500 error: {error}")
    return render_template('error.html', error_code=500, error_message="Internal server error"), 500

@app.errorhandler(413)
def too_large(error):
    logger.warning("File too large uploaded")
    flash("File too large. Maximum size is 16MB.", "error")
    return redirect(url_for('emp'))

if __name__ == "__main__":
    try:
        import sys
        import os

        # Get configuration from environment
        debug_mode = os.environ.get('FLASK_DEBUG', 'True').lower() == 'true'
        host = os.environ.get('FLASK_HOST', '127.0.0.1')
        port = int(os.environ.get('FLASK_PORT', 5000))

        logger.info(f"Starting Flask app on {host}:{port} (debug={debug_mode})")

        # ✅ Disable automatic reloader so it never re-runs on file changes
        app.run(debug=debug_mode, host=host, port=port, use_reloader=False)

    except Exception as e:
        logger.error(f"Failed to start Flask app: {e}")
        sys.exit(1)
